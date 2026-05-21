"""
Generates a clean deployment and sharing guide for the assignment.
Output: deliverables/Dashboard_Deployment_Guide.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

os.makedirs("deliverables", exist_ok=True)

DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY       = RGBColor(0x40, 0x40, 0x40)
GREEN      = RGBColor(0x37, 0x56, 0x23)
GREEN_LT   = RGBColor(0xE2, 0xEF, 0xDA)

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)

def set_cell_border(cell, color="CCCCCC", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = MID_BLUE

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

def step(doc, number, title, detail):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run(f"Step {number}:  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY
    if detail:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.left_indent = Cm(1.2)
        r3 = p2.add_run(detail)
        r3.font.size = Pt(10.5)
        r3.font.color.rgb = GREY
        r3.italic = True

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK_BLUE
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF4FB')
    pPr.append(shd)

def two_col_table(doc, headers, rows, widths=(4.0, 12.0)):
    tbl = doc.add_table(rows=1+len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, (cell, h) in enumerate(zip(hdr, headers)):
        set_cell_bg(cell, DARK_BLUE)
        set_cell_border(cell, '1F497D')
        cell.width = Cm(widths[i])
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
    for i, row_data in enumerate(rows):
        row = tbl.rows[i+1].cells
        bg = LIGHT_BLUE if i % 2 == 0 else WHITE
        for j, (cell, val) in enumerate(zip(row, row_data)):
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            cell.width = Cm(widths[j])
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10); r.font.color.rgb = GREY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()

# ── Build document ─────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Cover line
p = doc.add_paragraph()
r = p.add_run("ESADE MSc Finance  —  Sustainable Finance Group Assignment")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = RGBColor(0x80,0x80,0x80)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(18)
r2 = p2.add_run("Deliverable: Interactive Dashboard — Deployment and Sharing Guide")
r2.font.size = Pt(9); r2.italic = True; r2.font.color.rgb = RGBColor(0x80,0x80,0x80)

# ── Section 1 ──────────────────────────────────────────────────
heading1(doc, "1.  Overview")
body(doc,
    "The ESADE Sustainable Finance pipeline produces an interactive web dashboard that visualises "
    "all agent outputs — portfolio holdings, ESG scores, financial metrics, climate data, "
    "greenwashing assessments, and exclusion rationale. The dashboard is built with Streamlit "
    "and deployed in two ways: locally for development and pipeline testing, and publicly via "
    "Streamlit Community Cloud for sharing with classmates and the assessment panel.")

# ── Section 2 ──────────────────────────────────────────────────
heading1(doc, "2.  Dashboard Pages")
body(doc, "The dashboard contains eight interactive pages accessible from the left sidebar:")

two_col_table(doc,
    headers=["Page", "Content"],
    rows=[
        ("Overview",             "Key portfolio metrics, weights donut chart, ESG vs Sharpe scatter, sector breakdown"),
        ("Portfolio Holdings",   "Sortable table of all 20 holdings with E/S/G breakdown bar chart"),
        ("ESG Scores",           "Full universe ranking (57 companies), pillar averages, distribution histogram"),
        ("Risk & Returns",       "Risk-return scatter, Sharpe ratio bar chart, max drawdown per holding"),
        ("Climate & Biodiversity","Carbon intensity, nature risk score (ENCORE + WRI Aqueduct), renewable energy %"),
        ("Greenwashing",         "8-Test framework results per company — auto-populates when RAG JSONs are added"),
        ("Exclusions",           "Why each of the 19 companies was excluded, with ESG score comparison"),
        ("Mandate",              "Fund details, ESG pillar weights, hard exclusion rules, investment thesis"),
    ],
    widths=(4.5, 11.5)
)

# ── Section 3 ──────────────────────────────────────────────────
heading1(doc, "3.  Running the Dashboard Locally")
body(doc,
    "The local dashboard is used for development, pipeline testing, and live presentation. "
    "It reads directly from the outputs/ folder, so every time the pipeline is re-run the "
    "dashboard reflects the latest data on browser refresh.")

heading2(doc, "Prerequisites")
body(doc, "Python 3.11 installed. Project set up via setup.bat (one-time only).")

heading2(doc, "Launch")
step(doc, 1, "Open the project in VS Code", None)
step(doc, 2, "Press Ctrl+Shift+B", "Runs the full 13-agent pipeline (~63 seconds)")
step(doc, 3, "Open a terminal (Ctrl + `)", None)
step(doc, 4, "Run the dashboard", None)
code(doc, 'venv\\Scripts\\streamlit.exe run app.py --server.headless=true')
step(doc, 5, "Open browser at http://localhost:8501", None)

body(doc, "Alternatively, double-click launch_dashboard.bat in the project folder.")

# ── Section 4 ──────────────────────────────────────────────────
heading1(doc, "4.  Public Deployment via Streamlit Community Cloud")
body(doc,
    "The dashboard is deployed publicly so classmates and the assessment panel can access "
    "it from any device without installing Python. Streamlit Community Cloud (share.streamlit.io) "
    "hosts the app for free and automatically redeploys whenever the GitHub repository is updated.")

heading2(doc, "One-Time Setup")
step(doc, 1, "Push project to GitHub",
     "Repository: github.com/IVasileiadis00/esade-sf  (public)")
step(doc, 2, "Go to share.streamlit.io", "Sign in with GitHub account")
step(doc, 3, "Click Deploy an app → From GitHub", None)
step(doc, 4, "Fill in the deployment form", None)
code(doc, "Repository:      IVasileiadis00/esade-sf")
code(doc, "Branch:          main")
code(doc, "Main file path:  app.py")
step(doc, 5, "Click Deploy", "App builds automatically in ~2 minutes")

heading2(doc, "Updating the Live Dashboard")
body(doc,
    "After running the pipeline with new data, push to GitHub and the live URL updates automatically:")
code(doc, "git add .")
code(doc, "git push")

# ── Section 5 ──────────────────────────────────────────────────
heading1(doc, "5.  Sharing with Classmates")
two_col_table(doc,
    headers=["Method", "How"],
    rows=[
        ("Live URL (no install)",  "Share the Streamlit Cloud link — opens in any browser instantly"),
        ("GitHub repo",            "github.com/IVasileiadis00/esade-sf — download ZIP, run setup.bat"),
        ("Presentation day",       "Run locally on laptop, screen share — no internet dependency"),
    ],
    widths=(5.0, 11.0)
)

# ── Section 6 ──────────────────────────────────────────────────
heading1(doc, "6.  Data Update Process (Friday)")
body(doc,
    "When the professor's real CSV files are received, the dashboard updates in three steps:")
step(doc, 1, "Replace mock data",
     "Drop the 4 real CSV files into data/provided/, replacing the mock files")
step(doc, 2, "Re-run the pipeline",
     "Press Ctrl+Shift+B in VS Code — all 13 agents recalculate on real data (~63 seconds)")
step(doc, 3, "Push to GitHub",
     "git add .  then  git push — Streamlit Cloud redeploys automatically")
body(doc,
    "No code changes are required. The pipeline reads all files by column name, "
    "so the real data integrates without modification provided the column structure "
    "matches the data dictionary established in Agent 03.")

# ── Section 7 ──────────────────────────────────────────────────
heading1(doc, "7.  Technical Stack")
two_col_table(doc,
    headers=["Component", "Technology"],
    rows=[
        ("Dashboard framework",  "Streamlit 1.35+"),
        ("Charts",               "Plotly 5.18+"),
        ("Data processing",      "pandas 2.0+, numpy 1.26+"),
        ("Visualisations",       "matplotlib 3.7+"),
        ("Version control",      "Git + GitHub (github.com/IVasileiadis00/esade-sf)"),
        ("Cloud hosting",        "Streamlit Community Cloud (share.streamlit.io)"),
        ("Local execution",      "Python 3.11, VS Code, virtual environment (venv/)"),
    ],
    widths=(5.0, 11.0)
)

out = "deliverables/Dashboard_Deployment_Guide.docx"
doc.save(out)
print(f"Saved: {out}")
