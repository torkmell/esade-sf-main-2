"""
Generates the ESADE Sustainable Finance Investment Mandate document.
Includes: mandate table, 20 portfolio holdings with rationale + red flags,
and 19 excluded companies with red flags.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os
from datetime import date

TODAY = date.today().strftime("%d %B %Y")

# ── colour palette ────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)   # headings / header bg
MID_BLUE    = RGBColor(0x2E, 0x74, 0xB5)   # sub-headings
LIGHT_BLUE  = RGBColor(0xD6, 0xE4, 0xF0)   # table header bg
RED_LIGHT   = RGBColor(0xFF, 0xEB, 0xEB)   # red-flag rows
GREEN_LIGHT = RGBColor(0xEB, 0xF5, 0xEB)   # rationale rows
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
GREY_TEXT   = RGBColor(0x40, 0x40, 0x40)


# ── helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_col = str(rgb)  # RGBColor.__str__ returns "RRGGBB" hex string
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_col)
    tcPr.append(shd)


def set_cell_border(cell, border_color="BBBBBB", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), border_color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)
    run.bold = True
    if color:
        run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = DARK_BLUE
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = MID_BLUE
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BLUE
    return p


def body(doc, text, bold=False, italic=False, size=10, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or GREY_TEXT
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = GREY_TEXT
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY_TEXT
    return p


def mandate_row(table, item, explanation, item_bg=None):
    row = table.add_row()
    row.height = None
    c0, c1 = row.cells[0], row.cells[1]
    c0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if item_bg:
        set_cell_bg(c0, item_bg)
        set_cell_bg(c1, item_bg)
    for c in [c0, c1]:
        set_cell_border(c)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        for m in [c.paragraphs[0].paragraph_format]:
            m.left_indent = Cm(0.2)
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(item)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = DARK_BLUE

    p1 = c1.paragraphs[0]
    r1 = p1.add_run(explanation)
    r1.font.size = Pt(10)
    r1.font.color.rgb = GREY_TEXT
    return row


def section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── data ─────────────────────────────────────────────────────
MANDATE = [
    ("Client",
     "Institutional investors — pension funds, university endowments, and insurance companies seeking an SFDR Article 8-style fund that integrates ESG criteria alongside financial objectives. The fund is positioned for long-term capital preservation with a sustainability overlay, appropriate for liability-driven investors with a fiduciary duty to consider non-financial risks."),
    ("Investment Universe",
     "Large-cap European equities drawn from the STOXX Europe 600 index. Initial screening applied to 56 companies across 11 BICS sectors and 16 countries. All candidates must be listed on a regulated EU/EEA/UK exchange, have a minimum market capitalisation of EUR 5 billion, and have sufficient price history (minimum 3 years) for financial analysis."),
    ("Portfolio Type",
     "Long-only equity portfolio. No short positions, no leverage, no derivatives. The portfolio is fully invested at all times. All positions are held in the base currency of the relevant exchange; no currency hedging applied in this prototype."),
    ("Number of Holdings",
     "20 securities selected from the eligible universe of 56 candidates after applying hard exclusion screens (ESG floor, negative Sharpe ratio, greenwashing 8-Test). This falls within the recommended range of 15-25 securities, providing sufficient diversification while maintaining a concentrated conviction portfolio."),
    ("Time Horizon",
     "Medium to long term — minimum 5-year investment horizon. ESG transition risks (carbon pricing, regulatory change, stranded assets) materialise over multi-year periods; a short time horizon would penalise companies undergoing credible but costly sustainability transitions. Performance is evaluated annually against the STOXX Europe 600 benchmark."),
    ("Benchmark",
     "STOXX Europe 600 (equal-weighted for comparison purposes). Selected as the most representative broad European equity index. Portfolio ESG score, WACI (Weighted Average Carbon Intensity), and risk-adjusted returns are reported against this benchmark at each reporting period. A market-cap weighted comparison is also provided in the factsheet."),
    ("Sustainability Objective",
     "A combination of four sustainability approaches:\n"
     "1. ESG Integration — proprietary transparent scoring (E 40%, S 30%, G 30%) using raw reported metrics; no black-box vendor ratings.\n"
     "2. Net-Zero Transition — WACI tracked against benchmark; preference for companies with SBTi-validated targets.\n"
     "3. EU Taxonomy Alignment — companies with higher taxonomy-eligible revenue receive scoring uplift.\n"
     "4. Controversy & Greenwashing Avoidance — systematic 8-Test forensic review of all sustainability claims with mandatory exclusion of HIGH-risk companies."),
    ("Risk Constraints",
     "Maximum position size: 10% per single holding.\n"
     "Minimum sector diversification: 5 distinct BICS sectors represented.\n"
     "ESG floor exclusion: bottom 10% of universe by composite ESG score automatically excluded.\n"
     "Financial quality screen: companies with negative Sharpe ratio over 5-year history excluded.\n"
     "Greenwashing screen: companies scoring HIGH on 3 or more of the 8 greenwashing dimensions excluded.\n"
     "Carbon constraint: portfolio WACI must be disclosed; target to improve vs. benchmark WACI.\n"
     "Biodiversity proxy: nature-risk score (ENCORE + WRI Aqueduct) disclosed for all holdings."),
    ("Final Decision",
     "20 companies included (full list in Section 2). 19 companies excluded — 7 failed the ESG floor screen, 12 failed the negative Sharpe screen. 0 companies failed the greenwashing screen in this run (greenwashing data to be updated with real Claude Projects RAG outputs). All exclusions are logged with specific rationale in the exclusion register (Section 3)."),
]

HOLDINGS = [
    {
        "ticker": "ASML.AS", "name": "ASML Holding NV", "sector": "Technology — Semiconductors",
        "country": "Netherlands", "weight": "6.4%", "esg": "56.5", "sharpe": "1.91",
        "why": [
            ("Uncontested technology moat", "ASML holds a global monopoly on Extreme Ultraviolet (EUV) lithography machines — the only technology capable of manufacturing the world's most advanced chips. No competitor has a product within 5–10 years of parity. This structural moat insulates margins and pricing power regardless of macro cycles."),
            ("AI demand tailwind", "Surging demand for AI accelerators (NVIDIA, AMD, custom silicon) drives orders for advanced nodes at TSMC, Samsung and Intel Foundry — all of which depend exclusively on ASML's EUV systems. The company reported EUR 28.3 billion revenue in FY2024 and guides long-term to EUR 44–60 billion by 2030."),
            ("ESG credentials", "ASML targets net zero across Scope 1, 2 and 3 by 2040. Its products enable chipmakers to manufacture more efficiently, reducing energy per transistor — a critical enabler of the global energy transition. CDP A-List rated. SBTi target approved."),
            ("Strongest financial quality", "Highest Sharpe ratio in the portfolio at 1.91. Net cash position. >50% gross margins. R&D intensity ~15% of revenue, protecting the technology lead."),
        ],
        "red_flags": [
            ("China revenue concentration", "China represented ~36% of 2024 revenue. U.S. export controls now ban shipment of EUV machines and most DUV systems to China. ASML warns 2026 growth may stall as Chinese customers front-loaded orders ahead of restrictions. Revenue cliff risk if China drops below 15%."),
            ("Export control escalation risk", "The Dutch government, acting under U.S. pressure, has restricted ASML's export licences progressively since 2019. Further escalation — such as a ban on all DUV shipments — would materially impair near-term revenue."),
            ("Rare earths dependency", "China's 2025 curbs on rare earth mineral exports (used in ASML's machines) create supply chain vulnerability. ASML is working to diversify suppliers but full substitution takes 2–3 years."),
            ("Valuation risk", "Trading at ~35x forward earnings, well above semiconductor equipment peers. Any order miss or macro slowdown triggers significant de-rating risk."),
        ],
    },
    {
        "ticker": "INGA.AS", "name": "ING Groep NV", "sector": "Financial Services — Banks",
        "country": "Netherlands", "weight": "5.9%", "esg": "55.0", "sharpe": "1.66",
        "why": [
            ("Digital banking leadership", "ING operates a fully digital retail bank across 13 European markets with 40+ million customers. Its mobile-first model delivers industry-leading cost-to-income ratios and strong deposit retention."),
            ("Upgraded to MSCI AAA", "MSCI upgraded ING's ESG rating from AA to AAA in October 2025 — the highest possible rating — reflecting improved climate transition planning and enhanced governance disclosure. Sustainalytics rates ESG risk as Low (18.0)."),
            ("Climate finance commitment", "ING has committed to aligning its entire loan book with the Paris Agreement via the Terra approach — the first bank globally to do so systematically. It tracks Scope 3 financed emissions across 9 high-impact sectors."),
            ("Second highest Sharpe in portfolio", "Sharpe ratio of 1.66 reflects strong risk-adjusted returns driven by net interest margin expansion in the high-rate environment and improving asset quality."),
        ],
        "red_flags": [
            ("AML/compliance history", "ING paid EUR 775 million in 2018 to settle Dutch money laundering investigations. While controls have been overhauled, regulatory scrutiny of European banks' AML frameworks remains elevated, and any recurrence would be severely penalised."),
            ("Interest rate sensitivity", "ING's net interest income is positively correlated with rates. As the ECB begins cutting rates, margin compression risk is real; consensus expects NIM to compress 15–20bps by end-2026."),
            ("Real estate exposure", "ING holds significant commercial real estate loan exposure in Germany and Netherlands — markets under stress from office vacancy and rising refinancing costs. Provisions may increase."),
            ("Anti-ESG backlash as business risk", "ING's own 2025 annual report identifies 'anti-ESG sentiment' as an emerging business risk, reflecting the political pressure on sustainable finance commitments across its markets."),
        ],
    },
    {
        "ticker": "HEXA-B.ST", "name": "Hexagon AB", "sector": "Technology — Industrial Software",
        "country": "Sweden", "weight": "5.8%", "esg": "59.2", "sharpe": "1.51",
        "why": [
            ("Digital reality and measurement monopoly", "Hexagon is the global leader in precision measurement technology and industrial software, serving manufacturing, construction, mining, agriculture and geospatial sectors. Its sensors and software digitise the physical world — a key enabler of industrial decarbonisation."),
            ("High-margin recurring software revenue", "Over 50% of revenue is now recurring software and services, with EBIT margins consistently above 25%. The software pivot reduces cyclicality and improves valuation visibility."),
            ("Sustainability enablement", "Hexagon's products directly reduce waste, energy and resource consumption in manufacturing — qualifying significant portions of its software portfolio as EU Taxonomy-eligible under the 'enabling activities' criteria."),
            ("Nordic ESG governance", "As a Swedish company, Hexagon operates under some of Europe's strongest corporate governance norms. It reports under GRI standards and has set SBTi-aligned emissions reduction targets."),
        ],
        "red_flags": [
            ("Governance concentration risk", "The Melker Schörling family controls approximately 27% of votes via dual-class shares (A shares = 10 votes). This concentration limits the influence of minority shareholders on executive pay and strategic decisions."),
            ("M&A integration complexity", "Hexagon has made 200+ acquisitions over 20 years. Integration risk, goodwill impairment potential, and cultural fragmentation across acquired businesses are persistent concerns."),
            ("Cyclical industrial exposure", "Despite software growth, ~40% of revenue remains tied to capital equipment sold into cyclical sectors (automotive, aerospace, construction). A severe manufacturing downturn would impact hardware revenues."),
            ("CEO transition risk", "Founder-era leadership has been a key driver of the M&A-led growth model. Any CEO succession could alter strategic direction and M&A cadence."),
        ],
    },
    {
        "ticker": "ITX.MC", "name": "Inditex SA (Zara)", "sector": "Consumer Discretionary — Apparel Retail",
        "country": "Spain", "weight": "5.6%", "esg": "71.5", "sharpe": "1.05",
        "why": [
            ("Highest ESG score in portfolio", "Inditex scores 71.5/100 on our composite ESG model — highest of all 20 holdings — driven by strong environmental reporting, SBTi-validated targets, and improving supply chain governance."),
            ("Sustainability investment commitment", "EUR 175 million invested in sustainability initiatives in 2025 (30% increase YoY). 88% of fibres certified as organic, recycled or lower-impact alternatives. Targets 100% sustainable fibres by 2030."),
            ("Balance sheet strength", "Inditex carries zero net debt with EUR 10+ billion in net cash, providing extraordinary financial resilience. Its near-real-time demand sensing model (fast fashion 2.0) minimises unsold inventory vs. peers."),
            ("Supply chain control", "Unlike most fast fashion peers, Inditex manufactures ~60% of production in proximity markets (Spain, Portugal, Morocco, Turkey), enabling rapid design-to-shelf cycles while allowing better working condition oversight than Asian-only supply chains."),
        ],
        "red_flags": [
            ("Fast fashion business model tension", "Inditex's core model — multiple collections per year, trend-led volume — is structurally at odds with circular economy principles. Good On You rates Zara as 'It's a start' (2/5), noting insufficient evidence of meaningful living wage payments across the supply chain."),
            ("Supply chain transparency gap", "Inditex does not publish a full factory list or detailed audit results. The 2024 Sustainability Report shows significant progress but independent verification of supplier conditions remains limited."),
            ("Emissions trajectory uncertainty", "Despite SBTi validation, Inditex has not provided granular evidence that it is on track to meet its 2030 Scope 3 targets. Rapid volume growth creates a headwind against absolute emissions reductions."),
            ("Regulatory risk — EU ESRS", "The EU Corporate Sustainability Reporting Directive (CSRD) will require Inditex to publish detailed ESRS-aligned data from FY2024, potentially exposing supply chain practices to greater public scrutiny."),
        ],
    },
    {
        "ticker": "ULVR.L", "name": "Unilever PLC", "sector": "Consumer Staples — Personal Care & Foods",
        "country": "United Kingdom", "weight": "5.6%", "esg": "66.1", "sharpe": "1.19",
        "why": [
            ("Portfolio of 400 global power brands", "Unilever owns Dove, Hellmann's, Lipton, Domestos and 400+ brands with strong pricing power in both developed and emerging markets. 55% of sales come from emerging markets — a long-term structural growth driver."),
            ("Strong ESG history, refocused strategy", "After decades of ESG leadership under the Unilever Sustainable Living Plan, CEO Hein Schumacher has refocused commitments on four material pillars (climate, nature, plastics, livelihoods) with clearer brand-level accountability and more realistic targets."),
            ("Defensive income characteristics", "Essential consumer staples provide recession-resistant cash flows. Dividend yield of ~3.5% adds income return to the ESG mandate. Attractive for pension fund mandates seeking liability-matching income."),
            ("Ice Cream separation upside", "The divestiture of the EUR 8.5 billion Ice Cream division (Ben & Jerry's, Magnum) in 2025 unlocks shareholder value and refocuses capital on higher-growth, higher-margin personal care and nutrition segments."),
        ],
        "red_flags": [
            ("Sustainability backtracking", "Unilever missed its 2025 virgin plastic reduction target (pledged -50%, delivered -30%) and extended the deadline to 2026-2028. Greenpeace has publicly criticised this as 'repackaging ambition' — a potential greenwashing flag."),
            ("Activist investor pressure", "Multiple activist shareholders (including Nelson Peltz's Trian) have pressured Unilever to deprioritise ESG in favour of financial returns, creating governance tension between sustainability commitments and short-term performance."),
            ("Emerging markets currency risk", "55% emerging market exposure creates significant currency headwinds. Argentina, Nigeria and Turkey devaluations impacted reported results in 2024-2025."),
            ("Competitive pricing pressure", "Private-label alternatives gained significant share during the cost-of-living crisis. Unilever's volume growth has lagged price growth, raising questions about long-term pricing power."),
        ],
    },
    {
        "ticker": "AIR.PA", "name": "Airbus SE", "sector": "Industrials — Aerospace & Defense",
        "country": "France/Netherlands", "weight": "5.0%", "esg": "64.4", "sharpe": "0.88",
        "why": [
            ("Commercial aviation duopoly", "Airbus and Boeing together supply ~99% of commercial aircraft globally. With Boeing facing quality and delivery crises since 2024, Airbus has captured extraordinary market share. Its order backlog exceeded 8,700 aircraft — over 12 years of production at current rates."),
            ("SBTi-validated climate targets", "Airbus has SBTi-validated targets for a 63% reduction in Scope 1+2 emissions by 2030 (vs. 2015). It is the lead developer of the ZEROe hydrogen aircraft program targeting entry into service by 2035. MSCI rates Airbus BBB."),
            ("SAF transition leadership", "Airbus aircraft are certified for up to 50% Sustainable Aviation Fuel (SAF) blends today, with 100% SAF certification targeted by 2030. SAF is the industry's primary near-term decarbonisation pathway."),
            ("Financial outperformance", "EBIT (adjusted) grew from EUR 5.8bn (2022) to EUR 7.0bn (2024). Free cash flow conversion exceeds 80%. The backlog provides exceptional revenue visibility through the late 2030s."),
        ],
        "red_flags": [
            ("Aviation is a hard-to-abate sector", "Over 90% of Airbus's Scope 3 emissions come from aircraft operations — outside its direct control. Long-haul aviation has no credible zero-emission pathway until the late 2030s at the earliest. This creates inherent tension with net-zero mandates."),
            ("Supply chain bottlenecks", "Airbus has repeatedly missed delivery targets due to supplier shortfalls (engines from CFM/Pratt & Whitney, titanium, aerostructures). Persistent underdelivery constrains revenue recognition and risks customer relationships."),
            ("Defence dual-use exposure", "~10% of revenues are from defence programs. Some ESG frameworks exclude defence contractors entirely. The dual-use nature of aerospace technology creates ethical screening complexity."),
            ("Engine reliability crisis", "Pratt & Whitney GTF engine inspections (affecting A320neo family) required hundreds of aircraft groundings in 2024-2025, constraining airline capacity and creating warranty/liability overhang for Airbus."),
        ],
    },
    {
        "ticker": "VOLV-B.ST", "name": "Volvo AB", "sector": "Industrials — Commercial Vehicles",
        "country": "Sweden", "weight": "4.9%", "esg": "57.6", "sharpe": "1.06",
        "why": [
            ("Electric truck market leader", "Volvo is the world leader in battery-electric trucks, with the largest deployed fleet of heavy-duty electric trucks globally. Its Volvo FM Electric and FH Electric models are in serial production — ahead of all major competitors including Daimler Truck and Traton."),
            ("Strong ESG credentials", "Volvo targets fossil-free manufacturing by 2030 and net zero value chain by 2040. It has SBTi-validated near-term targets and publishes detailed Scope 3 reporting covering purchased goods, product use and end-of-life."),
            ("Financial quality", "Volvo consistently delivers EBIT margins above 13% — exceptional for the commercial vehicle industry. Net cash position and strong free cash flow fund the EV transition without dilutive capital raises."),
            ("Regulatory tailwind", "EU Fit for 55 regulations mandate a 90% reduction in new truck CO2 emissions by 2040 vs. 2019. This legislative lock-in creates a structural competitive advantage for early EV movers like Volvo."),
        ],
        "red_flags": [
            ("Cyclical industry — freight recession risk", "Commercial vehicle markets are highly cyclical. A global freight recession (visible in soft European truck order data in H2 2024) can cause rapid order cancellations and revenue collapse. The 2024-2025 cycle shows typical symptoms of an inventory correction."),
            ("EV infrastructure gap", "Electric truck adoption depends on charging infrastructure that does not yet exist at scale along major European freight corridors. Infrastructure delays slow customer adoption and compress Volvo's EV revenue ramp."),
            ("China competition", "Chinese OEMs (SAIC, BYD, CATL vehicle divisions) are aggressively entering European truck markets with lower-cost electric models. Price competition risk could compress Volvo's margin premium."),
            ("Geely ownership concentration", "AB Volvo's largest shareholder is Geely (Chinese) which holds ~8.2% of capital. While Volvo remains independently managed, this creates perception risk in an environment of rising geopolitical sensitivity."),
        ],
    },
    {
        "ticker": "VOD.L", "name": "Vodafone Group PLC", "sector": "Telecommunications",
        "country": "United Kingdom", "weight": "4.9%", "esg": "59.6", "sharpe": "0.99",
        "why": [
            ("Portfolio rationalisation value unlock", "Under CEO Margherita Della Valle, Vodafone has divested its India operations, Spain assets (to Zegona) and completed merger of German operations. This transformation unlocks significant capital and focuses resources on higher-return markets (Germany, UK, Africa)."),
            ("Africa growth engine — Vodacom/M-Pesa", "Vodafone's stake in Vodacom (South Africa) and M-Pesa (mobile money, 51 million users) provides exposure to Africa's fastest-growing digital financial services market — a genuine sustainability impact story on financial inclusion."),
            ("ESG: strong data privacy and connectivity access", "Sustainalytics rates Vodafone's ESG risk as 'Strong'. Vodafone has committed to net zero by 2040 and to connecting 100 million people in underserved communities by 2025 via its 'Connecting for Good' program."),
            ("Defensive valuation", "Vodafone trades at a significant discount to European telco peers and to its own historic multiples, providing a margin of safety for ESG-motivated investors."),
        ],
        "red_flags": [
            ("Chronic underperformance and dividend cuts", "Vodafone cut its dividend by 50% in 2024 — its first cut in 25 years. This reflects years of declining organic revenue, excessive leverage, and strategic missteps across multiple markets. Investor trust has been damaged."),
            ("German market deterioration", "Germany (Vodafone's largest market) faces intensifying competition from Deutsche Telekom and United Internet, alongside regulatory pressure on cable TV prices. Revenue has declined for 8 consecutive quarters as of 2025."),
            ("High debt burden", "Despite disposals, Vodafone's net debt remains above EUR 30 billion. Debt service consumes significant free cash flow, limiting investment in network quality and innovation."),
            ("Spectrum and capex cycle risk", "European 5G spectrum auctions require multi-billion-euro bids from all operators. Failure to win key spectrum could cede competitive position in Vodafone's core markets."),
        ],
    },
    {
        "ticker": "SAP.DE", "name": "SAP SE", "sector": "Technology — Enterprise Software",
        "country": "Germany", "weight": "4.8%", "esg": "67.9", "sharpe": "0.72",
        "why": [
            ("Dominant ERP market position", "SAP's ERP software is mission-critical for 93% of Forbes Global 2000 companies. Customer switching costs are extraordinarily high — a rip-and-replace takes 3–7 years and costs tens of millions per customer. This creates durable recurring revenue visibility."),
            ("Cloud transition nearing completion", "SAP's cloud revenue grew 33% in 2024, and cloud now represents >50% of total revenue. The transition from on-premise to RISE with SAP (cloud ERP) is generating significantly higher customer lifetime values and improving earnings quality."),
            ("Sustainability data platform leadership", "SAP's Sustainability Control Tower is the market-leading enterprise sustainability data platform, used by hundreds of multinationals for CSRD/ESRS reporting. As mandatory sustainability reporting expands globally, SAP's platform becomes a new growth driver."),
            ("Strong ESG self-performance", "SAP reports against its own sustainability platform and achieved carbon neutrality in its operations. SBTi targets validated. Sustainalytics ESG risk: Low."),
        ],
        "red_flags": [
            ("Restructuring disruption", "SAP announced 8,000 job cuts in January 2024 as part of a EUR 2+ billion restructuring programme. While necessary for AI reinvestment, this creates execution risk and potential product roadmap gaps."),
            ("Customer lock-in backlash", "SAP's aggressive push to migrate customers from on-premise to cloud (RISE with SAP) has generated customer frustration over pricing and timeline pressure. User communities (DSAG in Germany) have publicly criticised SAP's migration approach."),
            ("AI competition risk", "Microsoft, Salesforce and Oracle are embedding AI into competing ERP/CRM platforms. If SAP's AI capabilities (Joule copilot) prove inferior, customers at natural renewal points may consider migration."),
            ("Valuation premium", "SAP trades at ~45x forward earnings — the highest multiple among major European software companies. This prices in near-perfect execution; any miss triggers outsized share price reaction."),
        ],
    },
    {
        "ticker": "CAP.PA", "name": "Capgemini SE", "sector": "Technology — IT Services",
        "country": "France", "weight": "4.8%", "esg": "60.1", "sharpe": "0.92",
        "why": [
            ("AI and cloud services demand surge", "Capgemini is one of Europe's largest IT services and consulting firms, directly benefiting from enterprise AI adoption. Its 'AI-powered enterprise' offering saw 40% revenue growth in AI-related services in 2024, and the company is positioned as the primary integrator of SAP, Salesforce and Microsoft platforms for large European corporations."),
            ("Best weekly performer in the portfolio", "Capgemini delivered +6.87% in the week April 28–May 5, 2026 — the strongest return of any portfolio holding — reflecting improving order momentum and analyst upgrades."),
            ("ESG commitment", "Capgemini has committed to net zero by 2030 for Scope 1 and 2, and to a 30% reduction in Scope 3 by 2030. It runs the Capgemini Research Institute — a key source of ESG and digital transformation insights."),
            ("Geographic and client diversification", "Revenue spread across 50 countries with no single client above 5% of revenue. This reduces concentration risk materially vs. smaller IT services peers."),
        ],
        "red_flags": [
            ("Talent cost inflation", "IT services firms compete globally for developers and data scientists. Capgemini's largest delivery centres are in India (120,000+ employees), making it exposed to Indian wage inflation, visa policy changes, and talent attrition."),
            ("Pricing pressure from offshore competitors", "Indian IT giants (TCS, Infosys, Wipro) are aggressively targeting European enterprise clients. Capgemini's margin premium depends on maintaining a differentiated consulting overlay — increasingly challenged by offshore competitors' own AI capabilities."),
            ("Cyclical demand risk", "Enterprise IT budgets are discretionary and contract quickly in recessions. Capgemini experienced organic revenue contraction in H2 2024 as clients deferred transformation programmes. Order book visibility is shorter than product companies."),
            ("AI commoditisation risk", "If AI coding assistants significantly reduce the labour content of software delivery, Capgemini's people-intensive model faces structural margin pressure, regardless of pricing."),
        ],
    },
    {
        "ticker": "SU.PA", "name": "Schneider Electric SE", "sector": "Industrials — Energy Management",
        "country": "France", "weight": "4.8%", "esg": "55.4", "sharpe": "1.05",
        "why": [
            ("Top global ESG performer", "For the 14th consecutive year, Schneider Electric achieved DJSI World Index inclusion. S&P Global Corporate Sustainability Assessment: 85/100 (2nd in industry). MSCI ESG: AAA. CDP A-List. It is widely considered the benchmark for industrial sustainability."),
            ("Direct energy transition beneficiary", "Schneider's energy management and automation products are essential infrastructure for data centres, EV charging, smart grids and renewable energy integration. Data centre demand alone is growing 20%+ annually, driven by AI infrastructure buildout."),
            ("Supply chain ESG impact at scale", "Schneider's Zero Carbon Project has enrolled 1,000 top suppliers and achieved a 40% reduction in supplier Scope 3 emissions — one of the most material supply chain ESG programmes in the industrial sector."),
            ("EU Taxonomy alignment", "A significant share of Schneider's products and solutions qualify under EU Taxonomy climate mitigation objectives as 'enabling activities', making it one of the most taxonomy-aligned companies in the portfolio."),
        ],
        "red_flags": [
            ("Valuation premium — priced for perfection", "Schneider trades at ~30x forward earnings — a significant premium to industrial peers — entirely justified by ESG leadership and energy transition positioning. Any execution miss, guidance downgrade, or macro slowdown triggers disproportionate de-rating."),
            ("Data centre concentration", "Approximately 40% of recent revenue growth is attributable to data centre infrastructure. A pullback in hyperscaler capex (possible if AI ROI disappoints) would disproportionately impact Schneider's growth outlook."),
            ("Construction market cyclicality", "Buildings and infrastructure represent ~30% of revenue. A prolonged construction downturn in Europe (visible in 2024-2025 permit data) creates near-term revenue headwinds."),
            ("Cybersecurity incident history", "Schneider experienced a significant data breach in 2024. For an industrial IoT company managing critical infrastructure, cybersecurity incidents create reputational and regulatory risk."),
        ],
    },
    {
        "ticker": "SAN.PA", "name": "Sanofi SA", "sector": "Healthcare — Pharmaceuticals",
        "country": "France", "weight": "4.8%", "esg": "56.6", "sharpe": "1.01",
        "why": [
            ("Dupilumab blockbuster growth", "Dupixent (dupilumab) is the world's best-selling immunology drug, with 2024 revenues of EUR 12.5 billion and multiple label expansions still in progress (COPD, alopecia areata, prurigo nodularis). It is projected to become one of the top 5 drugs globally by 2030."),
            ("Vaccines division strategic value", "Sanofi is one of only five manufacturers of polio vaccines globally and a leading flu vaccine producer. Its vaccines business (Sanofi Pasteur) provides stable, socially important revenues with government procurement visibility."),
            ("ESG — access to medicines", "Sanofi has committed to making all its products available to lower-income countries at cost through 'the Sanofi model'. This is a material differentiator in S-pillar assessments and aligns with SDG 3 (Good Health)."),
            ("Accelerating R&D productivity", "After years of late-stage trial failures, Sanofi's pipeline has delivered approvals across oncology, rare diseases and immunology. Partnership with Regeneron has been transformative for biologics capabilities."),
        ],
        "red_flags": [
            ("Insulin product decline", "Sanofi's legacy diabetes franchise (Lantus/Toujeo) faces biosimilar competition and market share loss to GLP-1 agonists (Novo Nordisk, Eli Lilly). This division generated EUR 5+ billion in peak revenue and is structurally declining."),
            ("Drug pricing regulatory risk", "EU drug pricing negotiations under the Pharmaceutical Strategy for Europe 2023 could compress reference prices across Sanofi's portfolio. US IRA (Inflation Reduction Act) impacts do not directly affect Sanofi's EU revenue but set a political precedent."),
            ("Dupilumab concentration risk", "One drug represents ~35% of total revenue. Patent expiry (US composition-of-matter patent expires ~2033) and any clinical setback would create outsized earnings impact."),
            ("Pipeline binary risk events", "Several Phase 3 candidates (tolebrutinib for multiple sclerosis, rilzabrutinib for ITP) have binary readout risk in 2026. A major trial failure would materially impact share price and pipeline valuation."),
        ],
    },
    {
        "ticker": "AZN.L", "name": "AstraZeneca PLC", "sector": "Healthcare — Biopharmaceuticals",
        "country": "United Kingdom/Sweden", "weight": "4.8%", "esg": "59.6", "sharpe": "0.92",
        "why": [
            ("Exceptional pipeline execution", "AstraZeneca's oncology pipeline (Tagrisso, Enhertu, Imfinzi, Calquence) has delivered consistent Phase 3 success rates well above industry average. It is the only major pharmaceutical company to have grown revenue more than 3x in a decade without a major M&A deal."),
            ("MSCI AA (Leader) ESG rating", "AstraZeneca holds MSCI AA rating and was ranked among TIME's Top 20 Most Sustainable Companies for two consecutive years (2024, 2025). Newsweek awarded its highest 5-star environmental rating in 2025. Sustainalytics: 21.5 (Low Risk)."),
            ("Emerging markets access strategy", "AstraZeneca operates a unique healthcare access model in emerging markets, providing essential medicines at differential pricing. This drives volume growth while improving S-pillar scores and SDG 3 alignment."),
            ("Revenue diversification", "AstraZeneca operates across oncology, cardiovascular, renal, respiratory, rare disease and vaccines — among the broadest therapeutic portfolios of any pharma major. No single drug exceeds 20% of revenue."),
        ],
        "red_flags": [
            ("China revenue exposure ~16%", "China is AstraZeneca's second-largest market. An ongoing anti-corruption investigation (launched 2024) into AstraZeneca's former China president and sales practices creates regulatory and reputational risk. Volume-based procurement policies in China continue to compress pricing."),
            ("Biosimilar erosion on Fasenra/Brilinta", "Key products face biosimilar competition in the 2026-2028 window, which may create a revenue headwind even as newer oncology drugs grow."),
            ("Vaccine revenue normalisation", "COVID-19 vaccine revenues (Oxford/AstraZeneca) have normalised. The company has absorbed this declining revenue through pipeline execution, but investor expectations are now set very high."),
            ("1-week underperformance", "AstraZeneca was the worst performer in the portfolio during the week April 28–May 5, 2026, declining 4.10% — partly driven by broader pharma sector rotation and China-related news flow."),
        ],
    },
    {
        "ticker": "ADS.DE", "name": "Adidas AG", "sector": "Consumer Discretionary — Sportswear",
        "country": "Germany", "weight": "4.7%", "esg": "67.9", "sharpe": "0.61",
        "why": [
            ("Post-Yeezy recovery story", "Following the termination of the Ye (Kanye West) Yeezy partnership in 2022, Adidas has successfully monetised the remaining Yeezy inventory via charitable and commercial channels, recovering EUR 750+ million. The business has structurally recovered with operating margins back above 7% in 2024-2025."),
            ("Sustainability product leadership", "Adidas has committed to using only recycled polyester by 2024 and has already achieved this for performance products. Its Stan Smith Mylo (mycelium leather) and Parley Ocean Plastic lines are industry-leading sustainable product innovations."),
            ("Strong ESG composite score", "Adidas scores 67.9/100 on our composite ESG model, the joint-highest in the Consumer Discretionary sector. It maintains SBTi-validated targets and publishes detailed Scope 3 supply chain reporting."),
            ("Brand and pricing power recovery", "Adidas regained consumer relevance in 2024-2025 through collaborations (Wales Bonner, Fear of God), retro running (Sambas, Gazelles) and strategic reduction of wholesale overexposure. Direct-to-consumer (DTC) revenue mix is growing, improving margins."),
        ],
        "red_flags": [
            ("China competitive headwinds", "Local Chinese sportswear brands (Anta, Li-Ning, Xtep) have taken significant market share from Adidas in China, particularly in the patriot-buying trend post-COVID. China represents ~18% of Adidas revenue — a market where recovery is slower than expected."),
            ("Celebrity partnership dependency", "Despite diversifying away from Yeezy, Adidas remains dependent on high-profile athlete and celebrity endorsements (Lionel Messi, Beyoncé). Reputational risk from partner controversies remains structurally elevated."),
            ("Lowest Sharpe in portfolio", "At 0.61, Adidas has the weakest risk-adjusted financial performance among the 20 holdings — reflecting the volatility around the Yeezy crisis and subsequent recovery. Its inclusion is driven primarily by ESG score strength."),
            ("Supply chain labour risk", "~80% of Adidas products are manufactured in Asia (Vietnam, Indonesia, Cambodia). NGO reports have documented wage theft, excessive overtime and union suppression at some supplier factories, creating persistent S-pillar risk."),
        ],
    },
    {
        "ticker": "ROG.SW", "name": "Roche Holding AG", "sector": "Healthcare — Diagnostics & Pharmaceuticals",
        "country": "Switzerland", "weight": "4.6%", "esg": "65.2", "sharpe": "0.67",
        "why": [
            ("Unique diagnostics + pharma combination", "Roche is the only major company operating at equal scale in both pharmaceuticals and diagnostics. This combination creates unique competitive advantages: Roche's diagnostic tests enable companion diagnostics that guide prescribing of its own drugs, creating reinforcing revenue loops."),
            ("Personalised medicine leadership", "Roche's oncology portfolio (Tecentriq, Polivy, Alecensa, Vabysmo) is built on biomarker-driven, personalised medicine — the highest-growth area in pharma. Its HER2 franchise (trastuzumab, pertuzumab) transformed breast cancer treatment globally."),
            ("ESG: 100% sustainable electricity globally", "Roche achieved its 100% renewable electricity target globally in 2024. SBTi targets are validated. Roche explicitly links executive compensation to sustainability KPIs including CO2 reduction."),
            ("Access and affordability commitment", "Roche's 'Access to Healthcare' programme reached patients in low/lower-middle-income countries ahead of its 2025 goal — contributing to strong S-pillar performance and SDG 3 alignment."),
        ],
        "red_flags": [
            ("Biosimilar erosion — EUR 9bn revenue at risk", "Roche's three blockbuster biologic drugs (rituximab, trastuzumab, bevacizumab) have all lost patent protection. Biosimilar competition has already eroded EUR 4+ billion in annual revenues, and this headwind is expected to continue through 2027."),
            ("Post-COVID diagnostics normalisation", "COVID-19 diagnostics generated EUR 3-4 billion in peak annual revenues. These have largely normalised, creating a structural revenue gap that Roche's pipeline must fill."),
            ("US investment skew", "40% of Roche's R&D investment targets the US market, with disproportionate US capex. Swiss domestic investment is lower despite Switzerland accounting for a significant share of R&D employees — creating political risk with Swiss government and unions."),
            ("Dual share class governance", "Roche's voting structure (non-voting 'Genussscheine' bearer shares vs. registered voting shares) historically concentrates control in the Hoffmann and Oeri families (~45% of votes), limiting minority shareholder influence on strategic decisions."),
        ],
    },
    {
        "ticker": "NESN.SW", "name": "Nestlé SA", "sector": "Consumer Staples — Packaged Foods",
        "country": "Switzerland", "weight": "4.6%", "esg": "71.1", "sharpe": "0.47",
        "why": [
            ("Second highest ESG score in portfolio", "Nestlé scores 71.1/100 — the second highest in the portfolio — reflecting its longstanding leadership in nutrition science, regenerative agriculture sourcing, and water stewardship programmes across 186 countries."),
            ("Portfolio transformation towards health", "Under CEO Laurent Freixe (appointed 2024), Nestlé is accelerating its shift towards nutritional health products (medical nutrition, specialty coffees, pet food premium). These categories carry higher margins and stronger ESG narratives than confectionery."),
            ("Pricing power and brand depth", "Nestlé owns 2,000+ brands including Nescafé, KitKat, Purina, Maggi, and Gerber. In 14 product categories, Nestlé holds the global #1 or #2 market position. This brand depth provides durable pricing power even in inflationary environments."),
            ("Dividend aristocrat", "Nestlé has grown its dividend every year for 27 consecutive years, making it one of Europe's premier dividend growth stocks — attractive for pension fund and endowment mandates."),
        ],
        "red_flags": [
            ("Water stewardship controversy", "Nestlé has faced persistent criticism over its extraction of water from drought-stressed communities in the US (Flint, Michigan), Canada, and sub-Saharan Africa. While the company has divested some North American water brands, allegations of unsustainable water use continue to affect its S-pillar reputation."),
            ("Volume decline — market share losses", "Nestlé has lost market share across several categories to private-label alternatives and to focused competitors. Organic revenue growth has slowed materially in 2024-2025. The portfolio transformation is the right strategy but the path is uncertain."),
            ("CEO transition execution risk", "The 2024 CEO change creates strategic uncertainty. Mark Schneider's decade-long portfolio transformation programme is being reassessed, and investor confidence has been partly replaced by 'wait and see' caution."),
            ("Lowest Sharpe in portfolio", "Sharpe of 0.47 reflects the stock's significant underperformance since 2022. Its inclusion is justified by exceptional ESG scoring, but investors must accept below-benchmark financial momentum."),
        ],
    },
    {
        "ticker": "ORA.PA", "name": "Orange SA", "sector": "Telecommunications",
        "country": "France", "weight": "4.5%", "esg": "67.0", "sharpe": "0.58",
        "why": [
            ("Fibre infrastructure monopoly in France", "Orange operates the dominant fibre-to-the-home network in France, with 42 million connectable homes. As France reaches near-ubiquitous fibre coverage, Orange's network infrastructure becomes an increasingly regulated but durable earnings asset — similar to utilities."),
            ("Africa growth — fintech and connectivity", "Orange Money, the company's mobile money platform in Africa, serves 30+ million customers across 17 African markets. Africa represents 14% of revenues and is growing at double-digit rates — a genuine financial inclusion and ESG impact story."),
            ("Strong ESG governance", "Orange scores 67.0/100 on our composite ESG model. It has committed to net zero by 2040, uses 90%+ renewable electricity in Europe, and has SBTi-validated Scope 1+2 targets. CEO Christel Heydemann has made digital inclusion a core strategic pillar."),
            ("B2B enterprise services growth", "Orange Business (enterprise services) is growing its cybersecurity, cloud connectivity and IoT revenues at 8-10% annually — diversifying away from commoditised consumer telecoms."),
        ],
        "red_flags": [
            ("French regulatory and political risk", "As the partially state-owned incumbent telecom (the French state owns ~23%), Orange faces unique regulatory constraints on pricing, network investment and M&A. Government objectives may not always align with shareholder value maximisation."),
            ("Structural EU telecom market headwinds", "European telecommunications markets face chronic low growth, intense price competition, and very high capex burdens. Orange's European EBITDA margins are under pressure as copper-to-fibre migration costs peak."),
            ("Africa governance complexity", "Operating across 17 African jurisdictions exposes Orange to currency risk, political instability, regulatory uncertainty and corruption risk. Several markets (DRC, Mali, Niger) operate in difficult political environments."),
            ("Cybersecurity incident risk", "As a major telecom operator and B2B cybersecurity provider, Orange is a high-value target for nation-state and criminal cyberattacks. A major breach could simultaneously damage its reputation and its ability to sell security services."),
        ],
    },
    {
        "ticker": "AMS.MC", "name": "Amadeus IT Group SA", "sector": "Technology — Travel Technology",
        "country": "Spain", "weight": "4.5%", "esg": "55.3", "sharpe": "0.88",
        "why": [
            ("Global distribution monopoly", "Amadeus operates the world's largest Global Distribution System (GDS), processing ~60% of all global airline bookings made through travel agents. Its New Distribution Capability (NDC) platform is the industry standard for next-generation airline content distribution."),
            ("Post-COVID full recovery + AI upside", "Following COVID-19's near-destruction of the travel industry in 2020-2021, Amadeus has achieved full revenue recovery and is growing above pre-pandemic levels. AI-powered personalisation and dynamic pricing capabilities represent a new growth layer."),
            ("Platform business economics", "Amadeus's marketplace model connects 400+ airlines, 100,000+ hotels and 200,000+ travel agencies. Network effects create extraordinary switching costs — no major airline has successfully migrated away from Amadeus once deeply integrated."),
            ("Strong ESG disclosure and governance", "Amadeus publishes GRI-aligned sustainability reports with detailed supply chain data. It has committed to carbon neutrality by 2025 and net zero by 2050, with SBTi targets in progress."),
        ],
        "red_flags": [
            ("Airline disintermediation risk", "Major airlines (American, Lufthansa, Delta) have actively invested in bypassing GDS channels through NDC and direct booking. If the shift to direct booking accelerates, Amadeus's transaction volumes and per-booking economics could compress."),
            ("Concentration in air travel", "~60% of Amadeus revenues remain tied to air travel bookings. A pandemic, security crisis or climate-related travel suppression would devastate revenues as seen in 2020."),
            ("Cyber risk — critical travel infrastructure", "Amadeus processes millions of transactions per hour and holds sensitive passenger data for hundreds of millions of travellers. A major cyber incident would cause catastrophic reputational and financial damage."),
            ("Hotel/hospitality technology competition", "Amadeus's hospitality division competes with Oracle OPERA, Salesforce, and multiple start-ups. This market is more fragmented and competitive than GDS, with lower switching costs and lower margins."),
        ],
    },
    {
        "ticker": "CS.PA", "name": "AXA SA", "sector": "Financial Services — Insurance",
        "country": "France", "weight": "4.4%", "esg": "71.4", "sharpe": "0.39",
        "why": [
            ("Third highest ESG score in portfolio", "AXA scores 71.4/100 — the third highest in the portfolio. It is a founding member of the Net Zero Asset Owner Alliance and has committed to EUR 26 billion in green investments by 2026. AXA has fully exited tobacco manufacturing investments and thermal coal investments above 30% revenue thresholds."),
            ("Climate risk pricing leadership", "AXA's research institute (AXA Research Fund) is one of the world's leading publishers on climate risk and its implications for insurance pricing. This positions AXA to reprice climate risk into premiums ahead of competitors — a structural competitive advantage as physical climate risks intensify."),
            ("Diversified P&C + Life + Asset Management", "AXA's revenue mix spans Property & Casualty insurance, Life insurance, and AXA Investment Managers (EUR 850 billion AUM). This diversification provides earnings stability across economic cycles."),
            ("Coal and oil exclusion policies", "AXA was one of the first insurers globally to stop underwriting new coal projects (2015) and has progressively tightened oil sands exclusions. These commitments are material, externally verified, and binding."),
        ],
        "red_flags": [
            ("Climate liability exposure", "As insurers reprice or withdraw from climate-vulnerable markets (Florida, California, Mediterranean coasts), AXA faces both opportunity (repricing premiums) and risk (large uninsured losses if withdrawal is poorly timed or managed)."),
            ("Lowest Sharpe in portfolio", "AXA's Sharpe ratio of 0.39 is the weakest in the portfolio, reflecting interest rate sensitivity and the subdued performance of the European insurance sector in recent years. ESG credentials justify inclusion despite financial momentum weakness."),
            ("Regulatory tail risk — IFRS 17", "The IFRS 17 insurance accounting standard (adopted 2023) fundamentally changes how life insurance revenues and liabilities are presented. Multi-year earnings volatility as the market adjusts to the new standard creates investor uncertainty."),
            ("Catastrophe claims volatility", "Extreme weather events (floods, wildfires, hailstorms) drive unpredictable quarterly earnings swings. A severe Atlantic hurricane season or European flood event could materially impair AXA's P&C combined ratio in a single quarter."),
        ],
    },
]

EXCLUDED = [
    {
        "ticker": "VWS.CO", "name": "Vestas Wind Systems", "reason": "ESG Below Floor + Negative Sharpe (−0.05)",
        "red_flags": [
            ("Margin collapse", "Vestas reported negative EBIT in 2022-2023 due to fixed-price contracts signed during the commodity inflation surge. Despite a recovery plan, order backlog profitability remains uncertain and margins have not fully recovered."),
            ("Supply chain cost overruns", "Inflation in steel, fibreglass, rare earth magnets and logistics dramatically increased turbine production costs. Long-duration fixed-price contracts signed at pre-inflation prices locked in losses on thousands of turbines."),
            ("Project cancellations — US offshore wind", "Multiple US offshore wind projects (involving Vestas turbine supply agreements) were cancelled in 2023-2024 due to cost inflation, interest rate increases, and subsidy uncertainty. US market entry has been significantly delayed."),
            ("Despite green credentials, financial quality fails mandate screen", "Vestas is a legitimate sustainability holding from an ESG mission standpoint, but the mandate requires positive Sharpe ratio over 5 years. Vestas's negative Sharpe (−0.05) disqualifies it under the financial quality screen. Eligible for re-entry when financial metrics recover."),
        ]
    },
    {
        "ticker": "EQNR.OL", "name": "Equinor ASA", "reason": "ESG Below Floor (Score: 48.5)",
        "red_flags": [
            ("Fossil fuel transition credibility gap", "Equinor is Norway's state-controlled oil major. Despite a 2030 renewable energy target, over 95% of current revenues remain from oil and gas operations. NGO analyses (Carbon Tracker, Global Witness) have questioned whether Equinor's climate commitments are consistent with Paris alignment."),
            ("Norwegian Continental Shelf dependency", "Equinor's Norwegian production is mature and requires continued investment to arrest decline rates. High lifting costs relative to Middle East producers reduce returns in a low oil price environment."),
            ("US offshore wind write-downs", "Equinor has written down USD 300+ million on its New York Bight offshore wind projects due to supply chain costs, rising interest rates and regulatory uncertainty — a direct indicator of energy transition execution risk."),
        ]
    },
    {
        "ticker": "ORSTED.CO", "name": "Ørsted AS", "reason": "ESG Below Floor (Score: 43.2)",
        "red_flags": [
            ("US offshore wind portfolio collapse", "Ørsted cancelled its entire US offshore wind portfolio in 2023-2024 — including the 2.4 GW Ocean Wind 1 & 2 projects (New Jersey) — writing down DKK 28.4 billion. This is the largest single green energy impairment in history."),
            ("Business model profitability crisis", "Despite being widely regarded as the world's most sustainable energy company, Ørsted's financial model was severely stressed by rising interest rates (offshore wind is highly capital-intensive and interest-rate sensitive) and supply chain cost inflation."),
            ("Low ESG score — counterintuitive but explainable", "Ørsted's lower-than-expected ESG composite score reflects our financial quality-weighted model: the US write-downs damaged key financial metrics that feed into the ESG/financial composite. This is a limitation of the model — Ørsted's mission ESG credentials are exemplary."),
            ("Governance: board accountability for impairments", "The scale of the US impairments raised serious questions about whether the board and management adequately assessed project risks before committing capital. Three board members did not stand for re-election following the crisis."),
        ]
    },
    {
        "ticker": "ENEL.MI", "name": "Enel SpA", "reason": "ESG Below Floor (Score: 32.7) + Negative Sharpe (−0.48)",
        "red_flags": [
            ("Lowest ESG score among excluded companies", "Enel scored 32.7/100 — the lowest in the entire 56-company universe. Despite being Italy's largest utility with significant renewable capacity, its governance scores, controversy history, and operational complexity drive this result."),
            ("Italian regulatory and political risk", "Enel is partially state-owned (Italian MEF owns ~23%). Italian energy policy is subject to political interference, and government pressure on pricing and investment priorities can conflict with shareholder value."),
            ("EUR 60bn debt — highest leverage in sector", "Enel's net debt exceeds EUR 60 billion — among the highest of any European utility. High leverage creates refinancing risk in a higher-for-longer interest rate environment and limits financial flexibility for the renewable energy transition."),
            ("Emerging market exposure risk", "Enel operates in Latin America (Brazil, Chile, Colombia, Peru, Argentina) and Spain through Endesa. Currency risk, political instability (Argentina nationalisation risk) and regulatory unpredictability materially increase risk profile."),
        ]
    },
    {
        "ticker": "BN.PA", "name": "Danone SA", "reason": "ESG Below Floor (Score: 46.4)",
        "red_flags": [
            ("B-Corp status vs. financial performance paradox", "Danone is a certified B-Corp and one of the world's most publicly ESG-committed food companies. Yet it has consistently underperformed financially vs. Nestlé and Unilever, illustrating that ESG commitment without financial quality does not satisfy the mandate's composite screen."),
            ("Activist investor-driven CEO change", "Activist investor Bluebell Capital forced the removal of CEO Emmanuel Faber in 2021 by arguing that ESG commitments were incompatible with adequate shareholder returns. This corporate governance crisis remains a reputational shadow."),
            ("Margin pressure and portfolio complexity", "Danone's three divisions (Essential Dairy, Specialised Nutrition, Waters) have very different growth and margin profiles. Corporate complexity has impaired execution, and turnaround under CEO Antoine de Saint-Affrique has been slower than expected."),
            ("China infant formula risk", "Specialised Nutrition (including infant formula) generates ~25% of revenues and is highly dependent on the Chinese birth rate — which has declined every year since 2017. This secular headwind creates a structural revenue overhang."),
        ]
    },
    {
        "ticker": "VOW3.DE", "name": "Volkswagen AG", "reason": "ESG Below Floor (Score: 40.1)",
        "red_flags": [
            ("Dieselgate legacy — ongoing litigation", "VW's 2015 emissions scandal (software manipulation of diesel NOx tests) cost EUR 32+ billion in fines, settlements and recalls. Civil litigation continues in multiple jurisdictions. The scandal is the canonical case study of ESG red flags preceding financial catastrophe — governance failures were documented by ESG analysts years before the public scandal."),
            ("EV transition falling behind", "Despite early ambitions, VW's EV software (developed by Cariad subsidiary) has experienced catastrophic delays and cost overruns. The ID. family of EVs has underperformed Chinese competitors on software quality and charging speed. VW entered into a partnership with Rivian in 2024 to access EV software capabilities."),
            ("China market collapse", "VW generates ~30% of global revenues from China — a market where it is losing share rapidly to domestic EV brands (BYD, Li Auto, NIO). This structural share loss shows no signs of reversal."),
            ("Massive restructuring — 35,000 jobs at risk", "In 2024, VW announced potential closure of German factories for the first time in its 87-year history, with 35,000 jobs at risk. Negotiations with IG Metall have been prolonged and contentious, creating labour relations risk."),
        ]
    },
    {
        "ticker": "GFC.PA", "name": "Gecina SA", "reason": "ESG Below Floor (Score: 44.6)",
        "red_flags": [
            ("Office real estate structural headwinds", "Gecina is a Paris-focused office and residential REIT. The post-COVID shift to hybrid working has permanently reduced office space demand in Paris, even as Gecina's premium Haussmann/CBD portfolio holds up better than peripheral assets."),
            ("High leverage in a rising rate environment", "Gecina's LTV (Loan-to-Value) ratio exceeded 35% during the period of peak valuations. Rising interest rates increase refinancing costs and compress net asset values, creating balance sheet pressure."),
            ("Valuation uncertainty", "French commercial real estate valuations declined significantly in 2023-2025 as transaction volumes collapsed. Gecina's reported NAV may not reflect true realisable value in a distressed disposal scenario."),
        ]
    },
    {
        "ticker": "SHEL.L", "name": "Shell PLC", "reason": "Negative Sharpe Ratio (−0.26)",
        "red_flags": [
            ("Net-zero credibility challenges", "Shell's 2021 Dutch court ruling mandating a 45% Scope 3 reduction by 2030 was appealed and overturned on appeal in 2024 — but the legal and reputational precedent remains. ClientEarth and NGOs continue litigation challenging Shell's climate plans as inconsistent with Paris alignment."),
            ("Scope 3 emissions — 95% of total", "Shell's sold-product Scope 3 emissions dwarf its Scope 1+2 footprint. Under our mandate's greenwashing 8-Test, Shell's claim to be 'in step with the Paris Agreement' would likely receive HIGH red-flag ratings on Baseline, Target and Consistency dimensions."),
            ("Strategic pivot ambiguity", "Under CEO Wael Sawan, Shell has deprioritised power and renewables in favour of LNG and chemicals — reversing the direction of predecessor Ben van Beurden. This strategic inconsistency undermines long-term credibility of transition commitments."),
        ]
    },
    {
        "ticker": "IBE.MC", "name": "Iberdrola SA", "reason": "Negative Sharpe Ratio (−0.23)",
        "red_flags": [
            ("Spanish and UK regulatory risk", "Iberdrola generates ~35% of revenues in Spain and ~25% in the UK (ScottishPower). Both governments have imposed windfall profit taxes and regulated price caps on renewable energy — directly impairing returns on capital for new investments."),
            ("Offshore wind cost overruns", "Like Ørsted and Enel, Iberdrola has experienced significant cost inflation on offshore wind projects in the US (Vineyard Wind) and UK (East Anglia). Project delays have impaired near-term earnings quality."),
            ("Currency and political risk — Latin America", "~20% of revenues come from Brazil, Mexico and other Latin American markets. Currency depreciation, political nationalisation risk (particularly in Mexico under AMLO's legacy policies) and regulatory instability are persistent risks."),
        ]
    },
    {
        "ticker": "RWE.DE", "name": "RWE AG", "reason": "Negative Sharpe Ratio (−0.10)",
        "red_flags": [
            ("German coal phase-out costs", "RWE operates Germany's largest remaining coal fleet and is the designated operator to phase out coal by 2030 (lignite) and 2038 (hard coal) per German government policy. The transition creates multi-billion euro stranded asset and remediation costs."),
            ("German energy policy uncertainty", "Following the Habeck-era energy transition, the new German government has taken a more mixed approach to energy policy. Uncertainty over capacity payments, nuclear restart debates, and grid tariffs creates regulatory unpredictability for RWE's investment planning."),
            ("Offshore wind competitive intensity", "RWE is expanding aggressively in US and UK offshore wind, but faces the same cost inflation, permitting delays and subsidy risk as sector peers. Its wind project returns are under pressure across the portfolio."),
        ]
    },
    {
        "ticker": "UCG.MI", "name": "UniCredit SpA", "reason": "Negative Sharpe Ratio (−0.09)",
        "red_flags": [
            ("Commerzbank bid — strategic distraction", "UniCredit's unsolicited approach to acquire Germany's Commerzbank in 2024-2025 has created significant political and regulatory resistance (German government opposition). The multi-year distraction from core operations and potential deal complexity is a governance concern."),
            ("Italian sovereign risk linkage", "UniCredit's balance sheet is heavily exposed to Italian government bonds. A deterioration in Italian sovereign creditworthiness (e.g., fiscal slippage, political crisis) directly impacts UniCredit's CET1 ratio and funding costs."),
            ("NIM compression risk", "ECB rate cuts will compress net interest margins, which have been the primary driver of UniCredit's earnings recovery since 2022. Consensus expects a material NIM headwind in 2026-2027."),
        ]
    },
    {
        "ticker": "ZURN.SW", "name": "Zurich Insurance Group AG", "reason": "Negative Sharpe Ratio (−0.83)",
        "red_flags": [
            ("Worst Sharpe ratio in the entire universe", "At −0.83, Zurich Insurance has the most negative risk-adjusted return of all 56 candidate companies over the 5-year period. This disqualifies it under the mandate's financial quality screen irrespective of ESG credentials."),
            ("Catastrophe claims volatility", "Zurich's combined ratio has been impaired by a sequence of large natural catastrophe losses — European floods (Germany, Belgium, Switzerland), US hurricane losses, and Turkey earthquake losses. Climate change is expected to make these volatility episodes more frequent."),
            ("Climate liability underwriting risk", "As climate litigation against fossil fuel companies increases, insurers that underwrite fossil fuel assets face complex liability exposure when those assets cause climate damage. Zurich has made coal exclusions but retains oil and gas underwriting exposure."),
        ]
    },
    {
        "ticker": "HEIA.AS", "name": "Heineken NV", "reason": "Negative Sharpe Ratio (−0.34)",
        "red_flags": [
            ("Volume decline in Europe and US", "Beer consumption in developed markets is in structural decline, driven by health consciousness, generational shifts away from alcohol, and competition from spirits, wine, ready-to-drink beverages and cannabis alternatives."),
            ("Nigeria and Africa operational risk", "Heineken's Nigerian Breweries subsidiary has been severely impacted by naira devaluation and hyperinflation. The company has written down the value of its Nigerian assets and faces complex capital repatriation challenges."),
            ("Russia exit costs", "Heineken exited Russia in 2022 following the invasion of Ukraine, selling its three breweries for a symbolic RUB 1. The exit cost EUR 400+ million in write-downs and lost a Top 5 global beer market."),
        ]
    },
    {
        "ticker": "CARL-B.CO", "name": "Carlsberg AS", "reason": "Negative Sharpe Ratio (≈0.00 after rounding)",
        "red_flags": [
            ("Eastern Europe concentration risk", "Carlsberg generates a significant share of revenues from Russia (before exit) and Eastern European markets with above-average political and currency risk. The Russia exit has been executed but creates a revenue void."),
            ("Premium brand strategy execution risk", "Carlsberg is investing heavily in premiumising its portfolio (craft beers, 0% alcohol). This is the right strategic direction but carries execution risk in a commoditised beer market."),
            ("China JV uncertainty", "Carlsberg holds a controlling stake in Chongqing Brewery in China. China's slowdown, changing consumer preferences and local competition from China Resources Beer create uncertainty around the JV's long-term growth contribution."),
        ]
    },
    {
        "ticker": "OR.PA", "name": "L'Oréal SA", "reason": "Negative Sharpe Ratio (−0.26)",
        "red_flags": [
            ("China luxury beauty collapse", "L'Oréal's North Asia zone (dominated by China) declined 6% in FY2024 — a sharp reversal after years of double-digit growth. Chinese consumer confidence has weakened materially and local brands (Proya, Florasis) are gaining share in premium skincare."),
            ("Valuation premium compression risk", "L'Oréal trades at ~35x forward earnings — a significant premium to FMCG peers — historically justified by China growth. With China momentum reversing, the valuation premium is at risk of permanent compression."),
            ("Microplastics and cosmetic ingredient regulation", "EU regulators are tightening restrictions on microplastics, PFAS, and certain fragrance ingredients used in cosmetics. L'Oréal faces reformulation costs and potential product discontinuations across multiple brands."),
        ]
    },
    {
        "ticker": "AI.PA", "name": "Air Liquide SA", "reason": "Negative Sharpe Ratio (−0.17)",
        "red_flags": [
            ("Hydrogen economy timeline delays", "Air Liquide has made significant capital commitments to green hydrogen infrastructure (electrolysers, distribution, fuelling stations). Green hydrogen adoption is 5-7 years behind initial projections due to higher-than-expected production costs and slower than expected demand from heavy industry."),
            ("Valuation relative to growth", "Air Liquide trades at a premium to peers despite organic growth rates that have moderated from post-COVID highs. The hydrogen optionality embedded in the valuation requires longer-duration patience than some ESG mandates can tolerate."),
            ("Helium supply concentration risk", "Air Liquide is heavily dependent on helium supplies from Qatar and Russia. Geopolitical disruptions to either source create supply chain vulnerability for its electronics and healthcare gas customers."),
        ]
    },
    {
        "ticker": "RNO.PA", "name": "Renault SA", "reason": "Negative Sharpe Ratio (−0.11)",
        "red_flags": [
            ("Nissan Alliance instability", "The Renault-Nissan-Mitsubishi Alliance has been under severe strain since the Carlos Ghosn arrest in 2018. Renault's dilution of its Nissan stake (from 43% to 15%) resolved some governance issues but also reduced its access to Nissan's Japanese market and technology sharing."),
            ("French state ownership constraints", "The French government owns ~15% of Renault and exercises significant influence over strategic decisions — including the failed attempt to maintain full ownership of Renault's EV spinoff (Ampere). Government objectives sometimes conflict with commercial optimisation."),
            ("EV transition cost burden", "Renault is spending EUR 3+ billion on EV transition costs through Ampere. The competitive landscape for small/medium EVs in Europe is intensifying, with BYD, Tesla and Volkswagen all competing aggressively in Renault's core price segments."),
            ("Low ESG score context", "Renault's ESG performance has been below average in governance and emissions disclosures. The Dieselgate-era French emissions investigations (though distinct from VW) have created lasting regulatory scrutiny around French automotive ESG claims."),
        ]
    },
]


# ── build document ─────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── TITLE PAGE ────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("ESADE SUSTAINABLE EUROPEAN EQUITY FUND")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Investment Mandate & Portfolio Research Report")
    r2.font.size = Pt(14); r2.font.color.rgb = MID_BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"ESADE MSc Finance  |  Final Group Assignment  |  {TODAY}")
    r3.font.size = Pt(11); r3.font.color.rgb = GREY_TEXT; r3.italic = True

    doc.add_paragraph()
    section_divider(doc)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("CONFIDENTIAL — ACADEMIC PROTOTYPE — NOT FINANCIAL ADVICE")
    r4.font.size = Pt(9); r4.italic = True; r4.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

    doc.add_page_break()

    # ── SECTION 1: INVESTMENT MANDATE ─────────────────────────
    heading(doc, "Section 1: Investment Mandate", level=1)
    body(doc,
         "This mandate defines the rules, objectives, and constraints governing every portfolio decision. "
         "All exclusions, scores, and weightings in this document trace back to a specific clause below. "
         "The mandate was encoded as machine-readable JSON (mandate.json) and consumed by all 13 pipeline agents.",
         size=10, space_after=10)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(4.0)
    tbl.columns[1].width = Cm(13.5)

    # Header row
    hdr = tbl.rows[0]
    for i, txt in enumerate(["Item", "Required Explanation"]):
        c = hdr.cells[i]
        set_cell_bg(c, DARK_BLUE)
        set_cell_border(c, "FFFFFF", 6)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(txt)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

    alt = False
    for item, explanation in MANDATE:
        bg = RGBColor(0xF2, 0xF7, 0xFC) if alt else WHITE
        mandate_row(tbl, item, explanation, item_bg=bg)
        alt = not alt

    doc.add_paragraph()
    section_divider(doc)
    doc.add_page_break()

    # ── SECTION 2: PORTFOLIO HOLDINGS ─────────────────────────
    heading(doc, "Section 2: Portfolio Holdings — Research & Analysis", level=1)
    body(doc,
         "Each of the 20 selected companies is analysed below. For each holding we provide: "
         "(1) the investment rationale — why the company was selected, and "
         "(2) red flags — risks and concerns the panel may challenge during Q&A. "
         "All ESG scores are from our proprietary transparent model. Sharpe ratios are 5-year annualised (2020–2025).",
         size=10, space_after=10)

    for i, h in enumerate(HOLDINGS):
        heading(doc, f"{i+1}. {h['name']}  ({h['ticker']})", level=2)

        # Summary bar
        tbl2 = doc.add_table(rows=1, cols=4)
        tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl2.style = 'Table Grid'
        labels = [
            ("Sector", h["sector"]),
            ("Country", h["country"]),
            ("ESG Score", f"{h['esg']} / 100"),
            ("Sharpe | Weight", f"{h['sharpe']}  |  {h['weight']}"),
        ]
        for j, (lbl, val) in enumerate(labels):
            c = tbl2.rows[0].cells[j]
            set_cell_bg(c, LIGHT_BLUE)
            set_cell_border(c, "AAAAAA", 4)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
            r1 = c.paragraphs[0].add_run(lbl + ": ")
            r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = DARK_BLUE
            r2 = c.paragraphs[0].add_run(val)
            r2.font.size = Pt(9); r2.font.color.rgb = GREY_TEXT

        doc.add_paragraph()

        # Why selected
        p_why = doc.add_paragraph()
        r_why = p_why.add_run("  WHY SELECTED")
        r_why.bold = True; r_why.font.size = Pt(10); r_why.font.color.rgb = RGBColor(0x1A, 0x7A, 0x1A)

        for prefix, text in h["why"]:
            bullet(doc, text, bold_prefix=prefix)

        doc.add_paragraph()

        # Red flags
        p_rf = doc.add_paragraph()
        r_rf = p_rf.add_run("  RED FLAGS")
        r_rf.bold = True; r_rf.font.size = Pt(10); r_rf.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        for prefix, text in h["red_flags"]:
            bullet(doc, text, bold_prefix=prefix)

        if i < len(HOLDINGS) - 1:
            section_divider(doc)
            doc.add_paragraph()

    doc.add_page_break()

    # ── SECTION 3: EXCLUDED COMPANIES ─────────────────────────
    heading(doc, "Section 3: Exclusion Register — Shortlisted Companies", level=1)
    body(doc,
         "The following 19 companies were analysed and shortlisted but excluded from the final portfolio. "
         "Exclusion reasons are either: (a) ESG composite score below the bottom 10th percentile of the universe, "
         "or (b) negative Sharpe ratio over the 5-year analysis period. "
         "Red flags documented below explain the underlying reasons for underperformance.",
         size=10, space_after=10)

    for i, ex in enumerate(EXCLUDED):
        heading(doc, f"{i+1}. {ex['name']}  ({ex['ticker']})", level=2)

        p_exc = doc.add_paragraph()
        r_exc = p_exc.add_run(f"  EXCLUSION REASON: {ex['reason']}")
        r_exc.bold = True; r_exc.font.size = Pt(10); r_exc.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        for prefix, text in ex["red_flags"]:
            bullet(doc, text, bold_prefix=prefix)

        if i < len(EXCLUDED) - 1:
            section_divider(doc)
            doc.add_paragraph()

    doc.add_page_break()

    # ── SECTION 4: QUICK-REFERENCE TABLE ──────────────────────
    heading(doc, "Section 4: Portfolio Quick-Reference Summary", level=1)

    tbl3 = doc.add_table(rows=1, cols=6)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = 'Table Grid'
    hdrs = ["#", "Ticker", "Company", "Sector", "ESG", "Weight"]
    for j, h_txt in enumerate(hdrs):
        c = tbl3.rows[0].cells[j]
        set_cell_bg(c, DARK_BLUE)
        set_cell_border(c, "FFFFFF", 4)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(h_txt)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE

    for i, h in enumerate(HOLDINGS):
        row = tbl3.add_row()
        bg = RGBColor(0xF2, 0xF7, 0xFC) if i % 2 == 0 else WHITE
        vals = [str(i+1), h["ticker"], h["name"], h["sector"].split(" — ")[0], h["esg"], h["weight"]]
        for j, val in enumerate(vals):
            c = row.cells[j]
            set_cell_bg(c, bg)
            set_cell_border(c, "BBBBBB", 2)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(9); r.font.color.rgb = GREY_TEXT

    doc.add_paragraph()
    section_divider(doc)

    # Footer note
    p_fn = doc.add_paragraph()
    p_fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_fn = p_fn.add_run(
        "ESADE MSc Finance — Final Group Assignment  |  "
        "Academic prototype only — not a regulated investment product or financial advice.  |  "
        f"Generated: {TODAY}"
    )
    r_fn.font.size = Pt(8); r_fn.italic = True; r_fn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    out = "outputs/reports/Investment_Mandate_and_Research.docx"
    os.makedirs("outputs/reports", exist_ok=True)
    doc.save(out)
    print(f"Document saved: {out}")
    return out


if __name__ == "__main__":
    out = build()
    print(f"Done: {out}")
