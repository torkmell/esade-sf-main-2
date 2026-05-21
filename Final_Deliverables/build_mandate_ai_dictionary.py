"""
Generates: Final_Deliverables/Mandate_AI_Use_DataDictionary.docx
Contains:
  Part 1 — Investment Mandate
  Part 2 — AI Use Statement
  Part 3 — Data Dictionary
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Mandate_AI_Use_DataDictionary.docx")

# ── colour palette ──────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)   # headings
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # sub-headings / table header
LIGHT_BLUE = RGBColor(0xD5, 0xE8, 0xF0)   # table header fill (as hex string below)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT  = RGBColor(0x1A, 0x1A, 0x2E)

# ── helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body(doc, text, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size    = Pt(11)
    run.font.bold    = bold
    run.font.italic  = italic
    run.font.color.rgb = color if color else DARK_TEXT
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    return p

def add_page_break(doc):
    doc.add_page_break()

def make_table(doc, headers, rows, col_widths_cm, header_bg="1F497D", header_fg=WHITE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = header_fg
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = "F2F7FB" if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row_data):
            cell = row_cells[i]
            cell.width = Cm(col_widths_cm[i])
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "CCCCCC")
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_TEXT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()   # spacing after table
    return table

# ── document setup ───────────────────────────────────────────────────────────
doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("ESADE Sustainable European Equity Fund")
tr.font.size  = Pt(22)
tr.font.bold  = True
tr.font.color.rgb = DARK_BLUE

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Investment Mandate  ·  AI Use Statement  ·  Data Dictionary")
sr.font.size  = Pt(13)
sr.font.color.rgb = MID_BLUE

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run("ESADE MSc Finance  |  Final Group Assignment  |  May 2026")
dr.font.size  = Pt(11)
dr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()
doc.add_paragraph()
disc_p = doc.add_paragraph()
disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
disr = disc_p.add_run(
    "This document is an academic prototype produced for ESADE MSc Finance. "
    "It does not constitute financial advice or a regulated investment product."
)
disr.font.size   = Pt(9)
disr.font.italic = True
disr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# PART 1 — INVESTMENT MANDATE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 1: Investment Mandate", level=1)

add_body(doc,
    "Fund name:  ESADE Sustainable European Equity Fund",
    bold=True, space_after=2)
add_body(doc,
    "Strategy:  Long-only, concentrated European equity  |  "
    "Currency: EUR  |  Vintage: May 2026",
    space_after=10)

# ── investment thesis ────────────────────────────────────────────────────────
add_heading(doc, "Investment Thesis", level=2)
add_body(doc,
    "We construct a concentrated long-only European equity portfolio of 15–25 holdings "
    "by integrating fundamental financial quality with best-in-class ESG credentials, "
    "verified against EU Taxonomy alignment and independently screened for greenwashing. "
    "We believe companies with credible sustainability commitments carry lower regulatory, "
    "reputational, and transition risk over a five-year horizon. All exclusions and overrides "
    "are documented, creating an auditable process that can withstand scrutiny under SFDR "
    "Article 8 principles.",
    space_after=10)

# ── portfolio constraints ────────────────────────────────────────────────────
add_heading(doc, "Portfolio Constraints", level=2)
make_table(
    doc,
    headers=["Parameter", "Requirement"],
    rows=[
        ("Investable universe",      "STOXX Europe 600 subset — top 170 constituents by 10-year total return"),
        ("Minimum holdings",         "15 companies"),
        ("Maximum holdings",         "25 companies"),
        ("Target holdings",          "20 companies"),
        ("Maximum single weight",    "10% of portfolio NAV"),
        ("Minimum sectors",          "5 BICS Level-1 sectors"),
        ("Strategy",                 "Long-only equity"),
        ("Benchmark",                "STOXX Europe 600"),
        ("Currency",                 "EUR (all weights and returns EUR-denominated)"),
    ],
    col_widths_cm=[6, 11],
)

# ── composite scoring ────────────────────────────────────────────────────────
add_heading(doc, "Composite Score Construction", level=2)
add_body(doc,
    "Each candidate stock receives a composite score on a 0–100 scale. "
    "The score merges two components: ESG quality and financial efficiency. "
    "Weights reflect the fund's sustainability-first mandate.",
    space_after=6)

make_table(
    doc,
    headers=["Component", "Weight", "Sub-components"],
    rows=[
        ("ESG Score",       "60%", "Environmental 40% · Social 30% · Governance 30%"),
        ("Financial Score", "40%", "Sharpe ratio score · Biodiversity score inverse · EU Taxonomy eligibility"),
    ],
    col_widths_cm=[4, 2.5, 10.5],
)

add_heading(doc, "ESG Pillar Weights by SASB Sector", level=2)
add_body(doc,
    "Environmental, Social and Governance pillar weights are calibrated to SASB materiality "
    "standards for each BICS Level-1 sector. High-impact sectors (Energy, Materials) receive "
    "elevated Environmental weight; service sectors (Financials, Technology) receive higher "
    "Governance and Social weights.",
    space_after=6)

make_table(
    doc,
    headers=["Sector", "E Weight", "S Weight", "G Weight"],
    rows=[
        ("Energy",                    "55%", "25%", "20%"),
        ("Materials",                 "50%", "30%", "20%"),
        ("Industrials",               "45%", "35%", "20%"),
        ("Consumer Discretionary",    "35%", "40%", "25%"),
        ("Health Care",               "25%", "40%", "35%"),
        ("Technology",                "25%", "35%", "40%"),
        ("Financials",                "20%", "35%", "45%"),
    ],
    col_widths_cm=[6, 3, 3, 3],
)

# ── exclusions ───────────────────────────────────────────────────────────────
add_heading(doc, "Hard Exclusions", level=2)
add_body(doc,
    "The following criteria result in automatic removal from the investable universe, "
    "regardless of composite score:",
    space_after=6)
for excl in [
    "Thermal coal revenue exceeding 5% of total revenue",
    "Tobacco production (any revenue contribution)",
    "Involvement in controversial weapons (cluster munitions, anti-personnel mines)",
    "Greenwashing 8-Test score of HIGH on three or more of the eight dimensions",
    "No ESG data available (all three pillars NaN) — hallucination control policy",
    "Subsidiary of another universe company — excluded to prevent carbon double-counting",
]:
    add_bullet(doc, excl)

doc.add_paragraph()
add_heading(doc, "ESG Quality Floor", level=2)
add_body(doc,
    "Companies with a composite ESG score below the universe median floor "
    "(50th percentile of scored universe) are excluded from portfolio consideration. "
    "This floor was set at 50.1 in the current run, removing 16 companies.",
    space_after=10)

# ── watchlist ─────────────────────────────────────────────────────────────────
add_heading(doc, "Watchlist Triggers", level=2)
add_body(doc,
    "Companies meeting any of the following conditions are placed on the watchlist. "
    "Watchlist status does not trigger automatic exclusion but requires additional "
    "human review before inclusion in the final portfolio:",
    space_after=6)
for trigger in [
    "Greenwashing 8-Test score of HIGH on exactly two of the eight dimensions",
    "Science-Based Targets initiative (SBTi) target submitted but not yet approved",
    "Scope 3 emissions not disclosed in the most recent sustainability report",
    "Board gender diversity below 30% of director seats",
]:
    add_bullet(doc, trigger)

doc.add_paragraph()

# ── required metrics ──────────────────────────────────────────────────────────
add_heading(doc, "Required Portfolio Metrics", level=2)
add_body(doc,
    "The following metrics must be calculated and disclosed for the final portfolio "
    "as specified by the assignment mandate:",
    space_after=6)
for metric in [
    "Weighted Average Carbon Intensity (WACI, tCO₂e per €m revenue)",
    "ESG composite score broken down by Environmental, Social, and Governance pillars",
    "Biodiversity / nature-risk proxy score (ENCORE sector dependency + WRI Aqueduct water risk)",
    "EU Taxonomy eligibility percentage for each holding",
    "Greenwashing 8-Test result (JSON) for each portfolio holding",
    "Benchmark comparison vs STOXX Europe 600",
]:
    add_bullet(doc, metric)

doc.add_paragraph()

# ── human override ────────────────────────────────────────────────────────────
add_heading(doc, "Human Override Policy", level=2)
add_body(doc,
    "The investment team may override any quantitative ranking decision for any holding. "
    "Every override must be logged in the Human Review notebook with: the original "
    "quantitative decision, the override decision, a written rationale of at least two "
    "sentences, and the name of the team member responsible. A minimum of three documented "
    "overrides is required for Q&A defence. Override decisions are binding and supersede "
    "all model outputs.",
    space_after=10)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# PART 2 — AI USE STATEMENT
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 2: AI Use Statement", level=1)
add_body(doc,
    "This statement documents the use of artificial intelligence tools in the construction "
    "of the ESADE Sustainable European Equity Fund research pipeline. It is a required "
    "appendix to the written report, submitted in accordance with ESADE MSc Finance "
    "academic integrity guidelines.",
    space_after=10)

# 2.1 tools
add_heading(doc, "2.1  Tools Used", level=2)
make_table(
    doc,
    headers=["Tool", "Provider", "Purpose"],
    rows=[
        ("Claude (Sonnet / Opus)",
         "Anthropic",
         "Mandate drafting; ESG analysis; greenwashing 8-Test; document intelligence "
         "(RAG extraction from PDF sustainability reports); natural language generation "
         "for report sections; Python code generation"),
        ("Python / Jupyter Notebooks",
         "Open source",
         "Quantitative data processing, ESG scoring, financial metrics calculation, "
         "portfolio optimisation, chart generation"),
        ("yfinance",
         "Open source",
         "Market price data retrieval (adjusted closing prices, 2020–2025)"),
        ("n8n.cloud",
         "n8n GmbH",
         "Pipeline orchestration — connecting individual agent notebooks into a "
         "sequential automated workflow"),
        ("Claude Projects (RAG)",
         "Anthropic",
         "Structured extraction of quantitative ESG data from PDF sustainability reports, "
         "TCFD disclosures, and CSRD/ESRS filings with verbatim citations"),
    ],
    col_widths_cm=[3.5, 3, 10.5],
)

# 2.2 what AI did
add_heading(doc, "2.2  What AI Did", level=2)
for item in [
    "Generated all Python code in notebooks (team verified outputs, not source code)",
    "Extracted structured quantitative data from PDF sustainability reports with verbatim citations and page numbers",
    "Applied the 8-dimension greenwashing test to each company's sustainability claims, flagging dimensions as LOW / MED / HIGH / MISSING",
    "Drafted and revised report sections based on team-provided data and quantitative outputs",
    "Assisted in interpreting regulatory frameworks: EU Taxonomy, SFDR Article 8, CSRD/ESRS, PAI indicators",
    "Suggested scoring weights and exclusion thresholds which the team reviewed, amended, and approved",
]:
    add_bullet(doc, item)

doc.add_paragraph()

# 2.3 what humans did
add_heading(doc, "2.3  What Humans Did", level=2)
for item in [
    "Defined all investment mandate parameters, scoring weights, and hard exclusion rules",
    "Verified a 30% random sample of all AI-extracted ESG data against the source PDF with page number",
    "Verified 100% of exclusion and watchlist decisions against the primary source before finalising",
    "Exercised a minimum of three documented override decisions logged in the Human Review notebook",
    "Reviewed and approved all portfolio construction decisions including final stock selection and weights",
    "Validated all financial calculations against yfinance raw data or primary company source",
    "Made all judgement-based classifications in the Data Dictionary (reported / estimated / judgement-based)",
]:
    add_bullet(doc, item)

doc.add_paragraph()

# 2.4 hallucination controls
add_heading(doc, "2.4  Hallucination Controls", level=2)
for item in [
    "All Claude Projects RAG outputs are marked MISSING where information was absent in the source document — AI never invents a figure",
    "Every AI-extracted data point includes a verbatim source quote and page number to enable manual verification",
    "AI-estimated values are clearly classified as 'estimated' in the Data Dictionary; reported figures are classified as 'reported'",
    "ESG ratings are treated as quantitative indicators, not as objective ground truth — triangulation across two independent sources is required before inclusion",
    "The greenwashing 8-Test requires at least one direct quote per dimension; a dimension with no quotable evidence is scored MISSING, not LOW",
]:
    add_bullet(doc, item)

doc.add_paragraph()

# 2.5 limitations
add_heading(doc, "2.5  Limitations of AI-Assisted Analysis", level=2)
for item in [
    "Biodiversity scores use sector-level proxies (ENCORE dependency scores and WRI Aqueduct water risk) due to the absence of company-level TNFD disclosures in the dataset",
    "EU Taxonomy reported alignment data is sparse (8 of 167 companies); EU Taxonomy eligibility is used as a proxy and must not be conflated with confirmed alignment",
    "Market data is sourced from Yahoo Finance adjusted closing prices, which may differ from Bloomberg terminal data used in professional settings",
    "AI extraction quality depends on the text layer quality of source PDFs; scanned documents without OCR may produce lower-quality extractions",
    "Beta estimates are currently unavailable due to a STOXX 600 benchmark ticker download failure; portfolio market sensitivity cannot be reported",
    "The investable universe is the top 170 STOXX Europe 600 constituents by 10-year total return, which introduces survivorship bias and look-ahead bias — disclosed as a known limitation in the methodology",
]:
    add_bullet(doc, item)

doc.add_paragraph()

add_body(doc,
    "This portfolio is an academic prototype produced for ESADE MSc Finance. "
    "It does not constitute financial advice or a regulated investment product.",
    italic=True, color=RGBColor(0x70, 0x70, 0x70), space_after=10)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# PART 3 — DATA DICTIONARY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Part 3: Data Dictionary", level=1)
add_body(doc,
    "This dictionary documents every variable used in the pipeline. Each variable is "
    "classified using the following extraction-method taxonomy:",
    space_after=6)
make_table(
    doc,
    headers=["Classification", "Definition"],
    rows=[
        ("Reported",          "Value directly stated in a company's regulatory filing, sustainability report, or official data provider feed"),
        ("Observed",          "Value derived from observed market prices (e.g., returns, volatility calculated from yfinance data)"),
        ("Estimated",         "Value calculated by the pipeline using a model, proxy, or imputation method — not directly reported by the company"),
        ("AI-extracted",      "Value pulled from a PDF source document by Claude RAG with a verbatim citation; manually verified at 30% sample rate"),
        ("Judgement-based",   "Value or classification assigned by the investment team using professional judgement; no algorithmic rule fully determines the output"),
    ],
    col_widths_cm=[4, 13],
)

# ── 3.1 identity ──────────────────────────────────────────────────────────────
add_heading(doc, "3.1  Company Identity Fields", level=2)
make_table(
    doc,
    headers=["Variable", "Definition", "Unit / Format", "Source", "Classification", "Confidence"],
    rows=[
        ("ticker",                   "Bloomberg ticker (primary key across ESG, biodiversity and EU regulation files)",  "String (e.g. ASME)",   "Bloomberg / course dataset",    "Reported",  "High"),
        ("yf_ticker",                "Yahoo Finance ticker used to download market prices",                               "String (e.g. ASML.AS)","Bloomberg / course dataset",    "Reported",  "High"),
        ("idBbGlobalCompanyName",    "Full legal company name from Bloomberg Global Company ID",                          "String",               "Bloomberg / course dataset",    "Reported",  "High"),
        ("ISIN",                     "International Securities Identification Number",                                    "12-char alphanumeric", "Bloomberg / course dataset",    "Reported",  "High"),
        ("country",                  "Country of primary listing",                                                        "ISO 3166-2",           "Bloomberg / course dataset",    "Reported",  "High"),
        ("bics_sector",              "BICS Level-1 sector classification (classificationLevelName1 in raw data)",          "String (8 categories)","Bloomberg BICS taxonomy",       "Reported",  "High"),
        ("data_vintage",             "Date on which ESG/biodiversity/EU data was extracted and frozen",                   "YYYY-MM-DD",           "Pipeline metadata",             "Observed",  "High"),
    ],
    col_widths_cm=[3.5, 5.5, 2.8, 3, 2.3, 1.9],
)

# ── 3.2 esg ───────────────────────────────────────────────────────────────────
add_heading(doc, "3.2  ESG Scores", level=2)
make_table(
    doc,
    headers=["Variable", "Definition", "Unit / Range", "Source", "Classification", "Confidence"],
    rows=[
        ("E_score",                  "Environmental pillar score: SASB-weighted aggregate of GHG emissions intensity, water usage, and waste/hazardous-materials metrics", "0–100",  "Bloomberg ESG dataset + SASB weights", "Estimated",  "Medium"),
        ("S_score",                  "Social pillar score: SASB-weighted aggregate of workforce safety, diversity, and supply-chain metrics",                              "0–100",  "Bloomberg ESG dataset + SASB weights", "Estimated",  "Medium"),
        ("G_score",                  "Governance pillar score: SASB-weighted aggregate of board diversity, exec compensation, and anti-corruption metrics",                "0–100",  "Bloomberg ESG dataset + SASB weights", "Estimated",  "Medium"),
        ("ESG_score",                "Composite ESG score: E × pillar_weight_E + S × pillar_weight_S + G × pillar_weight_G (weights vary by BICS sector)",               "0–100",  "Pipeline calculation",                 "Estimated",  "Medium"),
        ("bloomberg_esg_disclosure", "Bloomberg ESG Disclosure Score — measures comprehensiveness of company's ESG reporting, not ESG performance",                       "0–100",  "Bloomberg",                            "Reported",   "High"),
        ("sustainalytics_risk_score","Sustainalytics ESG Risk Rating — unmanaged ESG risk score (lower = lower risk). NaN where not available.",                          "0–50",   "Sustainalytics (where available)",     "Reported",   "Medium"),
        ("tri_sources_available",    "Number of independent ESG data sources available for triangulation (max 2 in current build)",                                       "Integer 0–2", "Pipeline metadata",               "Observed",   "High"),
        ("tri_passes",               "Number of sources that pass the triangulation quality threshold",                                                                   "Integer 0–2", "Pipeline calculation",            "Estimated",  "Medium"),
        ("triangulation_result",     "Triangulation verdict: PASS (two sources agree), WATCHLIST (single source or marginal agreement)",                                  "PASS / WATCHLIST", "Pipeline judgement",          "Judgement-based", "Medium"),
        ("esg_data_flag",            "Data quality flag: OK = scored normally; LOW_DATA = one or more pillars missing (company excluded from portfolio)",                  "OK / LOW_DATA", "Pipeline metadata",              "Observed",   "High"),
    ],
    col_widths_cm=[3.5, 5.8, 2.5, 3, 2.3, 1.9],
)

# ── 3.3 climate ────────────────────────────────────────────────────────────────
add_heading(doc, "3.3  Climate & Carbon", level=2)
make_table(
    doc,
    headers=["Variable", "Definition", "Unit / Range", "Source", "Classification", "Confidence"],
    rows=[
        ("carbon_intensity",         "Company-level carbon intensity: total Scope 1+2 GHG emissions divided by revenue",                                   "tCO₂e / €m revenue", "Bloomberg ESG dataset",      "Reported",  "Medium"),
        ("ci_source",                "Source flag for carbon_intensity: bloomberg_calc = direct from Bloomberg; sector_median_imputed = filled from BICS sector median where direct data absent", "String", "Pipeline metadata", "Observed", "High"),
        ("WACI",                     "Weighted Average Carbon Intensity of the final portfolio: sum(weight_i × carbon_intensity_i). Portfolio-level metric required by SFDR.",  "tCO₂e / €m revenue", "Pipeline calculation", "Estimated", "Low–Medium"),
    ],
    col_widths_cm=[3.5, 5.8, 2.5, 3, 2.3, 1.9],
)
add_body(doc,
    "Note: 18 of 20 portfolio holdings use sector-median imputation for carbon_intensity "
    "due to absent Bloomberg direct data. WACI should be treated as an order-of-magnitude "
    "estimate. Confidence improves once primary-source Scope 1+2 data are obtained.",
    italic=True, space_after=10)

# ── 3.4 financial ──────────────────────────────────────────────────────────────
add_heading(doc, "3.4  Financial Metrics", level=2)
make_table(
    doc,
    headers=["Variable", "Definition", "Unit / Range", "Source", "Classification", "Confidence"],
    rows=[
        ("annual_return_pct",  "Annualised total return calculated from Yahoo Finance adjusted closing prices, 2020–2025",     "%",          "yfinance",              "Observed",   "High"),
        ("annual_volatility_pct", "Annualised return volatility (standard deviation of daily returns × √252), same period",   "%",          "yfinance",              "Observed",   "High"),
        ("sharpe_ratio",       "Sharpe ratio: (annual_return − risk_free_rate) / annual_volatility. Risk-free rate = 0%.",   "Dimensionless","Pipeline calculation",  "Observed",   "High"),
        ("max_drawdown_pct",   "Maximum peak-to-trough drawdown over the 2020–2025 observation window",                       "%",          "yfinance",              "Observed",   "High"),
        ("roe_pct",            "Return on Equity as reported in Bloomberg financial data",                                     "%",          "Bloomberg",             "Reported",   "Medium"),
        ("debt_to_equity",     "Total debt divided by total equity (balance sheet ratio)",                                     "Ratio",      "Bloomberg",             "Reported",   "Medium"),
        ("revenue_growth_pct", "Most recent annual revenue growth rate",                                                       "%",          "Bloomberg",             "Reported",   "Medium"),
        ("beta",               "Market beta vs STOXX Europe 600. Currently all NaN — benchmark download failure.",            "Dimensionless","yfinance (failed)",    "Observed",   "N/A"),
        ("sharpe_score",       "Sharpe ratio normalised to 0–100 scale by percentile rank within the scored universe",        "0–100",      "Pipeline calculation",  "Estimated",  "High"),
    ],
    col_widths_cm=[3.8, 5.5, 2.5, 2.7, 2.5, 2],
)

# ── 3.5 biodiversity ──────────────────────────────────────────────────────────
add_heading(doc, "3.5  Biodiversity & Nature Risk", level=2)
make_table(
    doc,
    headers=["Variable", "Definition", "Unit / Range", "Source", "Classification", "Confidence"],
    rows=[
        ("encore_score",       "ENCORE (Exploring Natural Capital Opportunities, Risks and Exposure) dependency score for the company's primary BICS sector. Measures how much the sector depends on ecosystem services.", "0–5 (5 = highest dependency)", "ENCORE / UNEP-WCMC — sector-level proxy", "Estimated", "Low"),
        ("aqueduct_score",     "WRI Aqueduct water-risk score for the company's primary BICS sector. Measures baseline water stress exposure.",                                                                              "0–5 (5 = highest stress)",      "WRI Aqueduct — sector-level proxy",       "Estimated", "Low"),
        ("biodiversity_score", "Composite nature-risk score: (encore_score × 10) + (aqueduct_score × 10). Higher = higher nature risk. Used inverted in composite scoring (lower risk = better).",                         "0–100",                          "Pipeline calculation",                    "Estimated", "Low"),
        ("nature_risk_tier",   "Categorical nature-risk classification: Low (0–29), Medium (30–49), High (50+)",                                                                                                            "Low / Medium / High",            "Pipeline classification",                 "Estimated", "Low"),
        ("bio_score_inv",      "Inverted biodiversity score used in composite: 100 − biodiversity_score. High value = low nature risk = positive attribute.", "0–100", "Pipeline calculation", "Estimated", "Low"),
    ],
    col_widths_cm=[3.5, 5.8, 2.5, 3, 2.3, 1.9],
)
add_body(doc,
    "Note: All biodiversity scores are sector-level proxies. Company-level TNFD disclosures "
    "are not yet available for this universe. Scores should be interpreted as directional "
    "indicators only. All 20 portfolio holdings are in the Low nature-risk tier except "
    "Diploma PLC and VAT Group (both Medium).",
    italic=True, space_after=10)

# ── 3.6 eu taxonomy ────────────────────────────────────────────────────────────
add_heading(doc, "3.6  EU Taxonomy & SFDR Compliance", level=2)
make_table(
    doc,
    headers=["Variable", "Definition", "Unit / Range", "Source", "Classification", "Confidence"],
    rows=[
        ("taxonomy_eligible_pct",           "Percentage of company revenue that is potentially eligible for EU Taxonomy classification (does not mean it is aligned)", "%",           "Bloomberg / legalEntityEuTaxonomy.csv", "Reported",  "Medium"),
        ("taxonomy_aligned_pct",            "Percentage of company revenue reported as EU Taxonomy-aligned. Sparse: only 8/167 companies have reported alignment data.", "%",         "Bloomberg / legalEntityEuTaxonomy.csv", "Reported",  "Low"),
        ("euTaxnmyEstmatdDnshMitgtnLevl1",  "Estimated DNSH (Do No Significant Harm) score for Climate Change Mitigation (Objective 1)",                                "0–100",      "Bloomberg",                             "Estimated", "Low"),
        ("euTaxnmyEstmatdDnshMitgtnLevl2",  "Estimated DNSH score for Climate Change Adaptation (Objective 2)",                                                         "0–100",      "Bloomberg",                             "Estimated", "Low"),
        ("euTaxnmyEstmatdDnshAdapttnLevl1", "Estimated DNSH score for Sustainable Use of Water (Objective 3)",                                                          "0–100",      "Bloomberg",                             "Estimated", "Low"),
        ("euTaxnmyEstmatdDnshAdapttnLevl2", "Estimated DNSH score for Circular Economy (Objective 4)",                                                                  "0–100",      "Bloomberg",                             "Estimated", "Low"),
        ("eu_score",                        "EU Taxonomy composite score used in portfolio ranking: derived from eligibility and DNSH estimates via pipeline formula",   "0–100",      "Pipeline calculation",                  "Estimated", "Low"),
        ("sfdr_compliant",                  "Whether the company meets the minimum SFDR Article 8 PAI indicator thresholds defined in the mandate",                     "True / False","Pipeline judgement",                    "Judgement-based", "Medium"),
    ],
    col_widths_cm=[4, 5.5, 2.3, 2.8, 2.5, 1.9],
)
add_body(doc,
    "Critical note: taxonomy_eligible_pct and taxonomy_aligned_pct must not be conflated. "
    "Eligibility means the activity could in principle qualify; alignment means the company "
    "has formally reported conformance with the Taxonomy technical screening criteria. "
    "The vast majority of companies in this universe have eligibility data but no reported alignment.",
    italic=True, space_after=10)

# ── 3.7 greenwashing ──────────────────────────────────────────────────────────
add_heading(doc, "3.7  Greenwashing Assessment", level=2)
make_table(
    doc,
    headers=["Variable", "Definition", "Unit / Range", "Source", "Classification", "Confidence"],
    rows=[
        ("gw_exclude",    "Hard exclusion flag from greenwashing 8-Test: True if company scores HIGH on three or more of the eight dimensions", "True / False", "Claude Projects RAG + human verification", "AI-extracted + Judgement-based", "Pending"),
        ("gw_watchlist",  "Watchlist flag from greenwashing 8-Test: True if company scores HIGH on exactly two dimensions",                    "True / False", "Claude Projects RAG + human verification", "AI-extracted + Judgement-based", "Pending"),
    ],
    col_widths_cm=[3.5, 6, 2, 3, 3, 1.5],
)
add_body(doc,
    "Note: All greenwashing flags are currently set to False (not yet screened). The RAG "
    "Operator must complete the 8-Test for all 20 portfolio holdings before the greenwashing "
    "agent (notebook 09) can run. Portfolio weights may shift by 1–3 holdings once "
    "screening is complete.",
    italic=True, space_after=10)

# ── 3.8 portfolio construction ────────────────────────────────────────────────
add_heading(doc, "3.8  Portfolio Construction Variables", level=2)
make_table(
    doc,
    headers=["Variable", "Definition", "Unit / Range", "Source", "Classification", "Confidence"],
    rows=[
        ("composite_score",  "Final ranking score: (ESG_score × 0.40) + (composite_financial_score × 0.60). Scale normalised to 0–100.", "0–100", "Pipeline calculation", "Estimated", "Medium"),
        ("rank",             "Rank within the post-exclusion universe by composite_score (1 = highest score)",                                                            "Integer 1–N", "Pipeline calculation", "Estimated", "High"),
        ("weight_raw",       "Raw portfolio weight before rounding: 1/N where N = number of eligible holdings above composite floor",                                     "Decimal 0–1", "Pipeline calculation", "Estimated", "High"),
        ("weight",           "Final portfolio weight, rounded to 4 decimal places. Weights sum to 1.0. No single weight exceeds 0.10.",                                  "Decimal 0–1", "Pipeline calculation", "Estimated", "High"),
    ],
    col_widths_cm=[3.5, 6.5, 2.3, 2.7, 2.3, 1.7],
)

doc.add_paragraph()
add_body(doc,
    "All scores and derived variables carry the vintage date of the underlying input data. "
    "Financial metrics vintage: 2026-05-12. ESG, biodiversity, and EU taxonomy vintage: 2026-05-14. "
    "Portfolio construction vintage: 2026-05-14.",
    italic=True, space_after=4)

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
