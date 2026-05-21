"""
Generates: Final_Deliverables/Pipeline_Section3_Section4.docx

Part 1 — Appendix: Pipeline Architecture & Agent Descriptions
Part 2 — Report Section 3: Universe Construction & Data Sources
Part 3 — Report Section 4: Quantitative Financial Screening
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Pipeline_Section3_Section4.docx")

# ── colours ──────────────────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
GREY_TEXT = RGBColor(0x70, 0x70, 0x70)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ── helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4")
        el.set(qn("w:space"),"0"); el.set(qn("w:color"), color)
        tcB.append(el)
    tcPr.append(tcB)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18 if level==1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body(doc, text, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11); run.font.bold = bold; run.font.italic = italic
    run.font.color.rgb = color if color else DARK_TEXT
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = DARK_TEXT
    run = p.add_run(text)
    run.font.size = Pt(11); run.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    return p

def make_table(doc, headers, rows, col_widths_cm, header_bg="1F497D"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = Cm(col_widths_cm[i])
        set_cell_bg(hdr[i], header_bg); set_cell_borders(hdr[i], "FFFFFF")
        p = hdr[i].paragraphs[0]; p.clear()
        run = p.add_run(h); run.font.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = WHITE; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        bg = "F2F7FB" if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row_data):
            cells[i].width = Cm(col_widths_cm[i])
            set_cell_bg(cells[i], bg); set_cell_borders(cells[i], "CCCCCC")
            p = cells[i].paragraphs[0]; p.clear()
            run = p.add_run(str(val)); run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_TEXT
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()
    return table

# ── document setup ────────────────────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(2.5)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── cover ─────────────────────────────────────────────────────────────────────
for _ in range(3): doc.add_paragraph()
for text, size, color in [
    ("ESADE Sustainable European Equity Fund", 22, DARK_BLUE),
    ("Pipeline Architecture  ·  Universe Construction  ·  Financial Screening", 13, MID_BLUE),
    ("ESADE MSc Finance  |  Final Group Assignment  |  May 2026", 11, GREY_TEXT),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color
    if size == 22: r.font.bold = True
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 1 — PIPELINE ARCHITECTURE (APPENDIX)
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Appendix: Pipeline Architecture & Agent Guide", level=1)
add_body(doc,
    "This appendix describes the AI-agent research pipeline built to screen 167 STOXX Europe 600 "
    "companies and construct the ESADE Sustainable European Equity Fund. The pipeline implements "
    "13 agent roles across 12 Jupyter notebooks. Agents run sequentially, communicating exclusively "
    "through CSV and JSON files stored in the outputs/ folder. No agent modifies the professor-provided "
    "source data, and any single notebook can be re-run independently without cascading side-effects.",
    space_after=10)

# A.1 execution order
add_heading(doc, "A.1  Pipeline Execution Order", level=2)
make_table(doc,
    headers=["Step", "Agent", "Notebook", "Type"],
    rows=[
        ("1",  "Agent 1 — Mandate",            "01_mandate.ipynb",               "Automated"),
        ("2",  "Agent 2 — Data Ingestion",      "02_data_ingestion.ipynb",        "Automated"),
        ("3",  "Agent 3 — Data Quality",        "03_data_quality.ipynb",          "Automated"),
        ("4",  "Agent 10 — Financial Analysis", "04_financial_analysis.ipynb",    "Automated"),
        ("5",  "Agents 5 & 6 — ESG + Climate",  "05_esg_climate.ipynb",           "Automated"),
        ("6",  "Agent 7 — Biodiversity",        "07_biodiversity.ipynb",          "Automated"),
        ("7",  "Agent 8 — EU Regulation",       "08_eu_regulation.ipynb",         "Automated"),
        ("8",  "Agent 4 — Document Intelligence","06_document_intelligence.ipynb","Manual — RAG Operator"),
        ("9",  "RAG Screening Sheet",           "data/rag/RAG_Screening_Sheet_Workbook_v1.xlsx", "Manual — RAG Operator"),
        ("10", "Agent 9 — Greenwashing",        "09_greenwashing.ipynb",          "Automated"),
        ("11", "Agent 11 — Portfolio Construction","10_portfolio_construction.ipynb","Automated"),
        ("12", "Agent 12 — Human Review",       "11_human_review.ipynb",          "Manual + Automated"),
        ("13", "Agent 13 — Reporting",          "12_reporting.ipynb",             "Automated"),
    ],
    col_widths_cm=[1.2, 4.5, 6.2, 5.1],
)

# A.2 folder structure
add_heading(doc, "A.2  Folder Structure & Data Handoffs", level=2)
make_table(doc,
    headers=["Folder", "Contents", "Read by"],
    rows=[
        ("data/provided/",   "Professor-provided CSVs — never modified",                      "Agent 2"),
        ("data/market/",     "yfinance price CSVs, date-stamped",                              "Agent 10"),
        ("data/rag/",        "RAG Screening Sheet Excel workbook",                             "Agent 9"),
        ("outputs/scores/",  "master_dataset, ESG scores, financial metrics, biodiversity, EU regulation", "Agents 9, 11, 13"),
        ("outputs/portfolio/","Final portfolio, exclusion log, universe scores",               "Agents 12, 13"),
        ("outputs/reports/", "Charts (PNG), factsheet text, pipeline diagram",                 "Presentation / report"),
        ("outputs/rag/",     "Greenwashing JSON extractions per company",                      "Agent 9"),
    ],
    col_widths_cm=[3.5, 8, 5.5],
)

# A.3 composite score formula
add_heading(doc, "A.3  Composite Score Formula", level=2)
add_body(doc,
    "The final ranking in Agent 11 (Portfolio Construction) uses the following formula to produce "
    "a composite score on a 0–100 scale for each eligible company:",
    space_after=6)
add_body(doc,
    "Composite Score  =  (ESG_score × 0.60)  +  (Sharpe_score × 0.24)  +  "
    "(Biodiversity_score_inverted × 0.12)  +  (EU_eligibility_score × 0.04)",
    bold=True, space_after=6)
add_body(doc,
    "The ESG score receives the dominant weight (60%) consistent with the fund's sustainability-first "
    "mandate. The financial component (40% total) is split across Sharpe ratio (24%), biodiversity "
    "risk (12%), and EU Taxonomy eligibility (4%). All components are normalised to 0–100 by "
    "percentile rank before weighting. The biodiversity score is inverted so that lower nature risk "
    "produces a higher score. If any input file is absent, its weight is redistributed proportionally "
    "to present components.",
    space_after=10)

# A.4 exclusion layers
add_heading(doc, "A.4  Exclusion Logic (Applied in Order)", level=2)
make_table(doc,
    headers=["Layer", "Agent", "Rule", "Companies excluded"],
    rows=[
        ("1 — Subsidiary filter",   "Agent 11", "Subsidiary of another universe company (prevents carbon double-counting)", "4"),
        ("2 — No ESG data",         "Agent 11", "All three ESG pillars NaN — hallucination control policy", "5"),
        ("3 — Greenwashing",        "Agent 9",  "HIGH rating on 3 or more of 8 greenwashing test dimensions", "TBD (post-RAG)"),
        ("4 — Very High nature risk","Agent 7",  "ENCORE nature risk tier = Very High", "0"),
        ("5 — ESG quality floor",   "Agent 11", "Composite ESG score below 50th-percentile floor (floor = 50.1)", "16"),
        ("6 — Volatility cap",      "Agent 10", "Annualised return volatility > 40%", "0"),
        ("7 — Manual override",     "Agent 12", "Captain decision with documented rationale", "≥ 3 for Q&A"),
    ],
    col_widths_cm=[4, 2.5, 7.5, 3],
)
add_body(doc,
    "Total exclusions in current run: 25 companies. Greenwashing exclusions are pending "
    "completion of the RAG Operator screening phase and may add 1–3 further removals.",
    italic=True, color=GREY_TEXT, space_after=10)

# A.5 agent summaries
add_heading(doc, "A.5  Agent Descriptions", level=2)

agents = [
    ("Agent 1 — Mandate (01_mandate.ipynb)",
     "Defines the fund's investment thesis, screening philosophy, sector exclusion rules, and composite "
     "score weights. Outputs mandate.json, which is read by Agents 8, 11, and 12 to ensure all "
     "downstream decisions are anchored to the agreed mandate.",
     "outputs/scores/mandate.json"),

    ("Agent 2 — Data Ingestion (02_data_ingestion.ipynb)",
     "Loads all four professor-provided CSV files and merges them on the Bloomberg company identifier "
     "(idBbCompany). Downloads five years of adjusted closing prices via yfinance for all 167 companies "
     "and caches them with a date stamp. Creates a ticker bridge column (yf_ticker) mapping Bloomberg "
     "tickers to Yahoo Finance exchange-qualified tickers. Produces the master dataset consumed by all "
     "downstream agents.",
     "outputs/scores/master_dataset_2026-05-12.csv\ndata/market/prices_2026-05-12.csv"),

    ("Agent 3 — Data Quality (03_data_quality.ipynb)",
     "Audits the master dataset before analysis. Calculates coverage rates for every column, flags "
     "columns below 50% coverage, runs IQR-based outlier detection on numerical ESG variables, and "
     "notes the EU Taxonomy eligibility vs alignment distinction. Produces a data dictionary "
     "classifying each variable by extraction method.",
     "outputs/scores/data_dictionary_*.csv\noutputs/scores/outliers_*.csv"),

    ("Agent 10 — Financial Analysis (04_financial_analysis.ipynb)",
     "Calculates five-year price-based risk/return metrics for all 167 companies: annualised return, "
     "annualised volatility, Sharpe ratio (risk-free rate = 0%), and maximum drawdown. Applies the "
     "first hard filter — the 40% volatility cap. No companies were excluded by this cap in the "
     "current run. Also fetches quality metrics (ROE, debt-to-equity, revenue growth) via the yfinance "
     ".info() endpoint. Beta calculation against the STOXX 600 currently returns NaN due to a "
     "benchmark ticker download failure; the volatility cap provides primary risk control.",
     "outputs/scores/financial_metrics_2026-05-12.csv\noutputs/portfolio/financial_exclusions_2026-05-12.csv"),

    ("Agents 5 & 6 — ESG + Climate (05_esg_climate.ipynb)",
     "Builds Environmental, Social, and Governance scores (0–100) for all 167 companies using "
     "SASB sector-adjusted materiality weights. Each ESG variable is min-max normalised within "
     "its pillar, then combined using sector-specific pillar weights (e.g. Energy: E=55%, "
     "Technology: G=40%). Runs ESG triangulation across two independent sources (Bloomberg "
     "Disclosure Score and Sustainalytics Risk Rating) using a 2-of-2 rule. Merges carbon intensity "
     "data from the master dataset. Five companies with all pillars NaN are flagged LOW_DATA "
     "and excluded.",
     "outputs/scores/esg_scores_2026-05-14.csv"),

    ("Agent 7 — Biodiversity (07_biodiversity.ipynb)",
     "Assigns a nature-risk score and tier (Low / Medium / High) to each company using two "
     "sector-level proxy frameworks: ENCORE (ecosystem service dependency, 0–5 scale) and "
     "WRI Aqueduct (water stress exposure, 0–5 scale). Scores are combined into a composite "
     "biodiversity_score (0–100). Very High nature risk is a hard exclusion trigger. All 167 "
     "companies fall in the Low or Medium tiers in the current universe.",
     "outputs/scores/biodiversity_scores_2026-05-14.csv"),

    ("Agent 8 — EU Regulation (08_eu_regulation.ipynb)",
     "Applies EU regulatory screening: EU Taxonomy eligibility vs alignment analysis, SFDR "
     "Article 8 fund-level compliance documentation, and PAI indicator coverage table. "
     "Critical design rule: taxonomy eligibility is never conflated with reported alignment. "
     "73 of 167 companies have eligibility data; only 8 of 167 have alignment data.",
     "outputs/scores/eu_regulation_2026-05-14.csv\noutputs/scores/pai_indicators_2026-05-14.csv\noutputs/scores/sfdr_compliance_2026-05-14.csv"),

    ("Agent 4 — Document Intelligence (06_document_intelligence.ipynb)",
     "Imports and structures ESG data extracted from PDF sustainability reports by the RAG "
     "Operator using Claude Projects. Acts as the bridge between manual PDF analysis and the "
     "automated pipeline. Claude Projects operates as a closed-domain retrieval system: it "
     "answers only from uploaded PDFs, preventing hallucination and ensuring every claim is "
     "traceable to a specific page.",
     "outputs/scores/rag_extractions_*.csv"),

    ("Agent 9 — Greenwashing (09_greenwashing.ipynb)",
     "Applies the 8-dimension greenwashing test to all portfolio holdings. Reads the RAG "
     "Operator's completed Excel workbook, converts entries to structured JSON, and scores "
     "each company on: Specificity, Metric, Baseline, Target, Time Horizon, Scope, External "
     "Validation, and Behavioural Consistency. Companies with HIGH on 3+ dimensions are "
     "excluded; HIGH on exactly 2 triggers watchlist status. Verification protocol: 30% "
     "of all extractions are manually verified; 100% of exclusion-driving findings are "
     "independently confirmed.",
     "outputs/scores/greenwashing_scores_*.csv (pending RAG Operator)"),

    ("Agent 11 — Portfolio Construction (10_portfolio_construction.ipynb)",
     "The convergence point of the pipeline. Merges all five upstream score files, applies "
     "all exclusion rules in sequence, ranks companies by composite score, and selects the "
     "final 20 holdings. Weights are equal-weighted across the qualifying universe (1/N), "
     "capped at 10% per holding, with excess redistributed proportionally. Designed for "
     "graceful degradation: if any upstream file is absent the notebook warns and continues "
     "with available data.",
     "outputs/portfolio/final_portfolio_2026-05-14.csv\noutputs/portfolio/universe_scores_2026-05-14.csv\noutputs/portfolio/exclusions.csv"),

    ("Agent 12 — Human Review (11_human_review.ipynb)",
     "Provides a structured interface for the investment team to document manual override "
     "decisions. Every override requires: ticker, action (include/exclude), original model "
     "decision, override decision, written rationale (minimum two sentences), and "
     "decision-maker name. A minimum of three documented overrides is required for Q&A "
     "defence. Also generates the final AI Use Statement.",
     "outputs/portfolio/human_overrides_*.csv"),

    ("Agent 13 — Reporting (12_reporting.ipynb)",
     "Generates all presentation-ready charts from the final portfolio: (1) horizontal bar "
     "chart of portfolio weights by company; (2) ESG pillar comparison — portfolio vs full "
     "universe; (3) sector allocation pie chart. Prints the portfolio factsheet summary "
     "(ESG 70.8/100, Sharpe 0.776, WACI 431.6 tCO₂e/€m rev, 7 sectors, 20 holdings). "
     "All outputs are PNG files for direct insertion into the presentation deck.",
     "outputs/reports/portfolio_weights.png\noutputs/reports/esg_comparison.png\noutputs/reports/sector_allocation.png"),
]

for title, body, outputs in agents:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = MID_BLUE
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    add_body(doc, body, space_after=3)
    add_body(doc, f"Key outputs: {outputs}", italic=True, color=GREY_TEXT, space_after=8)

# A.6 hallucination controls
add_heading(doc, "A.6  Hallucination Controls & Data Integrity", level=2)
make_table(doc,
    headers=["Control", "Where Applied"],
    rows=[
        ("Professor CSVs never modified — all outputs written to outputs/ folder", "Agent 2 onwards"),
        ("MISSING label used when data is absent — AI never invents a figure", "Agent 9 / RAG Operator"),
        ("30% random sample of RAG extractions manually verified against source PDF with page reference", "Agent 9 / RAG Operator"),
        ("100% of watchlist and exclusion-driving extractions independently verified before flag applied", "Agent 9 / RAG Operator"),
        ("ESG ratings treated as indicators, not objective truth; triangulation required across 2 sources", "Agents 5/6, 11"),
        ("Taxonomy eligibility never conflated with reported alignment — explicitly separated in code and report", "Agent 8"),
        ("All data vintage-tagged with download date on filename and in metadata column", "All agents"),
        ("Sustainalytics scores marked NaN where unavailable — not estimated or imputed", "Agents 5/6"),
    ],
    col_widths_cm=[10.5, 6.5],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 2 — REPORT SECTION 3: UNIVERSE CONSTRUCTION & DATA SOURCES
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Section 3: Universe Construction & Data Sources", level=1)

add_body(doc,
    "This section describes how the investable universe was defined, the data sources used to "
    "characterise it, and the data engineering steps required to merge heterogeneous inputs into "
    "a single analytical master dataset.",
    space_after=10)

add_heading(doc, "3.1  Universe Definition", level=2)
add_body(doc,
    "The investable universe is prescribed by the course assignment as the top 170 constituents "
    "of the STOXX Europe 600 index ranked by 10-year total return. This subset was provided as a "
    "pre-screened dataset by the course instructor. After data cleaning and merging, 167 companies "
    "were retained; three companies could not be matched across all four source files and were "
    "dropped before analysis began.",
    space_after=6)
add_body(doc,
    "The universe spans eight BICS Level-1 sectors — Technology, Financials, Industrials, "
    "Consumer Discretionary, Materials, Energy, Health Care, and Real Estate — and covers "
    "twelve Western European countries, with Denmark, Sweden, the Netherlands, and Germany "
    "most heavily represented. All companies carry Bloomberg exchange code GR, consistent with "
    "their primary European listing.",
    space_after=6)
add_body(doc,
    "Known bias disclosure: selecting the top 170 STOXX Europe 600 constituents by 10-year "
    "historical return introduces two structural biases. Survivorship bias arises because only "
    "companies that survived and remained in the index are included, systematically overstating "
    "historical returns. Look-ahead bias arises because using 10-year backward return to define "
    "the universe relies on information unavailable at the start of the measurement period. Both "
    "biases are disclosed here and in the limitations section. In a live mandate the full STOXX "
    "Europe 600 would serve as the starting universe.",
    italic=True, color=GREY_TEXT, space_after=10)

add_heading(doc, "3.2  Data Sources", level=2)
add_body(doc,
    "Four course-provided datasets and one market data feed form the foundation of the pipeline. "
    "All four CSV files are loaded read-only; the pipeline never modifies them.",
    space_after=6)
make_table(doc,
    headers=["File", "Contents", "Rows", "Key identifier"],
    rows=[
        ("equityBicsV2.csv",                         "Company identifiers: name, ticker, ISIN, country, BICS sector hierarchy (Levels 1–4)", "167", "idBbCompany"),
        ("esgEnvironmentalSocialConsolidatedV4.csv",  "Quantitative Environmental and Social metrics: Scope 1–3 GHG emissions, water usage, waste generation, workforce safety, diversity ratios", "167", "idBbCompany"),
        ("esgGovernanceConsolidatedV4.csv",           "Governance metrics: board gender diversity, executive compensation ratios, audit committee independence, anti-corruption policies", "167", "idBbCompany"),
        ("legalEntityEuTaxonomy.csv",                 "EU Taxonomy eligibility and estimated alignment percentages, DNSH indicator estimates (Objectives 1–4), green revenue proxy", "167", "idBbCompany"),
        ("yfinance (API)",                            "Daily adjusted closing prices, 1 January 2020 – 1 January 2025. Also provides ROE, debt-to-equity, revenue growth via .info() endpoint.", "166/167", "Yahoo Finance ticker"),
    ],
    col_widths_cm=[4.5, 7.5, 1.5, 3.5],
)

add_heading(doc, "3.3  Data Integration & Ticker Bridge", level=2)
add_body(doc,
    "A critical data engineering challenge is that the course-provided files use Bloomberg tickers "
    "(e.g. ASME for ASML Holding), while the yfinance market data API requires exchange-qualified "
    "Yahoo Finance tickers (e.g. ASML.AS). These ticker formats do not share a common syntax, so a "
    "manual mapping column (yf_ticker) was added to the master dataset during the data ingestion step. "
    "All downstream agents load this bridge column to ensure ESG scores (Bloomberg tickers) and "
    "financial metrics (Yahoo Finance tickers) can be joined without row loss.",
    space_after=6)
add_body(doc,
    "The four CSV files are merged sequentially on idBbCompany — the Bloomberg Global Company ID — "
    "using a left join anchored on equityBicsV2.csv. The result is a single master dataset of "
    "167 rows × 677 columns, vintaged 2026-05-12. All agent output files trace their lineage back "
    "to this master dataset.",
    space_after=10)

add_heading(doc, "3.4  Market Data", level=2)
add_body(doc,
    "Historical price data was downloaded via yfinance for all 167 Bloomberg tickers mapped to "
    "their Yahoo Finance equivalents. The download window is 1 January 2020 to 1 January 2025, "
    "providing a five-year observation period that captures two distinct market regimes: the "
    "COVID-19 drawdown and recovery (2020), the 2022 rate-rising environment, and the subsequent "
    "growth rebound (2023–2024). Prices are adjusted for dividends and splits.",
    space_after=6)
add_body(doc,
    "166 of 167 companies received price data. Sydbank A/S (Bloomberg ticker TM2) has no "
    "corresponding Yahoo Finance ticker and was excluded from the financial metrics calculation. "
    "Price data is cached locally as prices_2026-05-12.csv with a date stamp to ensure "
    "reproducibility and avoid repeated API calls.",
    space_after=6)
add_body(doc,
    "Note on GRANOLAS: of the seven GRANOLAS stocks (GSK, Roche, Nestlé, Novo Nordisk, L'Oréal, "
    "LVMH, SAP, ASML, AstraZeneca, Novartis, Sanofi), only ASML Holding (ASME) appears in the "
    "universe. The remaining GRANOLAS are STOXX Europe 600 members but their 10-year returns "
    "ranked outside the top 170 in the course dataset. AstraZeneca and Novartis are present in "
    "the universe but ranked below the composite score threshold required for final portfolio selection.",
    space_after=6)

make_table(doc,
    headers=["Metric", "Value"],
    rows=[
        ("Universe size (post merge)",              "167 companies"),
        ("Companies with price data",               "166 / 167 (Sydbank excluded)"),
        ("Price observation window",                "2020-01-01 to 2025-01-01 (5 years)"),
        ("Price data vintage",                      "2026-05-12"),
        ("ESG / biodiversity / EU data vintage",    "2026-05-14"),
        ("Master dataset dimensions",               "167 rows × 677 columns"),
        ("BICS Level-1 sectors covered",            "8"),
        ("Countries represented",                   "12 Western European"),
    ],
    col_widths_cm=[7, 10],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 3 — REPORT SECTION 4: QUANTITATIVE FINANCIAL SCREENING
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Section 4: Quantitative Financial Screening", level=1)
add_body(doc,
    "The pipeline applies two complementary financial screens before ESG analysis begins: a "
    "price-based efficiency screen that feeds directly into the composite ranking, and a "
    "fundamental quality framework used as a qualitative overlay during human review. Running "
    "financial screens first prevents ESG computation being wasted on companies that would be "
    "excluded for financial reasons regardless of their sustainability profile.",
    space_after=10)

add_heading(doc, "4.1  Screen A — Price-Based Efficiency (Automated)", level=2)
add_body(doc,
    "The price-based screen uses the five-year adjusted price series (2020–2025) to compute four "
    "metrics for each company. These metrics collectively capture how efficiently the market has "
    "priced each stock relative to its risk over the observation window.",
    space_after=6)
make_table(doc,
    headers=["Metric", "Formula", "Interpretation"],
    rows=[
        ("Annualised return",     "Geometric mean of daily returns, scaled to 252 trading days",        "Total return per year; higher is better"),
        ("Annualised volatility", "Standard deviation of daily returns × √252",                         "Risk proxy; capped at 40% as hard exclusion filter"),
        ("Sharpe ratio",          "(Annualised return − 0%) ÷ Annualised volatility",                   "Return per unit of risk; risk-free rate set to 0%"),
        ("Maximum drawdown",      "Largest peak-to-trough decline over the 5-year window",              "Downside risk severity; informational only"),
    ],
    col_widths_cm=[4, 6, 7],
)
add_body(doc,
    "The volatility cap (>40% annualised volatility) is the only hard filter in Screen A. "
    "No company in the current universe was excluded by this cap — all 164 companies with "
    "available price data fell below the 40% threshold. Three companies (missing Yahoo Finance "
    "tickers) could not be scored and received NaN financial metrics.",
    space_after=6)
add_body(doc,
    "The Sharpe ratio is normalised to a 0–100 scale by percentile rank within the 164 scored "
    "companies, producing the sharpe_score variable. This normalised score is the financial "
    "component that enters the composite ranking formula (weighted at 24% of the total composite). "
    "The raw Sharpe ratio is retained separately for the portfolio factsheet.",
    space_after=6)
make_table(doc,
    headers=["Metric", "Value"],
    rows=[
        ("Companies entering Screen A",         "166 (all with Yahoo Finance price data)"),
        ("Companies successfully scored",        "164 (3 missing yfinance data, 1 no YF ticker)"),
        ("Excluded by volatility cap (>40%)",    "0"),
        ("Beta vs STOXX 600 benchmark",          "All NaN — STOXX 600 ETF (EXW1.DE) download failed"),
        ("Weighted portfolio Sharpe ratio",      "0.776 (final 20 holdings)"),
        ("Highest individual Sharpe",            "Games Workshop Group (G7W): 1.123"),
        ("Lowest individual Sharpe (in portfolio)", "Norsk Hydro (NOH1): 0.543"),
    ],
    col_widths_cm=[7, 10],
)

add_heading(doc, "4.2  Screen B — Fundamental Quality Framework (Qualitative Overlay)", level=2)
add_body(doc,
    "The pipeline also incorporates a six-factor fundamental quality framework developed during "
    "the methodology design phase. This framework assesses accounting-based quality signals that "
    "are complementary to price-based metrics: while Screen A captures how the market has priced "
    "a stock historically, Screen B captures the underlying financial health of the business. "
    "In the current build, Screen B is applied as a qualitative overlay during human review "
    "(Agent 12) rather than as an automated exclusion filter.",
    space_after=6)
make_table(doc,
    headers=["Factor", "Metric", "Weight", "What it tests"],
    rows=[
        ("M-01", "ROIC minus WACC Spread",        "22%", "Economic value creation — does the business earn above its cost of capital?"),
        ("M-02", "Free Cash Flow Conversion",      "22%", "Earnings quality — how much reported profit converts to actual cash?"),
        ("M-03", "FCCR + Net Debt / EBITDA",       "18%", "Financial resilience — can the company service its debt without stress?"),
        ("M-04", "Sloan Accruals Ratio",            "13%", "Earnings manipulation signal — high accruals may indicate inflated earnings"),
        ("M-05", "EBITDA Margin Coefficient of Variation", "13%", "Earnings stability — consistent margins signal a durable business model"),
        ("M-06", "Dividend Sustainability Index",   "12%", "Capital return reliability — are dividends covered by cash flow and balance sheet?"),
    ],
    col_widths_cm=[1.5, 5, 2, 8.5],
)
add_body(doc,
    "The weights reflect the relative importance of each dimension within the fundamental quality "
    "assessment: value creation (M-01) and cash quality (M-02) together account for 44% of the "
    "score, reflecting their primacy as indicators of durable competitive advantage. Resilience "
    "(M-03) at 18% captures leverage risk. The remaining 38% covers quality-of-earnings and "
    "sustainability of returns.",
    space_after=6)
add_body(doc,
    "Note on data availability: the Bloomberg yfinance .info() endpoint provides ROE, "
    "debt-to-equity, and revenue growth for most companies, but ROIC, WACC, FCF conversion, "
    "and Sloan accruals require Bloomberg terminal or Compustat-quality data not available in "
    "the course dataset. These metrics are therefore applied judgementally during human review "
    "for the companies selected for Deep Review (10–12 holdings), rather than automated across "
    "all 167 companies.",
    italic=True, color=GREY_TEXT, space_after=10)

add_heading(doc, "4.3  Interaction Between Screens A and B", level=2)
add_body(doc,
    "Screen A (automated) and Screen B (qualitative) are intentionally designed as complements "
    "rather than substitutes. Screen A asks: how has this stock been priced relative to its risk? "
    "Screen B asks: is the underlying business of high accounting quality? A company can score "
    "well on one and poorly on the other — for example, a business with superb fundamentals "
    "that has experienced short-term market re-rating (low Sharpe), or a business with strong "
    "recent price performance driven by multiple expansion rather than earnings growth (high Sharpe, "
    "weak FCF conversion).",
    space_after=6)
add_body(doc,
    "In the portfolio construction step, Screen A contributes directly to the composite score "
    "via sharpe_score. Screen B findings are incorporated through the human override mechanism: "
    "where a Deep Review company is flagged by Screen B as having weak fundamental quality, the "
    "investment team may apply a negative override regardless of composite ranking. All such "
    "overrides are logged with written rationale in the Human Review notebook.",
    space_after=10)

add_heading(doc, "4.4  Financial Screen Results Summary", level=2)
make_table(doc,
    headers=["Stage", "Companies remaining"],
    rows=[
        ("Starting universe",                     "167"),
        ("After price data retrieval",            "166 (Sydbank excluded — no YF ticker)"),
        ("After financial metrics calculation",   "164 (3 missing yfinance data)"),
        ("After volatility cap (>40%)",           "164 (0 removed)"),
        ("After subsidiary deduplication",        "160 (4 parent/sub pairs resolved)"),
        ("After ESG no-data exclusion",           "155 (5 LOW_DATA companies removed)"),
        ("After ESG quality floor (score < 50.1)","139 (16 removed)"),
        ("Final investable universe for ranking", "139 companies"),
        ("Final portfolio selected",              "20 holdings"),
    ],
    col_widths_cm=[8, 9],
)

doc.add_paragraph()
add_body(doc,
    "Note: The financial screen contributes to the composite score ranking but is not the primary "
    "exclusion mechanism in the current build. The ESG quality floor (removing the bottom 16 "
    "companies) has greater numerical impact on universe size. This is consistent with the "
    "fund's sustainability-first mandate, in which ESG quality is the dominant ranking criterion "
    "(60% weight) and financial efficiency plays a supporting role (40%).",
    italic=True, color=GREY_TEXT, space_after=4)

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
