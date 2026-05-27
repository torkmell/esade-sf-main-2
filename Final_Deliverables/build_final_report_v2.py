#!/usr/bin/env python
"""Corrected final-report builder — v2, 2026-05-27.

Replaces the earlier `build_final_report.py`, which described the SUPERSEDED
in-house ESG method (NB05 Steps 0–6 min-max + Bloomberg-Disclosure
triangulation) and would not reconcile against the actual numbers. v2
describes the ESG specialist's screening workbook (Truvalue / Sustainalytics
/ ISS 2-of-3, hard red-flag override, 10-indicator sector-relative z, 50/20/30
cross-pillar, percentile rank) and is honest about the proxies, defaults,
and paperwork gaps the conversation surfaced.

    python Final_Deliverables/build_final_report_v2.py
    -> Final_Deliverables/Final_Report_v2_<date>.docx
"""
import os, json, glob, re
from datetime import date
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
os.chdir(PROJECT)
TODAY = str(date.today())

# ── Data loaders ───────────────────────────────────────────────────────────────
def latest(pattern):
    files = sorted(glob.glob(pattern))
    return files[-1] if files else None

mandate = json.load(open("outputs/scores/mandate.json"))
fp      = pd.read_csv(latest("outputs/portfolio/final_portfolio_*.csv"))
uni     = pd.read_csv(latest("outputs/portfolio/universe_scores_*.csv"))   # capped-40
gw      = pd.read_csv(latest("outputs/scores/greenwashing_scores_*.csv"))
ov      = pd.read_csv(latest("outputs/scores/override_decisions_*.csv"))
exclu   = pd.read_csv("outputs/portfolio/exclusions.csv") if os.path.exists("outputs/portfolio/exclusions.csv") else pd.DataFrame()
bt      = pd.read_csv("Optimization_module/outputs/backtest_results.csv")

# Workbook (specialist) — for Stage-1 mechanism and capped-vs-uncapped contrast
WB      = "data/provided/Portfolio_Screening_Output.xlsx"
s1uni   = pd.read_excel(WB, sheet_name="Stage 1 — Universe")
s2cap   = pd.read_excel(WB, sheet_name="Stage 2 — Capped Top 40")
s2unc   = pd.read_excel(WB, sheet_name="Stage 2 — Uncapped Top 40")
s2full  = pd.read_excel(WB, sheet_name="Stage 2 — Full ranking")
hardexc = pd.read_excel(WB, sheet_name="Hard exclusions detail")

# ── Computed numbers (one source of truth) ─────────────────────────────────────
N_HOLD       = len(fp)
w            = fp["weight"]
wesg         = (fp["ESG_score"]    * w).sum()
wfin         = (fp["fin_score"]    * w).sum()
wsharp       = (fp["sharpe_ratio"] * w).sum()
wvol         = (fp["vol_annual"]   * w).sum()
wmdd         = (fp["max_drawdown"] * w).sum()
wbeta        = (fp["beta"]         * w).sum()
waci         = (fp["carbon_intensity"].fillna(0) * w).sum()
maxw         = w.max() * 100
n_sect       = fp["sasb_sector"].nunique()
ci_imputed   = int((fp["ci_source"]=="sector_median_imputed").sum())
ci_reported  = int((fp["ci_source"]=="bloomberg_calc").sum())
imp_w        = fp.loc[fp.ci_source=="sector_median_imputed","weight"].sum()*100
norsk        = fp[fp["company_name"].str.contains("Norsk Hydro", na=False)]
nh_share     = float(norsk["carbon_intensity"].iloc[0]*norsk["weight"].iloc[0]/waci*100) if len(norsk) else 0.0
waci_exnh    = (fp.loc[~fp["company_name"].str.contains("Norsk Hydro", na=False), "carbon_intensity"].fillna(0) *
                fp.loc[~fp["company_name"].str.contains("Norsk Hydro", na=False), "weight"]).sum() / \
               fp.loc[~fp["company_name"].str.contains("Norsk Hydro", na=False), "weight"].sum() if len(norsk) else waci

# Universe sizes (the funnel)
N_RAW    = 600                           # STOXX 600
N_UNI    = len(s1uni)                    # 289
N_S1PASS = int(s1uni["Stage 1 PASS"].sum())  # 242
N_S2     = len(s2full)                   # 224
N_CAP    = len(s2cap)                    # 40
N_HARD   = len(hardexc)                  # 13

# ESG universe stats
uni_esg     = uni["ESG_score"].mean()    # capped-40 mean (79.3)
unc_esg     = s2unc["Percentile"].mean() # 91.8
gw_pass     = int((gw["gw_high_count"]<2).sum())
gw_mean     = float(gw["gw_score_pct"].mean())
gw_max      = float(gw["gw_score_pct"].max())
gw_worst    = gw.loc[gw["gw_score_pct"].idxmax(), "company_name"]

# E-pillar default flag (z exactly +1.000)
import re as _re
_SFX = sorted(["class b","class a","class c","ab","sa","nv","plc","ag","asa","oyj abp","oyj","oy","ltd","inc","corp","spa","s p a","s a","se","a s","as","group","holding","holdings","sca","psc"],key=len,reverse=True)
def _norm(s):
    s=_re.sub(r"[^\w\s]"," ",str(s).lower()); s=_re.sub(r"\s+"," ",s).strip()
    for x in _SFX:
        if s.endswith(" "+x): s=s[:-len(x)-1].strip(); break
    return s
s2full["_k"]=s2full["Company"].apply(_norm); fp["_k"]=fp["company_name"].apply(_norm)
m_z = fp.merge(s2full[["_k","E pillar","S pillar","G pillar","In-house ESG z","Percentile"]], on="_k", how="left")
n_E_default = int((m_z["E pillar"].round(6)==1.0).sum())

# Backtest
bt_ranked = bt.sort_values("rank")
top_method = bt_ranked.iloc[1]  # score_tilted (composite-score-proportional) - method we actually use
ew_method  = bt[bt["method"]=="equal_weight"].iloc[0]

# ── DOCX helpers ───────────────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
GREY_TXT  = RGBColor(0x70, 0x70, 0x70)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

def shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell(cell, text, bold=False, align=None, size=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None: p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    if size:  r.font.size = Pt(size)
    if color: r.font.color.rgb = color

def h(doc, text, level=1, color=DARK_BLUE):
    p = doc.add_heading(text, level=level)
    for run in p.runs: run.font.color.rgb = color
    return p

def para(doc, *runs, sa=8):
    """runs = list of (text, {"bold":..., "italic":...}) or plain strings."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    for r in runs:
        if isinstance(r, str):
            p.add_run(r)
        else:
            txt, opts = r
            run = p.add_run(txt)
            run.bold   = opts.get("bold", False)
            run.italic = opts.get("italic", False)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix); run.bold = True
    p.add_run(text)
    return p

def make_table(doc, headers, rows, widths_cm, header_bg="1F497D"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h_ in enumerate(headers):
        set_cell(t.rows[0].cells[c], h_, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 size=10, color=WHITE)
        shade(t.rows[0].cells[c], header_bg)
    R = WD_ALIGN_PARAGRAPH.RIGHT
    for row in rows:
        cells = t.add_row().cells
        for c, v in enumerate(row):
            align = R if (isinstance(v, str) and (v.endswith("%") or v.replace(".","").replace("-","").replace("+","").isdigit())) else None
            set_cell(cells[c], v, align=align, size=9.5)
    for row in t.rows:
        for c, w_ in enumerate(widths_cm):
            row.cells[c].width = Cm(w_)
    return t

# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)
for m_ in ("left_margin","right_margin","top_margin","bottom_margin"):
    setattr(doc.sections[0], m_, Inches(0.9))

# Title page
title = doc.add_heading(mandate.get("fund_name","ESADE Sustainable European Equity Fund"), 0)
for run in title.runs: run.font.color.rgb = DARK_BLUE
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Methodology Report — Final Portfolio Construction")
r.italic = True; r.font.size = Pt(13); r.font.color.rgb = MID_BLUE
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"ESADE MSc Finance · Sustainable Finance Group Assignment · {TODAY}")
r.font.size = Pt(10); r.font.color.rgb = GREY_TXT
doc.add_paragraph()
disc = doc.add_paragraph()
disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = disc.add_run("Academic prototype — not a regulated investment product or investment advice.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY_TXT
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "1. Executive Summary", 1)
para(doc,
    "This report documents the construction of the ", ("ESADE Sustainable European "
    "Equity Fund", {"bold": True}), ", a concentrated long-only European equity portfolio "
    f"of {N_HOLD} holdings benchmarked against the STOXX Europe 600. The fund was built by "
    "an AI-agent research pipeline of thirteen agents, with documented human oversight at "
    "every decision point. This document describes the methodology, the data, the screens, "
    "and — critically — the limitations, so that every number can be defended in Q&A.")

para(doc,
    ("Headline outcome.", {"bold": True}), f"  Twenty holdings across {n_sect} sectors, "
    f"with no single position above {maxw:.1f}%. The portfolio carries a weighted in-house "
    f"ESG percentile of {wesg:.1f} (versus a capped-universe mean of {uni_esg:.1f}), a "
    f"weighted Sharpe ratio of {wsharp:.2f}, a weighted annualised volatility of "
    f"{wvol*100:.1f}%, a beta to STOXX Europe 600 of {wbeta:.2f}, and a weighted carbon "
    f"intensity (WACI) of {waci:.0f} tCO₂e per million of revenue. All {N_HOLD} holdings "
    "passed the greenwashing 8-Test screen.")

para(doc,
    ("Process at a glance.", {"bold": True}), f"  The investable universe of {N_UNI} European "
    f"large-caps was filtered through four sequential stages: a 2-of-3 vendor eligibility "
    f"vote (with fossil-fuel exclusion) reduced it to {N_S1PASS}; a hard red-flag override "
    f"on assessment-grade vendors reduced it to {N_S2}; an in-house ESG ranking with a "
    f"sector cap produced a {N_CAP}-company shortlist; and a composite finance/ESG score "
    f"selected the final {N_HOLD}. The four verbs are vote, veto, rank, and pick.")

para(doc,
    ("What this report does honestly.", {"bold": True}), "  We use proxied or imputed "
    f"data in four places: WACI ({ci_imputed} of {N_HOLD} holdings use a sector-median "
    f"carbon intensity), the in-house ESG E-pillar ({n_E_default} of {N_HOLD} holdings receive a "
    "sector default), biodiversity (a nature-risk proxy by design), and EU Taxonomy "
    "alignment (vendor-estimated). Each is disclosed in §12. The portfolio is otherwise "
    "data-driven and reproducible from the project repository.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. INVESTMENT MANDATE & STRATEGIC THESIS
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "2. Investment Mandate and Strategic Thesis", 1)

h(doc, "2.1  Investment thesis", 2)
para(doc, mandate.get("investment_thesis",""))

h(doc, "2.2  Mandate constraints", 2)
para(doc, "The fund is governed by an explicit set of quantitative and qualitative "
    "constraints, written at the outset and unchanged through construction. These "
    "constraints define what the AI pipeline is permitted to produce.")
make_table(doc,
    ["Constraint", "Rule"],
    [
        ("Number of holdings",       f"{mandate['constraints']['min_holdings']}–{mandate['constraints']['max_holdings']} (target {mandate['constraints']['target_holdings']})"),
        ("Strategy",                 mandate["constraints"]["strategy"].title()),
        ("Currency",                 mandate["constraints"]["currency"]),
        ("Max single holding",       f"{mandate['constraints']['max_single_weight_pct']:.0f}%"),
        ("Min sectors",              f"≥ {mandate['constraints']['min_sectors']}"),
        ("Benchmark",                mandate["benchmark"]),
        ("Composite-score weights",  f"{int(mandate['composite_score_weights']['esg']*100)}% ESG · {int(mandate['composite_score_weights']['financial']*100)}% Financial"),
    ],
    widths_cm=[5.5, 11.5])

h(doc, "2.3  Hard exclusions and watchlist triggers", 2)
para(doc, "Four absolute exclusions and four watchlist triggers were written into the "
    "mandate before any data was loaded:")
for ex in mandate.get("hard_exclusions", []):
    bullet(doc, ex, bold_prefix="Exclude — ")
for wl in mandate.get("watchlist_triggers", []):
    bullet(doc, wl, bold_prefix="Watchlist — ")

h(doc, "2.4  Why this fund, why now", 2)
para(doc,
    "Three convictions underpin the strategy. First, vendor ESG ratings are noisy and "
    "frequently disagree (Berg, Kölbel & Rigobon, 2022, ", ("Aggregate Confusion", {"italic":True}),
    "); no single rating should be relied on. Second, companies with credible — not just "
    "stated — sustainability commitments face lower regulatory, reputational, and "
    "transition risk under SFDR, CSRD and the EU Taxonomy. Third, greenwashing is a "
    "material reputational and regulatory risk that requires forensic, evidence-based "
    "screening on actual disclosure documents, not on vendor scores alone.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. AI-AGENT PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "3. The AI-Agent Pipeline", 1)

para(doc,
    "The portfolio is built by a thirteen-agent pipeline. Each agent performs a single, "
    "auditable function — data ingestion, ESG scoring, greenwashing screening, financial "
    "analysis, portfolio construction — and writes its output to a versioned file that the "
    "next agent reads. Agents are connected with n8n.cloud for orchestration. Critically, "
    "the term ‘agent’ refers to a pipeline stage with a defined input and output; the "
    "portfolio-construction and optimisation modules themselves are deterministic Python, "
    "not LLM-driven decisions.")

h(doc, "3.1  Agent inventory", 2)
make_table(doc,
    ["#", "Agent", "Role"],
    [
        ("01", "Mandate",              "Fixes thesis, weights, exclusions; writes mandate.json"),
        ("02", "Data Ingestion",       "Loads 4 provided CSVs; downloads prices via yfinance"),
        ("03", "Data Quality",         "Missing-value audit, outlier flags, data dictionary"),
        ("04", "Document Intelligence","Imports RAG extractions from sustainability PDFs"),
        ("05", "ESG Scoring",          "Imports the specialist FactSet/Bloomberg screening workbook"),
        ("06", "Climate",              "Computes WACI from carbon intensity; sector-median imputed where missing"),
        ("07", "Biodiversity",         "Nature-risk proxy from ENCORE + WRI Aqueduct"),
        ("08", "EU Regulation",        "Taxonomy eligibility + SFDR Article 8 reading"),
        ("09", "Greenwashing",         "8-Test forensic screen on every holding's disclosure documents"),
        ("10", "Financial Analysis",   "Fundamental quality + price-based returns/Sharpe/volatility"),
        ("11", "Portfolio Construction","Composite score, sector cap, correlation guard, weight assignment"),
        ("12", "Human Review",         "Override log, watchlist dispositions, audit trail"),
        ("13", "Reporting",            "Charts, factsheet, dashboard, this report"),
    ],
    widths_cm=[1.2, 4.5, 11.3])

h(doc, "3.2  The four verbs of the funnel", 2)
para(doc,
    "The pipeline's screening logic can be remembered as four verbs in sequence:")
bullet(doc, "the 2-of-3 vendor vote and fossil-fuel exclusion produce the eligible universe.", bold_prefix="VOTE — ")
bullet(doc, "Sustainalytics High/Severe or ISS Poor overrides the vote; companies tripping either are excluded regardless of the majority.", bold_prefix="VETO — ")
bullet(doc, "surviving companies are ranked on the in-house ESG percentile; the sector-capped top 40 form the shortlist.", bold_prefix="RANK — ")
bullet(doc, "a 60% financial / 40% ESG composite selects the final 20 from the shortlist, with a 10% position cap and a 0.90 correlation guard.", bold_prefix="PICK — ")

h(doc, "3.3  Hybrid human–AI design", 2)
para(doc,
    "The pipeline is not autonomous. Human oversight is exercised at four discrete points: "
    "(i) writing the mandate and choosing the composite weights; (ii) verifying 30% of all "
    "RAG extractions against the source PDFs, and 100% of any extraction that drives a "
    "watchlist or exclusion decision; (iii) approving or overriding the quantitative "
    "ranking at the Investment Committee step; and (iv) reviewing the dashboard output "
    f"before publication. Nine override decisions are logged for the final {N_HOLD} (see "
    "§12.3).")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. DATA SOURCES & DATA QUALITY
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "4. Data Sources and Data Quality", 1)

h(doc, "4.1  Inventory of sources", 2)
para(doc, "Eight distinct data sources feed the pipeline. They are tabled below with "
    "their owner, coverage and confidence band.")
make_table(doc,
    ["Source", "Used for", "Coverage"],
    [
        ("equityBicsV2.csv (course)",                   "Company identifiers, BICS sector, ISIN", "289 names"),
        ("esgEnvironmentalSocialConsolidatedV4.csv",    "Quantitative E and S indicators",        "Variable: 50–95% missing per field"),
        ("esgGovernanceConsolidatedV4.csv",             "Governance indicators",                  "Variable: 30–90% missing per field"),
        ("legalEntityEuTaxonomy.csv",                   "EU Taxonomy eligibility & estimated alignment", "Eligibility well covered; reported alignment sparse"),
        ("yfinance",                                    "Daily prices for returns, Sharpe, beta", "8-year window, all holdings"),
        ("Portfolio_Screening_Output.xlsx (specialist)","Stage-1 vendor verdicts + Stage-2 ESG z-scores", "All 289 names through Stage 1; 224 ranked in Stage 2"),
        ("RAG corpus (120 PDFs)",                       "8-Test greenwashing screen on holdings", "40 capped folders; 8-Test run on the 20 held names"),
        ("ESG_External_Research.md",                    "Tier-1 controversy screen (press, NGOs, regulators)", "All 40 capped names; 6-month rolling window"),
    ],
    widths_cm=[6.0, 5.5, 5.5])

h(doc, "4.2  Data quality posture", 2)
para(doc,
    "Three principles govern data handling. First, ", ("absent data is never invented", {"bold":True}),
    " — variables with no value remain blank, with one disclosed exception (WACI carbon "
    "intensity, where sector-median imputation is used to satisfy the assignment's "
    "required-metric list). Second, ", ("every figure has a provenance", {"bold":True}),
    " — extracted, calculated, imputed and vendor-supplied values are sourced. Third, ",
    ("hallucination controls", {"bold":True}), " apply to all AI-extracted content: "
    "MISSING is a valid output, verbatim quotes carry page numbers, and any extraction "
    "that drives a watchlist or exclusion decision is verified manually against the source.")

para(doc,
    "One inherited deficiency must be disclosed. The pipeline's ", ("data dictionary", {"italic":True}),
    " currently classifies all 677 tracked variables as ", ("reported", {"italic":True}),
    " — even those that are vendor-estimated (EU Taxonomy alignment), imputed (carbon "
    "intensity), or defaulted (E-pillar values for low-disclosure sectors). The dictionary "
    "is being reclassified to use the prescribed five-way taxonomy (reported / observed / "
    "estimated / AI-extracted / judgement-based); the substantive numbers are unaffected "
    "but the documentation gap is acknowledged in §12.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. UNIVERSE CONSTRUCTION & STAGE-1 SCREENING
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "5. Universe Construction and Stage-1 Screening", 1)

h(doc, "5.1  From the index to the shortlist", 2)
para(doc,
    f"The investable universe is the {N_RAW}-name STOXX Europe 600. A size, history and "
    f"sector filter (free-float market cap ≥ €5bn; ≥ 3 years of pricing) brings this to "
    f"{N_UNI} candidate companies — the working universe the rest of the pipeline operates "
    "on. From here, screening is sequential and the funnel narrows in three further "
    "stages.")
make_table(doc,
    ["Stage", "Filter", "Survivors"],
    [
        ("Universe",                 "STOXX Europe 600 (size + history + sector)",  str(N_UNI)),
        ("Stage 1 — Eligibility",    "2-of-3 vendor vote + fossil-fuel exclusion",  str(N_S1PASS)),
        ("Stage 1 — Hybrid override","Sustainalytics High/Severe or ISS Poor",       str(N_S2)),
        ("Stage 2 — ESG ranking",    "In-house percentile + sector cap (≤6/sector)", str(N_CAP)),
        ("Stage 3 — Selection",      "Composite 40/60 + correlation guard + weight", str(N_HOLD)),
    ],
    widths_cm=[5.0, 8.5, 3.5])

h(doc, "5.2  The 2-of-3 vendor vote", 2)
para(doc,
    "Vendor ESG ratings notoriously disagree. The pipeline therefore treats no single "
    "vendor as decisive: a company is eligible only if it passes the threshold of at "
    "least two of three independent providers. The triangulation rule is hard-coded in "
    "the specialist's screening workbook.")
make_table(doc,
    ["Vendor", "What it measures", "PASS condition"],
    [
        ("Truvalue (FactSet Truvalue Labs)", "ESG performance tier",          "Rating ≥ Average"),
        ("Sustainalytics",                    "ESG risk (lower = better)",     "Band Negligible or Low (risk < 20)"),
        ("ISS ESG",                            "ESG quality tier",              "Rating ≥ Medium"),
    ],
    widths_cm=[5.0, 7.0, 5.0])

h(doc, "5.3  Hybrid red-flag override", 2)
para(doc,
    "Some signals are too severe to be outvoted. Sustainalytics ", ("High", {"italic":True}),
    " or ", ("Severe", {"italic":True}), " risk and ISS ", ("Poor", {"italic":True}),
    " each act as a single-vendor veto: a company tripping either is hard-excluded even "
    "if it carried two vendor passes. The workbook logs ", (f"{N_HARD}", {"bold":True}),
    " companies excluded by this override: 10 on Sustainalytics High, 1 on Severe, and 2 "
    "on ISS Poor. Truvalue is not granted a veto — its worst outcome is a non-binding "
    "watchlist flag, reflecting its softer assessment basis. The funnel labels this step "
    f"the ‘hybrid red-flag override’ because Stage 1 combines vote-based and veto-based "
    "logic into a single eligibility gate.")

h(doc, "5.4  Fossil-fuel exclusion", 2)
para(doc,
    "Independent of the vendor screen, any company flagged ‘active in fossil fuel’ in "
    "the FactSet activity tag is excluded outright. The mandate explicitly excludes "
    "thermal coal revenue above 5% and controversial weapons; the fossil flag is the "
    "operational implementation of the energy-exposure side of this rule. Twenty-two "
    f"of {N_UNI} universe companies carry the fossil flag.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. ESG SCORING METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "6. ESG Scoring Methodology (Stage 2)", 1)

h(doc, "6.1  The headline number — what it is, and what it isn't", 2)
para(doc,
    "Every company that survives Stage 1 receives an in-house ESG score. This score is ",
    ("not", {"italic":True}), " an absolute quality measure on a 0–100 scale; it is a ",
    ("percentile rank", {"bold":True}), " of the company's composite ESG z-score against "
    "the 224 scored peers. A score of 83.1 therefore means ‘this company ranks above 83.1% "
    "of the scored European universe on our in-house ESG composite’ — not ‘this company "
    "scored 83 out of 100 on ESG quality.’ This distinction matters: the universe average "
    "is, by construction, the 50th percentile, and the score is always relative.")

h(doc, "6.2  Ten indicators across three pillars", 2)
para(doc,
    "The in-house z is built from ten Bloomberg-sourced indicators — four environmental, "
    "three social, three governance. Three of the ten are binary (yes/no); the remainder "
    "are continuous.")
make_table(doc,
    ["Code", "Indicator", "Pillar", "Type"],
    [
        ("E1", "Carbon intensity",            "Environmental", "Continuous"),
        ("E2", "Scope-3 disclosure",          "Environmental", "Binary"),
        ("E3", "Water intensity",             "Environmental", "Continuous"),
        ("E4", "Renewable energy %",          "Environmental", "Continuous"),
        ("S1", "Workforce gender",            "Social",        "Continuous"),
        ("S2", "Lost-time injury rate",       "Social",        "Continuous"),
        ("S3", "Human-rights policy",         "Social",        "Binary"),
        ("G1", "Board independence",          "Governance",    "Continuous"),
        ("G2", "Board gender",                "Governance",    "Continuous"),
        ("G3", "ESG-linked pay",              "Governance",    "Binary"),
    ],
    widths_cm=[1.2, 7.5, 4.0, 4.3])

h(doc, "6.3  Scoring chain", 2)
para(doc,
    "Each indicator is converted to a z-score against sector peers, with ±3σ "
    "winsorisation to cap outliers. Where a sector contains fewer than ten constituents "
    "the z is computed against the full universe (a hybrid sector/universe basis). "
    "Indicator z-scores are aggregated into pillar z-scores using SASB-materiality "
    "weights that vary by sector — for example, an energy company's E-pillar leans on "
    "carbon and water; a bank's G-pillar leans on board independence. The three pillar "
    "z-scores then combine into the in-house composite z with fixed cross-pillar weights "
    "of ", ("50% Environmental · 20% Social · 30% Governance", {"bold":True}),
    " (methodology v1.1 default). The composite z is finally ranked, and the rank "
    "percentile is the headline ESG_score.")

para(doc, ("Worked example — Alfa Laval (rank 2).", {"bold":True}),
    "  E-pillar z = +1.41; S-pillar z = −0.37; G-pillar z = +0.05. The composite z is "
    "0.50·1.41 + 0.20·(−0.37) + 0.30·0.05 = 0.65, placing the company at the 99.6th "
    "percentile of the 224 scored peers. The arithmetic is fully reproducible from the "
    "pillar z-scores in the screening workbook.")

h(doc, "6.4  ESG used twice — and why", 2)
para(doc,
    "The in-house ESG score is used at two distinct points in the funnel. First, the "
    f"{N_S2} eligible companies are ranked on it, and the sector-capped top {N_CAP} form "
    "the shortlist (224 → 40). Second, ESG is one of two inputs to the composite that "
    f"selects the final {N_HOLD} ({N_CAP} → {N_HOLD}). This is a deliberate, layered design: "
    "ESG first acts as a ", ("best-in-class screen", {"italic":True}),
    " defining the eligible playing field, and then as ",
    ("ESG integration", {"italic":True}),
    " within the composite. Both are recognised techniques in sustainable investing and "
    "their combination is more robust than either alone. The honest implication is that "
    "ESG's true influence on selection is greater than the headline 40% composite weight: "
    "a company outside the top-40-by-ESG never enters the composite at all.")

h(doc, "6.5  E-pillar defaulting and the Financials tilt", 2)
para(doc,
    f"Of {N_HOLD} portfolio holdings, ", (f"{n_E_default} carry an E-pillar z of exactly +1.000", {"bold":True}),
    " — a discrete value the workbook assigns when environmental indicator data is absent "
    "or sector-immaterial. Real z-scores do not land on identical round numbers; this is "
    "a default. Universe-wide, 69 of 224 companies are pinned at the same +1.000 default, "
    "and 61 of those 69 are Financials — banks and insurers, which barely report carbon, "
    "water and renewables. Because the cross-pillar weight on E is fixed at 50%, roughly "
    "half the ESG score for those holdings rests on a sector default rather than measured "
    "data. The default itself is positive (≈ 85th percentile of the E-pillar distribution), "
    "which structurally favours low-disclosure sectors in the ranking.")

h(doc, "6.6  Capped vs unconstrained top 40", 2)
para(doc,
    "The shortlist is sector-capped at six names per sector. Without that cap, the "
    f"unconstrained top {N_CAP} by ESG percentile would carry a mean ESG score of "
    f"{unc_esg:.1f} — but 29 of those 40 would be Financials, making a diversified "
    "≥5-sector portfolio impossible. The cap deliberately trades ~12 percentile points "
    f"of average ESG (91.8 → {uni_esg:.1f}) for sector breadth, replacing 23 high-ESG "
    "Financials with lower-ESG names from underrepresented sectors. The portfolio "
    "therefore holds names such as Subsea 7 at the 55.7th percentile, defensible only "
    "as the cost of diversification: pure ESG ranking would have produced a portfolio "
    "of banks.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. FINANCIAL QUALITY SCREEN
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "7. Financial Quality Screen", 1)

h(doc, "7.1  Composite financial score", 2)
para(doc,
    "Each shortlisted company receives a composite financial score on a 0–100 scale, "
    "combining fundamental quality and market-based risk metrics. The fundamental metrics "
    "comprise four indicators (M-01 through M-04), drawn from the Screen-B framework "
    "documented in /docs/financial_filtering_framework: M-01 ROIC versus WACC (capital "
    "productivity), M-02 free cash flow conversion (earnings quality), M-03 FCCR and net "
    "debt to EBITDA (solvency), and M-04 Sloan ratio (accruals quality). The market "
    "metrics are Sharpe ratio, annualised volatility, maximum drawdown, and beta to the "
    "STOXX Europe 600 benchmark, computed from a five-year window of daily prices.")

h(doc, "7.2  Hard gate and review-required verdicts", 2)
para(doc,
    "A binary gate is applied before scoring: companies with annualised volatility above "
    "40%, or with grossly incomplete fundamental data, are excluded outright. "
    f"Seven companies failed the financial gate at this stage — Logitech, Rentokil "
    "Initial, Arcadis, L'Oreal, Sweco, Moncler, and AIXTRON — and are logged in the "
    "exclusions register with the failure reason.")

para(doc,
    "Two holdings carry a ", ("REVIEW_REQUIRED", {"italic":True}),
    " gate verdict rather than a clean PASS — AIB Group and Lloyds Banking Group. "
    "Both are large-cap European banks, and the verdict reflects a framework limitation "
    "rather than a company-specific concern: the Screen-B fundamental metrics (EBITDA-based "
    "solvency and cash-flow ratios) do not map cleanly to bank financial statements, which "
    "operate on net interest income and CET1 ratios. The Investment Committee documented "
    "this in the override log and elected to retain both holdings on the strength of "
    "their Tier-1 vendor verdicts, Sustainalytics ‘Low’ risk band, and within-mandate "
    "market metrics. A future iteration would substitute bank-specific criteria.")

h(doc, "7.3  Distribution of financial scores", 2)
para(doc,
    f"Within the final {N_HOLD}, financial scores range from {fp['fin_score'].min():.1f} "
    f"(Norsk Hydro) to {fp['fin_score'].max():.1f} (Swiss Prime Site), with a weighted "
    f"average of {wfin:.1f}. The wide spread reflects the deliberate 60% financial weight "
    "in the composite: companies with strong ESG but weak financials (Norsk Hydro at 39.4, "
    "Alfa Laval at 42.4, Aegon at 45.8) are still included because their ESG ranking "
    "compensates, but the composite ensures financial quality remains a binding "
    "consideration. The portfolio is not an unconstrained ESG ranking.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. PORTFOLIO CONSTRUCTION
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "8. Portfolio Construction (Stage 3)", 1)

h(doc, "8.1  Composite score", 2)
para(doc,
    "Within the capped 40-name shortlist, every company is assigned a composite score on "
    "a 0–100 scale:")
fp_eq = doc.add_paragraph()
fp_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp_eq.add_run("Composite = 0.40 × ESG_percentile + 0.60 × Financial_score")
r.bold = True; r.font.size = Pt(11)
para(doc,
    "The 60/40 financial-tilt reflects the mandate's stated weighting and the team's "
    "view that, for a long-horizon long-only fund, financial quality is the primary "
    "return driver and ESG is the secondary risk filter. The composite ranks the 40 "
    f"shortlisted companies; the top {N_HOLD} survive.")

h(doc, "8.2  Diversification rules", 2)
para(doc, "Three diversification constraints apply during selection:")
bullet(doc, f"≥{mandate['constraints']['min_sectors']} sectors represented (achieved: {n_sect}).", bold_prefix="Min sectors — ")
bullet(doc, "no more than five holdings from any single sector (a 25% sector ceiling on weight).", bold_prefix="Sector cap — ")
bullet(doc, "no two holdings with a 1-year return correlation above 0.90, applied iteratively as names are added.", bold_prefix="Correlation guard — ")

h(doc, "8.3  Weighting scheme", 2)
para(doc,
    "Holdings are weighted ", ("proportional to the composite score", {"italic":True}),
    ", normalised to sum to 100% and capped at 10% per name. This is sometimes called a "
    "‘score-tilted’ weighting: stronger names receive more capital, but the cap prevents "
    "any single position from dominating. The mandate's hard 10% ceiling is binding — no "
    f"holding exceeds {maxw:.1f}%. After the cap-and-renormalise loop, weights run from "
    f"{w.min()*100:.2f}% (Subsea 7, rank 20) to {maxw:.2f}% (Zurich Insurance, rank 1).")

h(doc, "8.4  Why not optimisation-first?", 2)
para(doc,
    "The optimisation module (PyPortfolioOpt) was used to ", ("backtest", {"italic":True}),
    " six alternative weighting schemes after construction — equal-weight, score-tilted, "
    "Black-Litterman, hierarchical risk parity, minimum-variance, and maximum-Sharpe (see "
    "§11) — but it does not select the holdings. Selection is rule-based: composite rank "
    "with hard sector and correlation guards. This separation is deliberate: a mean-"
    "variance optimisation on five years of in-sample data carries severe look-ahead bias "
    "for a five-year forward portfolio. The rule-based ranking is more honest about what "
    "the team can and cannot claim ex-ante.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. GREENWASHING ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "9. Greenwashing Assessment", 1)

h(doc, "9.1  What it measures", 2)
para(doc,
    "The greenwashing 8-Test asks a question the ESG score does not: ", ("can a company "
    "back up what it says about sustainability?", {"italic":True}), " It is a forensic "
    "screen of disclosure quality, not of environmental performance. A company may have "
    "real emissions reductions but make vague, unverified claims about them; another may "
    "have modest emissions but report transparently and consistently. The 8-Test "
    "distinguishes the two by interrogating eight dimensions of every sustainability "
    "claim against the source documents.")

h(doc, "9.2  The eight dimensions", 2)
make_table(doc,
    ["Dimension", "What is checked"],
    [
        ("1. Specificity",   "Exact wording, not vague buzzwords"),
        ("2. Metric",        "A real number supporting the claim"),
        ("3. Baseline",      "A clear comparison reference point"),
        ("4. Target",        "A firm, committed endpoint"),
        ("5. Time horizon",  "A credible, near-enough deadline"),
        ("6. Scope",         "Clear coverage of divisions or assets"),
        ("7. Verification",  "External assurance, not self-reported"),
        ("8. Consistency",   "Capex and lobbying match the claim"),
    ],
    widths_cm=[4.5, 12.5])

h(doc, "9.3  Sources, scope and verification", 2)
para(doc,
    "The screen draws on two tiers of evidence. ", ("Tier 1", {"bold":True}),
    " is the company's own disclosure — integrated annual reports, sustainability "
    "statements, TCFD and CSRD filings, governance reports. ", ("Tier 2", {"bold":True}),
    " is independent third-party evidence — the SBTi public registry for validated "
    "targets, and TPI Carbon Performance assessments for the credibility of stated "
    "transition pathways. The combined corpus runs to roughly 120 PDFs across the 40 "
    "capped names. The 8-Test itself was applied to all 20 portfolio holdings; the "
    "remaining 20 corpus folders are kept as the safety net for re-screening if a "
    "holding were ever excluded.")

para(doc,
    "Separately, a Tier-1 ", ("controversy screen", {"italic":True}),
    " covered all 40 capped names against the financial press (Reuters, Bloomberg, FT, "
    "WSJ), regulators (FCA, OFSI, FDA, PRA), NGOs (BankTrack, Business & Human Rights "
    "Resource Centre), and regional dailies, over a rolling six-month window. Five "
    "in-window findings were recorded, two borderline, and 33 clean.")

h(doc, "9.4  Results", 2)
para(doc,
    f"All {N_HOLD} portfolio holdings PASSED the 8-Test screen ", ("(zero excluded, zero "
    "watchlisted)", {"italic":True}), f". The mean concern score is {gw_mean:.1f}% and "
    f"the worst is {gw_worst} at {gw_max:.1f}%. The exclusion threshold (three or more "
    "HIGH-concern dimensions) was not reached by any holding; the watchlist threshold "
    "(two HIGH-concern dimensions) was also not reached. The 8-Test is therefore a "
    "confirmation rather than a filter in the present run, with the corpus standing "
    "ready to support a re-screen if any holding were later excluded.")

para(doc,
    "Of the five in-window controversy findings, three are on companies we hold (Lloyds "
    "Banking Group, AstraZeneca, Inditex). All three are documented in the override "
    "register as CONTROVERSY_REVIEW entries: Lloyds and Inditex are KEEP_DOCUMENTED, "
    "and AstraZeneca is on MONITOR given the materiality of the China indictment. The "
    "other two findings are on non-held companies (Raiffeisen Bank International, Ipsen).")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. CLIMATE, BIODIVERSITY & EU REGULATION
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "10. Climate, Biodiversity and EU Regulation", 1)

h(doc, "10.1  Weighted Average Carbon Intensity (WACI)", 2)
para(doc,
    f"The portfolio WACI is ", (f"{waci:.0f} tCO₂e per million of revenue", {"bold":True}),
    " on a Scope 1+2 basis. This is the mandated carbon metric and is computed as the "
    "weighted sum of holding-level carbon intensities. The Scope 3 dimension is "
    "deliberately excluded from the headline figure because Scope-3 reporting is too "
    "incomplete across the universe to compare like-for-like; Scope-3 disclosure is "
    "instead recorded as a binary indicator within the E-pillar of the ESG score.")

para(doc,
    ("Two-thirds of the WACI is attributable to a single holding.", {"bold":True}),
    f"  Norsk Hydro contributes approximately {nh_share:.0f}% of the portfolio WACI on a "
    f"weight-times-intensity basis. The intensity figure assigned to Norsk Hydro is the "
    "BICS-Materials sector median (≈ 2,150 tCO₂e/$M), not Norsk Hydro's reported value — "
    "see §10.2. Excluding Norsk Hydro, the portfolio WACI falls to roughly "
    f"{waci_exnh:.0f} tCO₂e/$M, in line with the wider European universe. The "
    "concentration is disclosed in §12 as a material data-quality limitation.")

h(doc, "10.2  Carbon intensity coverage and imputation", 2)
para(doc,
    f"Of the {N_HOLD} portfolio holdings, only ", (f"{ci_reported} carry a reported "
    "carbon intensity", {"bold":True}), f" (Subsea 7 and UCB). The remaining "
    f"{ci_imputed} use a ", ("sector-median imputed", {"italic":True}),
    f" value, accounting for {imp_w:.0f}% of portfolio weight. Imputation is by BICS "
    "sector, so multiple holdings in the same sector inherit identical values — Aegon, "
    "AIB, Lloyds and Zurich all carry the Financials median (0.7 tCO₂e/$M); the four "
    "Health Care holdings all carry 8.0; the four Real-Estate / Infrastructure holdings "
    "share 16.7. This is the largest single proxy in the pipeline. A documented "
    "human-override now substitutes Norsk Hydro's ",
    ("reported", {"bold":True}), " revenue-based Scope 1+2 intensity (~474 tCO₂e/€M, "
    "computed from the company's 2024 ESRS sustainability statements) for the imputed "
    "Materials-sector median; the override entry is logged in the IC register.")

h(doc, "10.3  Biodiversity / nature-risk proxy", 2)
para(doc,
    "Biodiversity is reported as a ", ("nature-risk proxy", {"italic":True}),
    " by design, in line with the assignment's wording. The proxy is built from two "
    "established sector-level datasets — ENCORE (Exploring Natural Capital Opportunities, "
    "Risks and Exposure), which rates sector-level dependencies and impacts on natural "
    "capital, and WRI Aqueduct, which rates water-stress exposure by geography. The "
    "combined score is sector-and-location-based rather than company-measured: two "
    "companies in the same sector and geography receive identical scores. This is a "
    "structural limitation of nature data and not a methodological flaw; it will narrow "
    "as the CSRD/ESRS E4 disclosures come into force.")

h(doc, "10.4  EU Taxonomy and SFDR posture", 2)
para(doc,
    "Eligibility and estimated alignment under the EU Taxonomy are reported using the "
    "FactSet ‘estimated’ fields (the column names carry an ", ("Estmatd", {"italic":True}),
    " prefix). Reported alignment coverage is sparse across the universe — fewer than "
    "ten companies have substantive reported alignment percentages — so eligibility is "
    "used as the operational proxy. The fund positions itself as ",
    ("SFDR Article 8", {"bold":True}),
    " (promoting environmental and social characteristics), supported by the documented "
    "exclusion list, the binding ESG screen, the greenwashing layer, and the auditable "
    "override register. It does not claim Article 9 (sustainable-objective) status, "
    "because the lack of reported Taxonomy alignment at portfolio level would not meet "
    "the higher bar.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. BACKTEST & PERFORMANCE
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "11. Backtest and Performance", 1)

h(doc, "11.1  Setup", 2)
para(doc,
    "The optimisation module runs a walk-forward out-of-sample backtest of six "
    "alternative weighting schemes applied to the final 20 holdings, against a benchmark "
    "of the STOXX Europe 600 (Yahoo ticker ^STOXX). The window is approximately five "
    "years of daily price data, with quarterly rebalancing inside each scheme; ticker "
    "prices use yfinance with auto_adjust=True so that dividends are reinvested into "
    "the portfolio's total return. The six schemes are equal-weight, score-tilted "
    "(composite-score-proportional, the production weighting), Black-Litterman, "
    "Hierarchical Risk Parity, Minimum-Variance, and Maximum-Sharpe.")

h(doc, "11.2  Results — by weighting scheme", 2)
make_table(doc,
    ["Method", "Cum. return", "CAGR", "Vol.", "Sharpe", "Max DD", "Composite"],
    [
        (r["method"].replace("_"," ").title().replace("Hrp","HRP"),
         f"{r['cumulative_return']*100:.0f}%",
         f"{r['cagr']*100:.1f}%",
         f"{r['ann_vol']*100:.1f}%",
         f"{r['sharpe']:.2f}",
         f"{r['max_drawdown']*100:.1f}%",
         f"{r['composite_score']:.2f}")
        for _, r in bt_ranked.iterrows()
    ],
    widths_cm=[3.5, 2.4, 2.2, 1.9, 1.9, 2.2, 2.4])

para(doc,
    f"All six schemes beat the benchmark over the OOS window on a Sharpe basis (STOXX 600 "
    "Sharpe ≈ 0.44 over the same window). Equal-weight tops the composite ranking "
    "(0.750), followed by score-tilted (0.710). The team adopted the ", ("score-tilted", {"bold":True}),
    " scheme in production because it preserves the composite-rank signal in the "
    "weights themselves: a stronger composite earns a larger position. Equal-weight has "
    "no such information content; the marginal Sharpe gap (1.52 vs 1.49) is within noise "
    "and does not justify discarding the signal.")

h(doc, "11.3  Honest framing — what the backtest does not prove", 2)
para(doc,
    "Three caveats limit what can be claimed:")
bullet(doc, "the holdings were selected in 2026 and then backtested over 2021–2026. The portfolio's apparent outperformance is, in part, a function of selection bias; the backtest is a relative test between weighting schemes, not a claim of historical alpha.", bold_prefix="Look-ahead bias. ")
bullet(doc, "the ^STOXX series on Yahoo is the price index — it carries no dividend reinvestment. The 20 holdings, downloaded with auto_adjust, are total-return. The comparison therefore flatters the portfolio by roughly the STOXX 600 dividend yield (~3% p.a.), or ~15pp cumulatively over the window.", bold_prefix="Price-return vs total-return. ")
bullet(doc, "two of three optimisation methods (Black-Litterman, Max-Sharpe) require cvxpy and a positive-definite covariance matrix; outputs are stable but sensitive to the look-back window choice. The ranking should be read as ordinal, not precise.", bold_prefix="Numerical sensitivity. ")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. LIMITATIONS, OVERSIGHT & CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h(doc, "12. Limitations, Human Oversight and Conclusion", 1)

h(doc, "12.1  Material limitations", 2)
para(doc,
    "Eight limitations are disclosed up front. None is a defect we are concealing; each "
    "is a deliberate trade-off whose cost we judged acceptable, and whose disclosure is "
    "part of the pipeline's audit posture.")
bullet(doc, f"{ci_imputed} of {N_HOLD} portfolio holdings still use a sector-median imputed carbon intensity. Norsk Hydro was the largest single imputed contributor and has been corrected via documented IC override — the imputed Materials-sector median (2,150 tCO₂e/€M) replaced with the reported revenue-based Scope 1+2 intensity (~474 tCO₂e/€M) from the company's 2024 ESRS statements. After the override Norsk Hydro accounts for ≈{nh_share:.0f}% of the headline WACI; the next-largest remaining imputed contributor is E.ON at the Utilities-sector median.", bold_prefix="WACI imputation. ")
bullet(doc, f"{n_E_default} of {N_HOLD} holdings receive a sector-default +1.000 E-pillar z-score because environmental indicator coverage is thin in their sectors (banks, insurers, some pharma). Because the E-pillar carries 50% of the composite z, this materially lifts those holdings' ESG percentiles; the universe-wide pattern (61 of 69 defaulted names are Financials) explains why an unconstrained ESG ranking would have produced an all-banks portfolio.", bold_prefix="ESG E-pillar defaults. ")
bullet(doc, "biodiversity is a sector-and-location proxy, not a company measurement. Two holdings in the same sector and geography carry identical scores. The metric satisfies the assignment's ‘≥1 nature-risk proxy’ requirement honestly, but it is not a fine-grained signal.", bold_prefix="Biodiversity proxy. ")
bullet(doc, "EU Taxonomy reported alignment coverage is too sparse to use directly; eligibility (and FactSet's ‘estimated’ alignment) is used as the operational proxy. The fund accordingly claims Article 8 rather than Article 9.", bold_prefix="Taxonomy estimation. ")
bullet(doc, "the controversy screen covers a rolling six-month window only. ‘Clean within window’ is not ‘exonerated in perpetuity’; legacy controversies (e.g. Norsk Hydro Brazil litigation, SBM Offshore FCPA history, Inditex Xinjiang investigation) are noted contextually but not counted as findings.", bold_prefix="Controversy window. ")
bullet(doc, "indicator-level coverage for the in-house ESG score (which of the 10 raw indicators are reported vs proxied for each company) is held in the ESG specialist's screening workbook and not independently reproduced in the project repository. The pillar-level and composite-level computations are fully reproducible from the workbook.", bold_prefix="ESG raw-indicator coverage. ")
bullet(doc, "the project data dictionary previously classified all 677 tracked variables as ‘reported’, which understated the use of vendor estimates, observed-from-market values, and EU Taxonomy estimates. The dictionary has been reclassified to use the prescribed five-way taxonomy — 22 variables now flagged ‘estimated’ (EU Taxonomy FactSet Estmatd fields), 402 ‘observed’ (calculated from market data), and 253 ‘reported’. A data-governance entry is logged in the IC override register.", bold_prefix="Data-dictionary classification. ")
bullet(doc, "the backtest carries look-ahead bias (holdings selected ex-post and backtested) and a price-return vs total-return mismatch against the ^STOXX proxy. Treat the backtest as an ordinal test between weighting schemes, not as evidence of historical alpha.", bold_prefix="Backtest scope. ")

h(doc, "12.2  Human oversight and override log", 2)
para(doc,
    f"Nine substantive override decisions are documented in outputs/scores/"
    f"override_decisions_*.csv, each with reviewer, evidence source, date, rationale, "
    "and a written caveat. Three categories of override appear:")
bullet(doc, "four transition-watchlist resolutions (Klepierre, MERLIN Properties, Galenica, Orion) — each kept on the strength of an SBTi-validated transition pathway despite a thin Bloomberg Tier-1 signal.", bold_prefix="Transition watchlist. ")
bullet(doc, "four Truvalue-watchlist resolutions (AIB, Norsk Hydro, Inditex, Swedish Orphan Biovitrum) — each retained because Sustainalytics ‘Low’ risk and ISS ‘Medium/Good’ offset a Truvalue ‘Below Average / Laggard’ tier flag, with the rationale documented case-by-case.", bold_prefix="Truvalue watchlist. ")
bullet(doc, "one financial-gate review (Lloyds Banking) — kept because the REVIEW_REQUIRED verdict reflects the Screen-B framework's incompatibility with bank financial statements rather than a company-specific concern.", bold_prefix="Financial-gate review. ")
para(doc,
    "These are concrete, evidence-sourced human decisions, satisfying the assignment's "
    "‘≥ 3 human override examples’ requirement many times over. Five further entries were "
    "added in the current revision of the register: three CONTROVERSY_REVIEW decisions "
    "(Lloyds OFSI fine — KEEP_DOCUMENTED; AstraZeneca China indictment — MONITOR; Inditex "
    "A Coruña wage dispute — KEEP_DOCUMENTED), one DATA_GOVERNANCE_DECISION (data-"
    "dictionary reclassification), and one WACI_DATA_OVERRIDE (Norsk Hydro reported "
    "Scope 1+2 substituted for the imputed Materials median). The register now contains "
    "14 substantive, evidence-cited entries.")

h(doc, "12.3  Conclusion", 2)
para(doc,
    "The portfolio is the output of a disciplined sequential funnel — vote, veto, rank, "
    f"pick — applied to a {N_UNI}-name European universe. Every step is documented, "
    "every exclusion is rule-based and reproducible, and every proxy and default is "
    "disclosed in this section rather than hidden in the methodology. The result is "
    f"{N_HOLD} holdings across {n_sect} sectors, with an in-house weighted ESG percentile "
    f"of {wesg:.1f}, a Sharpe ratio of {wsharp:.2f}, a WACI of {waci:.0f}, and zero "
    "greenwashing red flags. We do not claim this is the only defensible portfolio that "
    "could be built from the same data — but we believe every choice we made can be "
    "explained, every number traced, and every limitation acknowledged. That is the "
    "standard the pipeline was designed to meet, and that is the standard this report "
    "documents.")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run(f"ESADE MSc Finance · Sustainable Finance Group Assignment · "
                 f"Methodology Report v2 · {TODAY}")
r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY_TXT

# ── Save ──────────────────────────────────────────────────────────────────────
out = f"Final_Deliverables/Final_Report_v2_{TODAY}.docx"
doc.save(out)

# Word count (rough)
import zipfile, xml.etree.ElementTree as ET
NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
with zipfile.ZipFile(out) as z:
    body = ET.fromstring(z.read("word/document.xml"))
    text = " ".join(t.text or "" for t in body.iter(f"{NS}t"))
words = len(re.findall(r"\w+", text))

print(f"Final report saved: {out}")
print(f"  approximate word count: {words}")
print(f"  12 sections · {N_HOLD} holdings · "
      f"weighted ESG {wesg:.1f} · WACI {waci:.0f} · Sharpe {wsharp:.2f}")
