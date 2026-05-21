#!/usr/bin/env python
"""Export the final 20-stock portfolio as a formatted Word document.

    python scripts/export_portfolio_docx.py
    -> outputs/portfolio/final_portfolio_<date>.docx
"""
import glob, json, os
import pandas as pd
from datetime import date

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
os.chdir(PROJECT)
TODAY = str(date.today())

fp = pd.read_csv(sorted(glob.glob("outputs/portfolio/final_portfolio_*.csv"))[-1])
mandate = json.load(open("outputs/scores/mandate.json")) if os.path.exists(
    "outputs/scores/mandate.json") else {}
fund  = mandate.get("fund_name", "ESADE Sustainable European Equity Fund")
bench = mandate.get("benchmark", "STOXX Europe 600")

fp = fp.sort_values("weight", ascending=False).reset_index(drop=True)
w = fp["weight"]

# ── Headline metrics ───────────────────────────────────────────────────────────
wesg   = (fp["ESG_score"]    * w).sum()
wfin   = (fp["fin_score"]    * w).sum()
wsharp = (fp["sharpe_ratio"] * w).sum()
waci   = (fp["carbon_intensity"].fillna(0) * w).sum()
sec_w  = fp.groupby("sasb_sector")["weight"].sum()
n_ic   = int(fp["override_disposition"].notna().sum())
gw_pass = int((fp["gw_high_count"] < 2).sum())

# ════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell(cell, text, bold=False, align=None, size=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if size:
        run.font.size = Pt(size)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(9.5)

# Landscape — a holdings sheet is naturally wide
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(sec, m, Inches(0.7))

doc.add_heading("Final Portfolio", 0)
p = doc.add_paragraph()
p.add_run(f"{fund}   ").bold = True
p.add_run(f"·  20 long-only European equity holdings  ·  benchmark {bench}  "
          f"·  data vintage {TODAY}").italic = True

# ── Key metrics table ──────────────────────────────────────────────────────────
doc.add_heading("Key metrics", 2)
metrics = [
    ("Number of holdings", "20"),
    ("Sectors represented", f"{fp['sasb_sector'].nunique()}  (mandate minimum 5)"),
    ("Largest sector weight", f"{sec_w.max()*100:.1f}%  (ceiling 25%)"),
    ("Largest single holding", f"{w.max()*100:.1f}%  (ceiling 10%)"),
    ("Weighted ESG score", f"{wesg:.1f} / 100"),
    ("Weighted financial score", f"{wfin:.1f} / 100"),
    ("Weighted Sharpe ratio (5y)", f"{wsharp:.2f}"),
    ("WACI (carbon intensity)", f"{waci:.0f} tCO2e/$M revenue"),
    ("Greenwashing 8-Test", f"{gw_pass}/20 PASS  (0 excluded, 0 watchlisted)"),
    ("Holdings with an IC override", f"{n_ic} of 20"),
]
mt = doc.add_table(rows=0, cols=4)
mt.style = "Light Grid Accent 1"
for i in range(0, len(metrics), 2):
    cells = mt.add_row().cells
    set_cell(cells[0], metrics[i][0], bold=True)
    set_cell(cells[1], metrics[i][1])
    if i + 1 < len(metrics):
        set_cell(cells[2], metrics[i+1][0], bold=True)
        set_cell(cells[3], metrics[i+1][1])

# ── Holdings table ─────────────────────────────────────────────────────────────
doc.add_heading("Holdings", 2)
cols = ["#", "Company", "Ticker", "Sector", "Weight", "ESG", "Financial",
        "Composite", "8-Test", "IC review"]
t = doc.add_table(rows=1, cols=len(cols))
t.style = "Light Grid Accent 1"
for c, h in enumerate(cols):
    set_cell(t.rows[0].cells[c], h, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(t.rows[0].cells[c], "1F4E78")
    t.rows[0].cells[c].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

R = WD_ALIGN_PARAGRAPH.RIGHT
for i, (_, r) in enumerate(fp.iterrows(), 1):
    cells = t.add_row().cells
    ic = r["override_disposition"] if pd.notna(r["override_disposition"]) else "—"
    vals = [(str(i), None), (r["company_name"], None), (r["yf_ticker"], None),
            (r["sasb_sector"], None), (f"{r['weight']*100:.2f}%", R),
            (f"{r['ESG_score']:.1f}", R), (f"{r['fin_score']:.1f}", R),
            (f"{r['composite_score']:.1f}", R), (f"{r['gw_score_pct']:.1f}%", R),
            (ic, None)]
    for c, (val, al) in enumerate(vals):
        set_cell(cells[c], val, align=al)
# Total row
tot = t.add_row().cells
set_cell(tot[1], "TOTAL", bold=True)
set_cell(tot[4], f"{w.sum()*100:.1f}%", bold=True, align=R)
for c in tot:
    shade(c, "EAF0F6")

widths = [0.4, 2.5, 0.8, 2.0, 0.8, 0.6, 0.85, 0.95, 0.7, 1.4]
for row in t.rows:
    for c, wd in enumerate(widths):
        row.cells[c].width = Inches(wd)

doc.add_paragraph()
note = doc.add_paragraph()
note.add_run("Method.  ").bold = True
note.add_run(
    "Holdings are selected from a sector-capped 40-company ESG shortlist by a "
    "composite score (60% financial quality + 40% ESG), with a hard "
    "max-5-holdings-per-sector cap and a 0.90 return-correlation guard. Weights "
    "are proportional to the composite score, capped at 10% per name. ESG, "
    "Financial and Composite are 0–100 scores; 8-Test is the greenwashing "
    "concern score (0% = clean, 100% = severe).").italic = True

# ── Sector allocation ──────────────────────────────────────────────────────────
doc.add_heading("Sector allocation", 2)
g = (fp.groupby("sasb_sector")["weight"].agg(["sum", "count"])
       .sort_values("sum", ascending=False))
st = doc.add_table(rows=1, cols=3)
st.style = "Light Grid Accent 1"
for c, h in enumerate(["Sector", "Holdings", "Weight"]):
    set_cell(st.rows[0].cells[c], h, bold=True)
for sname, r in g.iterrows():
    cells = st.add_row().cells
    set_cell(cells[0], sname)
    set_cell(cells[1], int(r["count"]), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(cells[2], f"{r['sum']*100:.1f}%", align=R)
for c, wd in enumerate([3.0, 1.0, 1.0]):
    for row in st.rows:
        row.cells[c].width = Inches(wd)

doc.add_paragraph()
disc = doc.add_paragraph()
disc.add_run(
    "Academic prototype — not a regulated investment product or investment "
    "advice. ESADE MSc Finance, final group assignment.").italic = True
disc.runs[0].font.size = Pt(8.5)
disc.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

out = f"outputs/portfolio/final_portfolio_{TODAY}.docx"
doc.save(out)
print(f"Word document saved: {out}")
print(f"  20 holdings | weighted ESG {wesg:.1f} | Sharpe {wsharp:.2f} | "
      f"WACI {waci:.0f} | {fp['sasb_sector'].nunique()} sectors")
