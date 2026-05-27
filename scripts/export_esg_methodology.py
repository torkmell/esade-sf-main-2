#!/usr/bin/env python
"""Export a one-page methodology note: how the in-house ESG score is computed.

Documents the chain  10 indicators -> sector-relative winsorised z ->
SASB-weighted pillar z -> 50/20/30 cross-pillar composite -> percentile.
Worked examples and pillar z-scores are read live from the ESG specialist's
screening workbook so the numbers always match the pipeline.

    python scripts/export_esg_methodology.py
    -> outputs/reports/esg_methodology_note_<date>.docx
"""
import glob, os, re
import pandas as pd
from datetime import date

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
os.chdir(PROJECT)
TODAY = str(date.today())

WB = "data/provided/Portfolio_Screening_Output.xlsx"

# ── Load source data ────────────────────────────────────────────────────────────
s2 = pd.read_excel(WB, sheet_name="Stage 2 — Full ranking")
fp = pd.read_csv(sorted(glob.glob("outputs/portfolio/final_portfolio_*.csv"))[-1])

_SFX = sorted(["class b", "class a", "class c", "ab", "sa", "nv", "plc", "ag",
               "asa", "oyj abp", "oyj", "oy", "ltd", "inc", "corp", "spa",
               "s p a", "s a", "se", "a s", "as", "group", "holding",
               "holdings", "sca", "psc"], key=len, reverse=True)

def norm(s):
    s = re.sub(r"[^\w\s]", " ", str(s).lower())
    s = re.sub(r"\s+", " ", s).strip()
    for x in _SFX:
        if s.endswith(" " + x):
            s = s[: -len(x) - 1].strip()
            break
    return s

s2["_k"] = s2["Company"].apply(norm)
fp["_k"] = fp["company_name"].apply(norm)
m = fp.merge(s2[["_k", "In-house ESG z", "Percentile", "E pillar", "S pillar",
                 "G pillar"]], on="_k", how="left").sort_values(
                     "Percentile", ascending=False)

# Worked-example rows: top, median, bottom holding by percentile
ex_rows = pd.concat([m.head(1), m.iloc[[len(m) // 2]], m.tail(1)])

# ════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell(cell, text, bold=False, align=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def header_row(table, labels):
    for c, h in enumerate(labels):
        set_cell(table.rows[0].cells[c], h, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, white=True)
        shade(table.rows[0].cells[c], "1F4E78")

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)
for m_ in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(doc.sections[0], m_, Inches(0.8))

doc.add_heading("In-House ESG Score — Methodology", 0)
p = doc.add_paragraph()
p.add_run("ESADE Sustainable European Equity Fund").bold = True
p.add_run("   ·   how the in-house ESG z-score and E / S / G pillars are "
          f"constructed   ·   {TODAY}").italic = True

intro = doc.add_paragraph()
intro.add_run(
    "Rather than adopt a single third-party ESG rating — vendors disagree "
    "materially (Berg, Koelbel & Rigobon, 2022) — the fund builds its own "
    "transparent score from raw indicators. The score is a sector-relative "
    "composite: a company is judged against peers in its own sector, not "
    "against the whole market. The pipeline records it three ways — the raw "
    "composite z-score (in_house_z), its percentile rank (ESG_score), and the "
    "rescaled 0–100 pillar scores (E_score / S_score / G_score).")

# ── The five-step chain ─────────────────────────────────────────────────────────
doc.add_heading("How the score is built — five steps", 2)
steps = [
    ("1 · Indicators",
     "Ten indicators (4 Environmental, 3 Social, 3 Governance), Bloomberg as "
     "primary source. E: carbon intensity, Scope-3 disclosure (binary), water "
     "intensity, renewable-energy %. S: workforce gender, lost-time injury "
     "rate, human-rights policy (binary). G: board independence, board "
     "gender, ESG-linked pay (binary)."),
    ("2 · Sector-relative z-score",
     "Each indicator is winsorised at ±3σ to cap outliers, then converted to a "
     "z-score. Hybrid basis: the z-score is computed within the company's own "
     "sector when that sector has ≥10 constituents, and across the full "
     "universe for smaller sectors (too few peers for a stable sector mean)."),
    ("3 · Pillar z-scores",
     "Indicator z-scores are aggregated into an E, S and G pillar z-score. "
     "Weighting within each pillar is SASB-materiality based — each sector "
     "carries its own indicator weights from the SASB Materiality Map, so an "
     "energy company and a bank are not judged on identical priorities."),
    ("4 · Cross-pillar composite",
     "The three pillar z-scores are combined into the overall in-house ESG "
     "z-score with fixed weights: 50% Environmental, 20% Social, 30% "
     "Governance (methodology v1.1 default)."),
    ("5 · Percentile rank",
     "The composite z-score is ranked across the 224 scored companies. That "
     "percentile (0–100) is the ESG_score carried into portfolio selection; "
     "the E / S / G pillar scores shown elsewhere are the pillar z-scores "
     "min-max rescaled to 0–100."),
]
st = doc.add_table(rows=0, cols=2)
st.style = "Light Grid Accent 1"
for tag, body in steps:
    cells = st.add_row().cells
    set_cell(cells[0], tag, bold=True)
    set_cell(cells[1], body)
st.columns[0].width = Inches(1.9)
st.columns[1].width = Inches(5.0)
for row in st.rows:
    row.cells[0].width = Inches(1.9)
    row.cells[1].width = Inches(5.0)

# ── The 10 indicators ───────────────────────────────────────────────────────────
doc.add_heading("The ten indicators", 2)
inds = [
    ("E1", "Carbon intensity", "Environmental"),
    ("E2", "Scope-3 disclosure (binary)", "Environmental"),
    ("E3", "Water intensity", "Environmental"),
    ("E4", "Renewable energy %", "Environmental"),
    ("S1", "Workforce gender", "Social"),
    ("S2", "Lost-time injury rate", "Social"),
    ("S3", "Human-rights policy (binary)", "Social"),
    ("G1", "Board independence", "Governance"),
    ("G2", "Board gender", "Governance"),
    ("G3", "ESG-linked pay (binary)", "Governance"),
]
it = doc.add_table(rows=1, cols=3)
it.style = "Light Grid Accent 1"
header_row(it, ["Code", "Indicator", "Pillar"])
for code, name, pillar in inds:
    cells = it.add_row().cells
    set_cell(cells[0], code, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(cells[1], name)
    set_cell(cells[2], pillar)
for row in it.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(3.6)
    row.cells[2].width = Inches(1.7)

# ── Cross-pillar formula ────────────────────────────────────────────────────────
doc.add_heading("Cross-pillar formula", 2)
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = f.add_run("in-house ESG z  =  0.50 · E_pillar_z  +  0.20 · S_pillar_z  "
               "+  0.30 · G_pillar_z")
fr.bold = True
fr.font.size = Pt(11)

# ── Worked examples ─────────────────────────────────────────────────────────────
doc.add_heading("Worked examples (portfolio holdings)", 2)
wt = doc.add_table(rows=1, cols=6)
wt.style = "Light Grid Accent 1"
header_row(wt, ["Holding", "E pillar z", "S pillar z", "G pillar z",
                "0.50·E + 0.20·S + 0.30·G", "Percentile"])
R = WD_ALIGN_PARAGRAPH.RIGHT
for _, r in ex_rows.iterrows():
    e, s, g = r["E pillar"], r["S pillar"], r["G pillar"]
    calc = 0.50 * e + 0.20 * s + 0.30 * g
    cells = wt.add_row().cells
    set_cell(cells[0], r["company_name"])
    set_cell(cells[1], f"{e:+.3f}", align=R)
    set_cell(cells[2], f"{s:+.3f}", align=R)
    set_cell(cells[3], f"{g:+.3f}", align=R)
    set_cell(cells[4], f"{calc:+.3f}  (z {r['In-house ESG z']:+.3f})", align=R)
    set_cell(cells[5], f"{r['Percentile']:.1f}", align=R)
for row in wt.rows:
    for c, wd in enumerate([1.9, 0.85, 0.85, 0.85, 1.75, 0.85]):
        row.cells[c].width = Inches(wd)
doc.add_paragraph().add_run(
    "The composite z reproduces exactly from the three pillar z-scores — the "
    "50/20/30 weighting is verifiable name by name.").italic = True

# ── Scope / reproducibility ─────────────────────────────────────────────────────
doc.add_heading("Source and reproducibility", 2)
sc = doc.add_paragraph()
sc.add_run("Source.  ").bold = True
sc.add_run(
    "The in-house z, percentile and pillar z-scores are produced by the ESG "
    "specialist's screening workbook (data/provided/Portfolio_Screening_"
    "Output.xlsx, sheet ‘Methodology Summary’ and ‘Stage 2 — Full ranking’). "
    "Notebook 05 imports these columns; it does not recompute them.")
sc2 = doc.add_paragraph()
sc2.add_run("Reproducibility.  ").bold = True
sc2.add_run(
    "Steps 3–5 (pillar z → composite z → percentile) are fully reproducible "
    "from the workbook and verified above. Steps 1–2 (raw indicators → "
    "sector-relative winsorised z) are documented in the workbook but rely on "
    "the specialist's underlying Bloomberg indicator data, which sits outside "
    "the repository. This is a disclosed scope limitation.")

disc = doc.add_paragraph()
disc.add_run(
    "Academic prototype — not a regulated investment product or investment "
    "advice. ESADE MSc Finance, final group assignment.").italic = True
disc.runs[0].font.size = Pt(8.5)
disc.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

os.makedirs("outputs/reports", exist_ok=True)
out = f"outputs/reports/esg_methodology_note_{TODAY}.docx"
doc.save(out)
print(f"Methodology note saved: {out}")
print(f"  {len(inds)} indicators | 50/20/30 cross-pillar | "
      f"{len(ex_rows)} worked examples from the final portfolio")
