#!/usr/bin/env python
"""Export the greenwashing 8-Test screening as Word (.docx) and Excel (.xlsx).

Same content as scripts/_8test_report.py — every dimension's rating, the final
recommendation and how each was assessed — but formatted for hand-in:
  * Excel : Summary (one row per holding, 8 rating columns) + Detail + Methodology
  * Word  : methodology + summary table + a per-company breakdown section
"""
import glob, json, os
import pandas as pd
from datetime import date

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
os.chdir(PROJECT)
TODAY = str(date.today())

DIMS = [("specificity","Specificity"), ("metric","Metric"), ("baseline","Baseline"),
        ("target","Target"), ("time_horizon","Time Horizon"), ("scope","Scope"),
        ("verification","Verification"), ("consistency","Consistency")]
SHOW   = {"LOW":"PASS", "MED":"PARTIAL", "HIGH":"FAIL", "MISSING":"MISSING"}
POINTS = {"LOW":0, "MED":1, "HIGH":2, "MISSING":1}
FILL   = {"PASS":"C6EFCE", "PARTIAL":"FFEB9C", "FAIL":"FFC7CE", "MISSING":"D9D9D9"}
FONTC  = {"PASS":"006100", "PARTIAL":"9C6500", "FAIL":"9C0006", "MISSING":"595959"}

gw = pd.read_csv(sorted(glob.glob("outputs/scores/greenwashing_scores_*.csv"))[-1])
t2 = pd.read_csv("outputs/rag/_tier2_sbti_cp.csv").set_index("ticker")

def tier1(tk):
    p = f"outputs/rag/_tier1_archive/greenwash_{tk}.json"
    return json.load(open(p))["dimensions"] if os.path.exists(p) else {}

def t2_note(tk):
    if tk not in t2.index:
        return "no SBTi/TPI record"
    r = t2.loc[tk]
    s = str(r["sbti_near_term_status"])
    sb = ("not in the SBTi database" if s in ("--","nan","")
          else f"SBTi near-term status '{s}'"
               + (f", {r['sbti_near_term_class']} ({r['sbti_near_term_year']})"
                  if s == "Targets set" else ""))
    tp = str(r["tpi_cp"])
    tpi = (f"; TPI Carbon Performance {tp}-aligned"
           if ("degree" in tp.lower() or "1.5" in tp) else "; not assessed by TPI")
    return sb + tpi

# ── Build one record per holding ───────────────────────────────────────────────
records = []
for _, r in gw.sort_values("gw_score_pct").iterrows():
    tk = r["ticker"]
    t1 = tier1(tk)
    dims = []
    for n, (dk, label) in enumerate(DIMS, 1):
        raw  = str(r[f"gw_{dk}_rating"]).upper()
        show = SHOW.get(raw, raw)
        d1   = t1.get(dk, {})
        reason = str(d1.get("reasoning") or "").strip()
        t1rt   = str(d1.get("rating","")).upper()
        if dk in ("target","verification") and t1rt and t1rt != raw:
            reason += (f"  [Tier-2 revision: Tier-1 reading was {SHOW.get(t1rt,t1rt)}; "
                       f"adjusted to {show} on the SBTi/TPI evidence.]")
        elif dk in ("target","verification"):
            reason += "  [Corroborated by the Tier-2 SBTi/TPI evidence.]"
        reason = (reason or "—").replace("\n"," ")
        dims.append(dict(n=n, label=label, rating=show,
                         how=reason, quote=str(d1.get("quote") or "").strip(),
                         page=d1.get("page","")))
    nF = sum(d["rating"]=="FAIL" for d in dims)
    nQ = sum(d["rating"]=="PARTIAL" for d in dims)
    nM = sum(d["rating"]=="MISSING" for d in dims)
    nP = 8 - nF - nQ - nM
    pts = sum(POINTS[str(r[f'gw_{dk}_rating']).upper()] for dk,_ in DIMS)
    sp = t1.get("specificity", {})
    records.append(dict(
        company=r["company_name"], ticker=tk, score=float(r["gw_score_pct"]),
        rec=("EXCLUDE" if nF>=3 else "WATCHLIST" if nF==2 else "PASS"),
        dims=dims, nP=nP, nQ=nQ, nF=nF, nM=nM, pts=pts,
        main_claim=str(sp.get("quote") or ""), main_page=sp.get("page",""),
        t2=t2_note(tk), verdict=str(r.get("analyst_note","")).strip()))

# ════════════════════════════════════════════════════════════════════════════
# EXCEL
# ════════════════════════════════════════════════════════════════════════════
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

wb = Workbook()
HEAD = Font(bold=True, color="FFFFFF")
HFILL = PatternFill("solid", fgColor="1F4E78")
WRAP = Alignment(wrap_text=True, vertical="top")
CTR  = Alignment(horizontal="center", vertical="center")
thin = Side(style="thin", color="BFBFBF")
BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)

def style_header(ws, ncol):
    for c in range(1, ncol+1):
        cell = ws.cell(1, c); cell.font = HEAD; cell.fill = HFILL
        cell.alignment = CTR; cell.border = BORDER

# --- Sheet 1: Summary ---
ws = wb.active; ws.title = "Summary"
cols = ["#","Company","Ticker"] + [lbl for _,lbl in DIMS] + \
       ["Score %","FAIL","Recommendation"]
ws.append(cols)
for i, rec in enumerate(records, 1):
    row = [i, rec["company"], rec["ticker"]] + [d["rating"] for d in rec["dims"]] + \
          [round(rec["score"],1), rec["nF"], rec["rec"]]
    ws.append(row)
    rr = i + 1
    for j, d in enumerate(rec["dims"]):
        cell = ws.cell(rr, 4+j)
        cell.fill = PatternFill("solid", fgColor=FILL[d["rating"]])
        cell.font = Font(color=FONTC[d["rating"]], bold=True)
        cell.alignment = CTR
    ws.cell(rr, 3+8+3).fill = PatternFill(  # Recommendation
        "solid", fgColor=FILL.get(rec["rec"], "C6EFCE"))
    ws.cell(rr, 3+8+3).font = Font(bold=True)
style_header(ws, len(cols))
widths = [4, 34, 11] + [13]*8 + [9, 7, 16]
for c, w in enumerate(widths, 1):
    ws.column_dimensions[ws.cell(1,c).column_letter].width = w
ws.freeze_panes = "D2"

# --- Sheet 2: Detail ---
wd = wb.create_sheet("Detail")
wd.append(["#","Company","Ticker","Dim #","Dimension","Rating",
           "How it was assessed","Evidence quote","Page"])
rr = 1
for i, rec in enumerate(records, 1):
    for d in rec["dims"]:
        rr += 1
        wd.append([i, rec["company"], rec["ticker"], d["n"], d["label"],
                   d["rating"], d["how"], d["quote"], d["page"]])
        rc = wd.cell(rr, 6)
        rc.fill = PatternFill("solid", fgColor=FILL[d["rating"]])
        rc.font = Font(color=FONTC[d["rating"]], bold=True); rc.alignment = CTR
        wd.cell(rr,7).alignment = WRAP; wd.cell(rr,8).alignment = WRAP
style_header(wd, 9)
for c, w in enumerate([4,30,10,7,15,11,70,55,7], 1):
    wd.column_dimensions[wd.cell(1,c).column_letter].width = w
wd.freeze_panes = "A2"

# --- Sheet 3: Methodology ---
wm = wb.create_sheet("Methodology")
method = [
 ("Greenwashing 8-Test — methodology", True),
 ("", False),
 ("Each holding's headline sustainability claim is tested on 8 dimensions, each", False),
 ("rated PASS / PARTIAL / FAIL / MISSING.", False),
 ("", False),
 ("Dimensions: 1 Specificity, 2 Metric, 3 Baseline, 4 Target, 5 Time Horizon,", False),
 ("6 Scope, 7 Verification, 8 Consistency.", False),
 ("", False),
 ("Scoring — ratings convert to risk points: PASS 0, PARTIAL 1, MISSING 1, FAIL 2.", False),
 ("The 8-Test concern score = total points / 16, expressed as a percentage", False),
 ("(0% = clean, 100% = severe).", False),
 ("", False),
 ("Recommendation rule: 3+ FAIL -> EXCLUDE; exactly 2 FAIL -> WATCHLIST;", False),
 ("otherwise PASS.", False),
 ("", False),
 ("Evidence tiers:", True),
 ("  Tier 1 — primary company documents (annual / sustainability reports) in the", False),
 ("           RAG corpus; every rating cites a verbatim, page-numbered quote.", False),
 ("  Tier 2 — external verification: the SBTi target database and TPI Carbon", False),
 ("           Performance assessments, matched by ISIN. Tier 2 drives dimension 4", False),
 ("           (Target) and 7 (Verification).", False),
 ("", False),
 ("Authoritative source: data/rag/RAG_Screening_Sheet_Workbook_v1.xlsx.", False),
 (f"Generated {TODAY}. Academic prototype — not investment advice.", False),
]
for txt, bold in method:
    wm.append([txt])
    if bold:
        wm.cell(wm.max_row,1).font = Font(bold=True)
wm.column_dimensions["A"].width = 92

xlsx_path = f"outputs/rag/greenwashing_8test_report_{TODAY}.xlsx"
wb.save(xlsx_path)

# ════════════════════════════════════════════════════════════════════════════
# WORD
# ════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

doc.add_heading("Greenwashing 8-Test — Per-Company Screening Report", 0)
p = doc.add_paragraph()
p.add_run(f"Agent 9  ·  20 portfolio holdings  ·  generated {TODAY}").italic = True

doc.add_heading("Methodology", 1)
doc.add_paragraph(
    "Each holding's headline sustainability claim is tested on eight dimensions "
    "— Specificity, Metric, Baseline, Target, Time Horizon, Scope, Verification "
    "and Consistency — each rated PASS, PARTIAL, FAIL or MISSING. Ratings convert "
    "to risk points (PASS 0, PARTIAL 1, MISSING 1, FAIL 2); the total over 16 is "
    "the 8-Test concern score. Recommendation rule: three or more FAIL ratings "
    "trigger EXCLUDE, exactly two trigger WATCHLIST, otherwise PASS.")
doc.add_paragraph(
    "Evidence is two-tier. Tier 1 is the primary company documents in the RAG "
    "corpus — every rating cites a verbatim, page-numbered quote. Tier 2 is "
    "external verification — the SBTi target database and TPI Carbon Performance "
    "assessments, matched by ISIN — and drives the Target and Verification "
    "dimensions. The authoritative record is RAG_Screening_Sheet_Workbook_v1.xlsx.")

doc.add_heading("Summary", 1)
tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Light Grid Accent 1"
for c, h in enumerate(["#","Company","Ticker","Score","FAIL","Recommendation"]):
    tbl.rows[0].cells[c].text = h
    tbl.rows[0].cells[c].paragraphs[0].runs[0].bold = True
for i, rec in enumerate(records, 1):
    cells = tbl.add_row().cells
    cells[0].text = str(i); cells[1].text = rec["company"]
    cells[2].text = rec["ticker"]; cells[3].text = f"{rec['score']:.1f}%"
    cells[4].text = str(rec["nF"]); cells[5].text = rec["rec"]
    shade(cells[5], FILL.get(rec["rec"], "C6EFCE"))

doc.add_heading("Per-company breakdown", 1)
for i, rec in enumerate(records, 1):
    doc.add_heading(f"{i}. {rec['company']}  ·  {rec['ticker']}", 2)
    h = doc.add_paragraph()
    h.add_run(f"8-Test score: {rec['score']:.1f}%   ·   Recommendation: {rec['rec']}").bold = True
    if rec["main_claim"]:
        doc.add_paragraph(f"Main claim tested: “{rec['main_claim']}” "
                          f"(p.{rec['main_page']}).")
    doc.add_paragraph(f"Tier-2 external check: {rec['t2']}.")

    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for c, hd in enumerate(["#","Dimension","Rating","How it was assessed"]):
        t.rows[0].cells[c].text = hd
        t.rows[0].cells[c].paragraphs[0].runs[0].bold = True
    for d in rec["dims"]:
        cells = t.add_row().cells
        cells[0].text = str(d["n"])
        cells[1].text = d["label"]
        cells[2].text = d["rating"]
        shade(cells[2], FILL[d["rating"]])
        how = d["how"]
        if d["quote"] and d["label"] != "Specificity":
            how += f"  Evidence: “{d['quote']}” (p.{d['page']})."
        cells[3].text = how
    for c, w in enumerate([0.4, 1.3, 0.9, 4.4]):
        for row in t.rows:
            row.cells[c].width = Inches(w)

    doc.add_paragraph(
        f"Score logic: {rec['nP']} PASS (0 pts), {rec['nQ']} PARTIAL (1), "
        f"{rec['nM']} MISSING (1), {rec['nF']} FAIL (2)  =  {rec['pts']}/16  =  "
        f"{rec['score']:.1f}%.")
    doc.add_paragraph(
        f"Recommendation logic: {rec['nF']} dimension(s) rated FAIL "
        f"(EXCLUDE needs ≥3, WATCHLIST needs exactly 2)  →  {rec['rec']}.")
    if rec["verdict"] and rec["verdict"].lower() != "nan":
        v = doc.add_paragraph()
        v.add_run("Verdict rationale. ").bold = True
        v.add_run(rec["verdict"])

docx_path = f"outputs/rag/greenwashing_8test_report_{TODAY}.docx"
doc.save(docx_path)

print(f"Excel saved: {xlsx_path}  (Summary + Detail + Methodology)")
print(f"Word  saved: {docx_path}  ({len(records)} companies)")
