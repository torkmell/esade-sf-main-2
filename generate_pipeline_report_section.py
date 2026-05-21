"""
Generates the Pipeline Architecture section of the report.
Output: deliverables/Section_Pipeline_Architecture.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

os.makedirs("deliverables", exist_ok=True)

# ── Colours ────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
BLACK      = RGBColor(0x00, 0x00, 0x00)
GREY       = RGBColor(0x40, 0x40, 0x40)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)

def set_cell_border(cell, color="CCCCCC", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = MID_BLUE

def body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    return p

def body_with_bold(doc, parts, space_after=6):
    """parts = list of (text, bold) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GREY
        run.bold = bold
    return p

def bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = DARK_BLUE
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.color.rgb = GREY
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GREY

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1.0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    # Light background via shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF4FB')
    pPr.append(shd)

def table_2col(doc, headers, rows, col_widths=(4.5, 11.0)):
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    # Header row
    hdr = tbl.rows[0].cells
    for i, (cell, h) in enumerate(zip(hdr, headers)):
        set_cell_bg(cell, DARK_BLUE)
        set_cell_border(cell, color="1F497D")
        cell.width = Cm(col_widths[i])
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1].cells
        bg = LIGHT_BLUE if r_idx % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for i, (cell, val) in enumerate(zip(row, row_data)):
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            cell.width = Cm(col_widths[i])
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = GREY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Cover line ─────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run("ESADE MSc Finance  —  Sustainable Finance Group Assignment")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run.italic = True

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(18)
run2 = p2.add_run("Report Section: Pipeline Architecture")
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run2.italic = True

# ══════════════════════════════════════════════════════════════
heading1(doc, "Section: Pipeline Architecture")
# ══════════════════════════════════════════════════════════════

body(doc,
    "This section describes the technical architecture of the AI-agent research pipeline developed "
    "for the ESADE Sustainable Finance group assignment. The pipeline automates the full investment "
    "process from raw data ingestion through to portfolio construction and reporting, and was "
    "implemented entirely from scratch using Python and Visual Studio Code.")

# ── 1. Technology Choice ───────────────────────────────────────
heading2(doc, "1.  Technology Choice: Python and Visual Studio Code")

body(doc,
    "The pipeline was implemented using Python 3.11 executed within Visual Studio Code (VS Code), "
    "rather than a graphical no-code orchestration tool. This decision was made deliberately for "
    "three reasons.")

bullet(doc,
    "All 13 agents, the master runner, and the interactive dashboard are contained within a single "
    "project folder. There is no dependency on external cloud services, no drag-and-drop workflow "
    "editor, and no separate platform to log into. Every component is version-controlled and "
    "inspectable in one place.",
    bold_prefix="Single environment.")

bullet(doc,
    "Each agent was written from scratch as a Jupyter notebook, meaning the team has full visibility "
    "into every transformation applied to the data. There are no black-box connectors or pre-built "
    "modules — every scoring formula, exclusion rule, and weight is explicitly coded and auditable.",
    bold_prefix="Full auditability.")

bullet(doc,
    "The pipeline runs end-to-end in approximately 63 seconds with a single keyboard shortcut "
    "(Ctrl+Shift+B in VS Code). This makes live demonstration straightforward and eliminates "
    "dependency on internet connectivity or third-party service availability during the presentation.",
    bold_prefix="Reproducibility.")

doc.add_paragraph()
body_with_bold(doc, [
    ("The master runner script (", False),
    ("run_pipeline.py", True),
    (") acts as the orchestrator. It executes each of the 13 Jupyter notebooks in sequence using "
     "nbconvert, captures the exit code of each agent, and prints a structured summary report "
     "showing execution time and pass/fail status per agent. The pipeline handles missing files "
     "gracefully — agents whose notebook file is not found are marked SKIP rather than causing "
     "the entire run to abort.", False),
])

# ── 2. Agent Communication ─────────────────────────────────────
heading2(doc, "2.  Agent Communication: CSV File Passing")

body(doc,
    "Agents communicate exclusively through flat files saved to the outputs/ directory. Each agent "
    "reads the output of the preceding agent, performs its transformation, and writes a new dated "
    "CSV or JSON file. This design mirrors a data warehouse staging pattern and has two key "
    "properties: agents are stateless (they do not hold data in memory between runs), and any "
    "agent can be re-run individually without restarting the full pipeline.")

doc.add_paragraph()
body(doc, "The file-passing chain is as follows:")

file_chain = [
    ("Agent 01 — Mandate",               "outputs/scores/mandate.json",                         "Investment thesis, weights, exclusion rules"),
    ("Agent 02 — Data Ingestion",         "outputs/scores/master_dataset_DATE.csv",              "57 companies x 84 columns, vintage-tagged"),
    ("Agent 03 — Data Quality",           "outputs/scores/data_dictionary_DATE.csv",             "Missing-value audit, outlier flags"),
    ("Agent 04 — Document Intelligence",  "outputs/rag/doc_intel_TICKER.json",                   "RAG extractions from sustainability PDFs"),
    ("Agent 05/06 — ESG & Climate",       "outputs/scores/esg_scores_DATE.csv",                  "E, S, G and composite ESG scores (0-100), WACI"),
    ("Agent 07 — Biodiversity",           "outputs/scores/biodiversity_scores_DATE.csv",         "Nature risk score (ENCORE + WRI Aqueduct)"),
    ("Agent 08 — EU Regulation",          "outputs/scores/eu_regulation_DATE.csv",               "SFDR Article 8 compliance, 14 PAI indicators"),
    ("Agent 09 — Greenwashing",           "outputs/rag/greenwash_TICKER.json",                   "8-Test ratings per company, watchlist/exclusions"),
    ("Agent 10 — Financial Analysis",     "outputs/scores/financial_metrics_DATE.csv",           "Return, volatility, Sharpe ratio, max drawdown"),
    ("Agent 11 — Portfolio Construction", "outputs/portfolio/final_portfolio_DATE.csv",          "20 holdings with weights; exclusions.csv"),
    ("Agent 12 — Human Review",           "outputs/scores/human_overrides_DATE.csv",             "Override log and AI Use Statement"),
    ("Agent 13 — Reporting",              "outputs/reports/*.png",                               "Portfolio factsheet charts"),
]

table_2col(doc,
    headers=["Agent", "Output File", "Contents"],
    rows=[(a, f, c) for a, f, c in file_chain],
    col_widths=(4.5, 5.5, 6.0)
)

# ── 3. Execution Model ─────────────────────────────────────────
heading2(doc, "3.  Execution Model")

body(doc,
    "The pipeline is invoked via a single VS Code build task. Internally, run_pipeline.py calls "
    "nbconvert for each notebook in sequence:")

code_block(doc,
    "python -m jupyter nbconvert --to notebook --execute --inplace\n"
    "       --ExecutePreprocessor.timeout=300\n"
    "       --ExecutePreprocessor.kernel_name=sustainable-finance\n"
    "       notebooks/agentXX_name.ipynb")

body(doc,
    "Each notebook is executed in-place, meaning outputs are written back into the .ipynb file "
    "itself and are inspectable in Jupyter after the run. The timeout is set to 300 seconds per "
    "agent. In practice, the full 13-agent pipeline completes in approximately 63 seconds on "
    "standard hardware.")

body_with_bold(doc, [
    ("Results are reported as ", False),
    ("[OK]", True),
    (" or ", False),
    ("[!!]", True),
    (" per agent, with elapsed time and the last five lines of error output printed for any "
     "failures. This makes debugging straightforward — a data engineer can identify the failing "
     "agent, open the relevant notebook, and correct the issue without touching any other part "
     "of the pipeline.", False),
])

# ── 4. Data Replacement ────────────────────────────────────────
heading2(doc, "4.  Data Source Replacement")

body(doc,
    "The pipeline was developed and tested against a synthetic mock dataset that mirrors the "
    "structure and column naming of the four course-provided CSV files. When the real data files "
    "are received, replacement requires three steps:")

bullet(doc, "Drop the four real CSV files into the data/provided/ directory, replacing the mock files.")
bullet(doc, "Run the full pipeline once (Ctrl+Shift+B). All agents will automatically recalculate "
            "using the real data.")
bullet(doc, "Refresh the Streamlit dashboard browser tab. All charts and tables update instantly.")

body(doc,
    "No code changes are required. The pipeline reads all input files by column name, not by "
    "position, so the real files will integrate correctly provided the professor's column naming "
    "is consistent with the data dictionary established in Agent 03.")

# ── 5. Dashboard ───────────────────────────────────────────────
heading2(doc, "5.  Interactive Dashboard")

body(doc,
    "Agent outputs are visualised through a Streamlit web dashboard (app.py) that reads directly "
    "from the outputs/ directory. The dashboard provides eight interactive views: portfolio "
    "overview, holdings detail, ESG scores across the full universe, risk and return analytics, "
    "climate and biodiversity metrics, greenwashing assessment, exclusion rationale, and the "
    "investment mandate summary.")

body_with_bold(doc, [
    ("The dashboard is launched with a single command (", False),
    ("venv/Scripts/streamlit.exe run app.py --server.headless=true", True),
    (") and is accessible at http://localhost:8501. Because it reads the latest file in each "
     "outputs/ subfolder at startup, it always reflects the most recent pipeline run without "
     "any manual refresh of the underlying data.", False),
])

# ── 6. Comparison to n8n ───────────────────────────────────────
heading2(doc, "6.  Comparison to n8n")

body(doc,
    "The original assignment specification suggested n8n.cloud as the orchestration layer. "
    "After evaluation, the team opted for the Python-native implementation described above. "
    "The table below summarises the trade-offs considered.")

table_2col(doc,
    headers=["Dimension", "n8n.cloud", "Python + VS Code (chosen)"],
    rows=[
        ("Setup",           "Requires cloud account, browser-based workflow editor",   "Single Python file, runs locally"),
        ("Auditability",    "Connectors are black-box; logic not inspectable",         "Every line of code visible and editable"),
        ("Internet dependency", "Requires live connection to n8n.cloud during demo",  "Fully offline, no external services"),
        ("Agent logic",     "Pre-built nodes; custom logic requires JavaScript",       "Full Python — any library, any formula"),
        ("Demo reliability","Service outage or auth failure could break live demo",    "Ctrl+Shift+B — always works"),
        ("Academic value",  "Demonstrates tool usage",                                 "Demonstrates pipeline engineering from scratch"),
    ],
    col_widths=(4.0, 6.5, 6.0)
)

body(doc,
    "The Python implementation provides a stronger technical demonstration for the Q&A defence, "
    "as the team can explain every architectural decision — from how agents pass data to each "
    "other, to how exclusion rules are enforced at the portfolio construction stage — without "
    "relying on a third-party platform's internal behaviour.")

# ── 7. Reproducibility Statement ──────────────────────────────
heading2(doc, "7.  Reproducibility Statement")

body(doc,
    "The pipeline is fully reproducible. All dependencies are listed in requirements.txt and "
    "installable via a single setup.bat script. The virtual environment (venv/) pins exact "
    "package versions. Any team member with Python 3.11 installed can clone the project folder, "
    "run setup.bat once, and execute the full pipeline without additional configuration.")

bullet(doc, "Python 3.11.9")
bullet(doc, "pandas 2.x, numpy, matplotlib, plotly — data processing and visualisation")
bullet(doc, "yfinance 1.3.0 — market price downloads")
bullet(doc, "jupyter, nbconvert — notebook execution engine")
bullet(doc, "streamlit 1.57.0 — interactive dashboard")
bullet(doc, "python-docx — Word document generation")

doc.add_paragraph()
body(doc,
    "All random seeds are fixed where applicable. Mock data is generated deterministically "
    "from generate_mock_data.py, ensuring that test runs produce identical outputs across "
    "machines and sessions.")

# ── Save ───────────────────────────────────────────────────────
out = "deliverables/Section_Pipeline_Architecture.docx"
doc.save(out)
print(f"Saved: {out}")
